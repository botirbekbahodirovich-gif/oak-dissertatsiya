"""Dissertation workspace — advisor↔student yozish/ko'rib chiqish moduli.

NOM O'ZGARISHI: loyihalar jadvali `diss_projects` (spec'dagi "dissertations"
nomi bu bazada OAK himoya korpusi bilan band). Sahifalar /workspace ostida
(/dissertation/<int:id> ham OAK yozuvi sahifasi sifatida band); API yo'llari
spec bo'yicha /api/dissertation/*, /api/blocks/*, /api/advisor/*.

Xavfsizlik invariantlari (har endpointda):
  - login_required + egalik tekshiruvi: faqat egasi, yoki qabul qilingan
    advisor_links + diss_projects.advisor_id mos kelgan rahbar. Adminlar HAM
    kira olmaydi (intellektual mulk). Aks holda 403.
  - Barcha SQL parametrlangan (%s).
  - Yozuvchi endpointlar POST + csrf.exempt (kodbaza patterni: sessiya auth,
    JSON fetch) yoki forma+CSRF.
  - Blok saqlash: server tomonda bleach bilan sanitizatsiya, 5s rate-limit
    (per-block, in-memory — worker boshiga; spec ruxsat bergan soddalik),
    3 daqiqalik soft-lock, 20 talik versiya retensiyasi.
  - Fayllar lokal static/uploads ostida saqlanadi (bu kodbazaning mavjud
    patterni — cabinet avatarlari, universitet galereyasi shu yo'lda; .env da
    SUPABASE_KEY yo'q, mavjud Supabase faqat read-only avatar CDN). Saqlash
    _store_upload() ichida izolyatsiya qilingan — keyin Supabase'ga
    almashtirish bitta funksiya.
"""
import json
import os
import re
import secrets
import time
import uuid
import html as html_mod
from datetime import datetime, timezone

from flask import (Blueprint, jsonify, request, render_template, redirect,
                   abort, send_file, current_app, url_for)
from flask_login import login_required, current_user

from app import csrf

dissertation_bp = Blueprint('dissertation', __name__)

_schema_ready = False


@dissertation_bp.after_request
def _no_store(response):
    """Muharrir sahifalari va API — shaxsiy kontent, hech qachon keshlanmasin:
    eski keshdan qaytgan HTML/JSON foydalanuvchiga "matnim yo'qoldi" bo'lib
    ko'rinadi (ED_STATE.content sahifaga bikilgan)."""
    response.headers['Cache-Control'] = 'no-store'
    return response


_diss_js_version = None


@dissertation_bp.context_processor
def _inject_asset_version():
    """dissertation.js uchun mtime asosidagi cache-buster. nginx /static/ ni
    7 kun (Flask 1 kun) keshlaydi — versiyasiz URL'da tuzatilgan JS
    foydalanuvchi brauzerlariga bir haftagacha yetib bormaydi (b686522
    tuzatishi shu sabab "ishlamay" ko'ringan)."""
    global _diss_js_version
    if _diss_js_version is None:
        try:
            p = os.path.join(current_app.static_folder, 'js', 'dissertation.js')
            _diss_js_version = str(int(os.stat(p).st_mtime))
        except OSError:
            _diss_js_version = '1'
    return {'diss_js_v': _diss_js_version}

MAX_DEPTH = 3
MAX_CONTENT_BYTES = 2 * 1024 * 1024   # 2MB per save
MAX_TITLE = 500
MAX_ANNOT_BODY = 5000
LOCK_TTL = 180                        # soniya — shu vaqtdan eski lock "o'lik"
SAVE_MIN_INTERVAL = 5                 # soniya — autosave rate limit
VERSIONS_KEPT = 20
IMG_MAX_BYTES = 5 * 1024 * 1024       # 5MB (free tarif)

DEGREE_LABELS = {'magistr': 'Magistrlik', 'phd': 'PhD', 'dsc': 'DSc'}
STATUS_LABELS = {'draft': 'Qoralama', 'in_review': "Ko'rib chiqilmoqda",
                 'revision': 'Tuzatish', 'approved': "Ma'qullangan",
                 'archived': 'Arxiv'}
REVIEW_LABELS = {'not_reviewed': "Ko'rilmagan", 'deficiencies': 'Kamchiliklar bor',
                 'task_assigned': 'Topshiriq berildi', 'approved': "Ma'qullandi"}
ANNOT_TYPES = {'comment': 'Izoh', 'correction': 'Tuzatish', 'task': 'Topshiriq'}

# OAK tuzilmasida RAQAMLANMAYDIGAN maxsus bo'limlar (normallashtirilgan,
# to'g'ri apostrof bilan). Sarlavha shu ro'yxatga tushsa block_type='special'.
_SPECIAL_TITLES = {
    'kirish', 'annotatsiya', 'xulosa', 'umumiy xulosa',
    'xulosa, taklif va tavsiyalar',
    'foydalanilgan adabiyotlar', "foydalanilgan adabiyotlar ro'yxati",
    "adabiyotlar ro'yxati", 'ilova', 'ilovalar', 'mundarija',
    'qisqartmalar', 'shartli belgilar',
}

# har xil apostroflarni to'g'ri ' ga keltirish (Kirill/lotin/curly variantlari)
_APOS_TRANSLATE = {ord(c): "'" for c in "`ʻʼ‘’"}

# "I BOB", "II-bob", "1 bob", "1-BOB" kabi sarlavhalar allaqachon o'z raqamiga
# ega — displayда ikki karra raqamlamaslik uchun aniqlanadi.
_BOB_LABEL_RE = re.compile(r'^\s*(?:[ivxlcdm]+|\d+)\s*[-.–\s]*bob\b', re.I)


def _norm_title(title):
    return (title or '').strip().lower().translate(_APOS_TRANSLATE)


def _is_special_title(title):
    return _norm_title(title) in _SPECIAL_TITLES


def _heading_label(numbering, title):
    """Ko'rsatiladigan sarlavha: raqam bo'lsa "1.2. Nom", lekin sarlavha
    allaqachon 'I BOB' bilan boshlansa yoki raqam yo'q bo'lsa — faqat nom."""
    title = title or ''
    if numbering and not _BOB_LABEL_RE.match(title):
        return f'{numbering}. {title}'
    return title


# per-worker rate-limit xotirasi: {block_id: last_save_ts}, {user_id: export_ts}
_last_save = {}
_last_export = {}


# ── schema (lazy, idempotent — migrations/add_dissertation_module.sql aksi) ──

def _ensure_schema(cur):
    global _schema_ready
    if _schema_ready:
        return
    cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS is_premium BOOLEAN DEFAULT FALSE")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS advisor_links (
            id SERIAL PRIMARY KEY,
            student_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
            advisor_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
            status VARCHAR(20) NOT NULL DEFAULT 'pending'
                CHECK (status IN ('pending', 'accepted', 'declined', 'removed')),
            invited_by INTEGER REFERENCES users(id),
            invite_message TEXT,
            created_at TIMESTAMP DEFAULT NOW(),
            responded_at TIMESTAMP,
            UNIQUE(student_id, advisor_id),
            CHECK (student_id <> advisor_id)
        )""")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_advisor_links_student "
                "ON advisor_links(student_id) WHERE status = 'accepted'")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_advisor_links_advisor "
                "ON advisor_links(advisor_id) WHERE status = 'accepted'")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS diss_projects (
            id SERIAL PRIMARY KEY,
            owner_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
            title VARCHAR(600) NOT NULL,
            degree_type VARCHAR(30) DEFAULT 'phd'
                CHECK (degree_type IN ('magistr', 'phd', 'dsc')),
            specialty_code VARCHAR(30),
            language VARCHAR(10) DEFAULT 'uz',
            status VARCHAR(30) DEFAULT 'draft'
                CHECK (status IN ('draft', 'in_review', 'revision', 'approved', 'archived')),
            advisor_id INTEGER REFERENCES users(id),
            last_submitted_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT NOW(),
            updated_at TIMESTAMP DEFAULT NOW()
        )""")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_diss_projects_owner ON diss_projects(owner_id)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_diss_projects_advisor ON diss_projects(advisor_id)")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS dissertation_blocks (
            id SERIAL PRIMARY KEY,
            dissertation_id INTEGER NOT NULL REFERENCES diss_projects(id) ON DELETE CASCADE,
            parent_id INTEGER REFERENCES dissertation_blocks(id) ON DELETE CASCADE,
            title VARCHAR(500) NOT NULL,
            numbering VARCHAR(50),
            sort_order INTEGER NOT NULL DEFAULT 0,
            depth INTEGER NOT NULL DEFAULT 0,
            content TEXT DEFAULT '',
            content_plain TEXT DEFAULT '',
            word_count INTEGER DEFAULT 0,
            review_status VARCHAR(30) DEFAULT 'not_reviewed'
                CHECK (review_status IN ('not_reviewed', 'deficiencies', 'task_assigned', 'approved')),
            review_status_by INTEGER REFERENCES users(id),
            review_status_at TIMESTAMP,
            is_locked_by INTEGER REFERENCES users(id),
            locked_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT NOW(),
            updated_at TIMESTAMP DEFAULT NOW(),
            CHECK (depth >= 0 AND depth <= 3)
        )""")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_blocks_dissertation "
                "ON dissertation_blocks(dissertation_id, parent_id, sort_order)")
    # block_type: 'special' bo'limlar (Kirish, Xulosa, Adabiyotlar...) RAQAMLANMAYDI
    cur.execute("ALTER TABLE dissertation_blocks "
                "ADD COLUMN IF NOT EXISTS block_type VARCHAR(20) DEFAULT 'chapter'")
    # word_target: bo'lim uchun so'z maqsadi (Akademik Reja integratsiyasi —
    # blueprints/roadmap.py; wizard OAK skeletida to'ldiradi, muharrir ko'rsatadi)
    cur.execute("ALTER TABLE dissertation_blocks "
                "ADD COLUMN IF NOT EXISTS word_target INTEGER")
    cur.execute("""
        DO $$ BEGIN
          IF NOT EXISTS (SELECT 1 FROM pg_constraint
                         WHERE conname = 'dissertation_blocks_block_type_chk') THEN
            ALTER TABLE dissertation_blocks
              ADD CONSTRAINT dissertation_blocks_block_type_chk
              CHECK (block_type IN ('chapter', 'special'));
          END IF;
        END $$;""")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS block_versions (
            id SERIAL PRIMARY KEY,
            block_id INTEGER NOT NULL REFERENCES dissertation_blocks(id) ON DELETE CASCADE,
            content TEXT NOT NULL,
            word_count INTEGER DEFAULT 0,
            saved_by INTEGER REFERENCES users(id),
            save_type VARCHAR(20) DEFAULT 'manual'
                CHECK (save_type IN ('manual', 'autosave', 'pre_restore')),
            created_at TIMESTAMP DEFAULT NOW()
        )""")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_block_versions_block "
                "ON block_versions(block_id, created_at DESC)")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS block_annotations (
            id SERIAL PRIMARY KEY,
            block_id INTEGER NOT NULL REFERENCES dissertation_blocks(id) ON DELETE CASCADE,
            author_id INTEGER NOT NULL REFERENCES users(id),
            annotation_type VARCHAR(20) DEFAULT 'comment'
                CHECK (annotation_type IN ('comment', 'correction', 'task')),
            anchor_text TEXT NOT NULL,
            anchor_prefix VARCHAR(100),
            anchor_suffix VARCHAR(100),
            anchor_offset INTEGER,
            body TEXT NOT NULL,
            status VARCHAR(20) DEFAULT 'open'
                CHECK (status IN ('open', 'resolved', 'orphaned')),
            resolved_by INTEGER REFERENCES users(id),
            resolved_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT NOW()
        )""")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_annotations_block "
                "ON block_annotations(block_id, status)")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS annotation_replies (
            id SERIAL PRIMARY KEY,
            annotation_id INTEGER NOT NULL REFERENCES block_annotations(id) ON DELETE CASCADE,
            author_id INTEGER NOT NULL REFERENCES users(id),
            body TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT NOW()
        )""")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_annotation_replies "
                "ON annotation_replies(annotation_id)")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS diss_notifications (
            id SERIAL PRIMARY KEY,
            user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
            event_type VARCHAR(40) NOT NULL,
            dissertation_id INTEGER REFERENCES diss_projects(id) ON DELETE CASCADE,
            block_id INTEGER REFERENCES dissertation_blocks(id) ON DELETE CASCADE,
            actor_id INTEGER REFERENCES users(id),
            payload JSONB DEFAULT '{}',
            is_read BOOLEAN DEFAULT FALSE,
            created_at TIMESTAMP DEFAULT NOW()
        )""")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_diss_notif_user "
                "ON diss_notifications(user_id, is_read, created_at DESC)")
    # AI copilot placeholder — v1 da ishlatilmaydi, v2 shu jadvalga ulanadi.
    cur.execute("""
        CREATE TABLE IF NOT EXISTS ai_review_requests (
            id SERIAL PRIMARY KEY,
            block_id INTEGER REFERENCES dissertation_blocks(id) ON DELETE CASCADE,
            user_id INTEGER REFERENCES users(id),
            request_type VARCHAR(30),
            status VARCHAR(20) DEFAULT 'pending',
            result JSONB,
            created_at TIMESTAMP DEFAULT NOW()
        )""")
    # havola orqali taklif (advisor / student / collaborator) — bir martalik token
    cur.execute("""
        CREATE TABLE IF NOT EXISTS advisor_invite_tokens (
            id SERIAL PRIMARY KEY,
            token VARCHAR(64) UNIQUE NOT NULL,
            created_by INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
            role VARCHAR(15) NOT NULL DEFAULT 'advisor'
                CHECK (role IN ('advisor', 'student', 'collaborator')),
            dissertation_id INTEGER REFERENCES diss_projects(id) ON DELETE CASCADE,
            can_comment BOOLEAN DEFAULT TRUE,
            can_edit BOOLEAN DEFAULT FALSE,
            can_review_status BOOLEAN DEFAULT FALSE,
            expires_at TIMESTAMP DEFAULT (NOW() + INTERVAL '7 days'),
            used_by INTEGER REFERENCES users(id),
            used_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT NOW()
        )""")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_invite_tokens_token "
                "ON advisor_invite_tokens(token)")
    # qo'shimcha hamkorlar (loyihaga max 3) — granular ruxsatlar bilan
    cur.execute("""
        CREATE TABLE IF NOT EXISTS diss_collaborators (
            id SERIAL PRIMARY KEY,
            dissertation_id INTEGER NOT NULL REFERENCES diss_projects(id) ON DELETE CASCADE,
            user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
            invited_by INTEGER REFERENCES users(id),
            status VARCHAR(20) DEFAULT 'pending'
                CHECK (status IN ('pending', 'accepted', 'declined', 'removed')),
            can_comment BOOLEAN DEFAULT TRUE,
            can_edit BOOLEAN DEFAULT FALSE,
            can_review_status BOOLEAN DEFAULT FALSE,
            created_at TIMESTAMP DEFAULT NOW(),
            UNIQUE(dissertation_id, user_id)
        )""")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_diss_collab_user "
                "ON diss_collaborators(user_id) WHERE status = 'accepted'")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_diss_collab_diss "
                "ON diss_collaborators(dissertation_id)")
    _migrate_block_types(cur)
    _repair_highlight_pollution(cur)
    _schema_ready = True


def _repair_highlight_pollution(cur):
    """Bir martalik tozalash (worker boshiga): eski JS annotatsiya highlight
    spanlarini kontent bilan birga saqlab yuborgan — ularni zararsiz <span>ga
    aylantiramiz (bleach data-annotation-id ni allaqachon olib tashlagan,
    shuning uchun saqlangan shakl aynan shu satr). Xato bo'lsa initni buzmaydi."""
    try:
        cur.execute("""
            UPDATE dissertation_blocks
            SET content = replace(content, '<span class="annotation-highlight">', '<span>')
            WHERE content LIKE '%annotation-highlight%'
        """)
    except Exception:
        pass


def _migrate_block_types(cur):
    """Bir martalik migratsiya (worker boshiga): mavjud bloklarning sarlavhasi
    special ro'yxatga tushsa block_type='special' qilinadi va shu loyihalarning
    raqamlashi qayta hisoblanadi. Xato bo'lsa sxema initini buzmaydi."""
    try:
        cur.execute("""
            UPDATE dissertation_blocks SET block_type = 'special'
            WHERE block_type <> 'special'
              AND translate(lower(btrim(title)), $$`ʻʼ‘’$$, $$'''''$$) = ANY(%s)
            RETURNING dissertation_id
        """, (list(_SPECIAL_TITLES),))
        affected = {r[0] for r in cur.fetchall()}
        for diss_id in affected:
            _recompute_numbering(cur, diss_id)
    except Exception:
        pass


# ── umumiy yordamchilar ──────────────────────────────────────────────────────

def _conn():
    from data import get_connection
    return get_connection()


def _notify(cur, user_id, event_type, diss_id=None, block_id=None, payload=None):
    cur.execute("""
        INSERT INTO diss_notifications
            (user_id, event_type, dissertation_id, block_id, actor_id, payload)
        VALUES (%s, %s, %s, %s, %s, %s)
    """, (user_id, event_type, diss_id, block_id, current_user.id,
          json.dumps(payload or {}, ensure_ascii=False)))


def _usernames(cur, ids):
    ids = [i for i in set(ids) if i]
    if not ids:
        return {}
    cur.execute("SELECT id, username FROM users WHERE id = ANY(%s)", (ids,))
    return {r[0]: r[1] or f'user{r[0]}' for r in cur.fetchall()}


_PROJECT_COL_NAMES = ('id', 'owner_id', 'title', 'degree_type', 'specialty_code',
                      'language', 'status', 'advisor_id', 'last_submitted_at',
                      'created_at', 'updated_at')


def _project_cols(alias=''):
    """SELECT uchun ustunlar ro'yxati. JOIN'li so'rovlarda ustun nomi bir necha
    jadvalda uchrasa "ambiguous column" xatosi bo'ladi — shuning uchun alias
    berib prefikslaymiz, masalan _project_cols('d') → 'd.id, d.owner_id, ...'."""
    prefix = f'{alias}.' if alias else ''
    return ', '.join(prefix + c for c in _PROJECT_COL_NAMES)


def _project_dict(row):
    d = dict(zip(_PROJECT_COL_NAMES, row))
    d['degree_label'] = DEGREE_LABELS.get(d['degree_type'], d['degree_type'])
    d['status_label'] = STATUS_LABELS.get(d['status'], d['status'])
    return d


def _fetch_project(cur, diss_id):
    cur.execute(f"SELECT {_project_cols('d')} FROM diss_projects d WHERE d.id = %s",
                (diss_id,))
    row = cur.fetchone()
    if not row:
        abort(404)
    return _project_dict(row)


def _caps(role, view, edit, structure, comment, review, manage):
    """Chaqiruvchi qobiliyatlar to'plami."""
    return {'role': role, 'can_view': view, 'can_edit': edit,
            'can_structure': structure, 'can_comment': comment,
            'can_review_status': review, 'can_manage': manage}


def _collaborator_caps(cur, diss_id, uid):
    """Qabul qilingan hamkorning ruxsatlari yoki None."""
    cur.execute("""SELECT can_comment, can_edit, can_review_status
                   FROM diss_collaborators
                   WHERE dissertation_id = %s AND user_id = %s AND status = 'accepted'""",
                (diss_id, uid))
    r = cur.fetchone()
    if not r:
        return None
    return {'can_comment': bool(r[0]), 'can_edit': bool(r[1]),
            'can_review_status': bool(r[2])}


def get_access(cur, diss_id):
    """(project, caps) — chaqiruvchining loyihaga qobiliyatlari. Kirish
    umuman yo'q bo'lsa 403 (loyiha topilmasa 404). Adminlar uchun istisno
    YO'Q: dissertatsiya intellektual mulk.
      - owner: hamma narsa
      - advisor: ko'rish + izoh + qism holati (tuzilma/tahrir/boshqaruv YO'Q)
      - collaborator: ko'rish + bayroqlarga qarab izoh/tahrir/holat"""
    p = _fetch_project(cur, diss_id)
    uid = current_user.id
    if p['owner_id'] == uid:
        # ega: tuzilma/tahrir/izoh/boshqaruv HAMMASI; lekin qism "holati" —
        # ko'ruvchining (rahbar/holat-huquqli hamkor) ishi, o'zini o'zi
        # ma'qullamasligi uchun can_review_status = False.
        return p, _caps('owner', True, True, True, True, False, True)
    if p['advisor_id'] == uid:
        cur.execute("""SELECT 1 FROM advisor_links
                       WHERE student_id = %s AND advisor_id = %s AND status = 'accepted'""",
                    (p['owner_id'], uid))
        if cur.fetchone():
            return p, _caps('advisor', True, False, False, True, True, False)
    c = _collaborator_caps(cur, diss_id, uid)
    if c is not None:
        return p, _caps('collaborator', True, c['can_edit'], False,
                        c['can_comment'], c['can_review_status'], False)
    abort(403)


def get_dissertation_or_403(cur, diss_id, allow_advisor=True):
    """Orqaga moslik: (project, role). allow_advisor=False => faqat EGA
    (tuzilma/boshqaruv endpointlari uchun). allow_advisor=True => ega/rahbar/
    hamkor (ko'rish darajasi). Tahrir/izoh/holat kabi nozik ruxsatlar uchun
    endpointlar get_access/get_block_access + caps'dan foydalanadi."""
    p, caps = get_access(cur, diss_id)
    if caps['role'] == 'owner':
        return p, 'owner'
    if not allow_advisor:
        abort(403)          # faqat ega ruxsat etilgan endpoint
    return p, caps['role']  # 'advisor' yoki 'collaborator' (ko'rish)


def _fetch_block(cur, block_id):
    cur.execute("""
        SELECT id, dissertation_id, parent_id, title, numbering, sort_order,
               depth, content, content_plain, word_count, review_status,
               is_locked_by, locked_at, updated_at, block_type, word_target
        FROM dissertation_blocks WHERE id = %s
    """, (block_id,))
    row = cur.fetchone()
    if not row:
        abort(404)
    cols = ('id', 'dissertation_id', 'parent_id', 'title', 'numbering',
            'sort_order', 'depth', 'content', 'content_plain', 'word_count',
            'review_status', 'is_locked_by', 'locked_at', 'updated_at',
            'block_type', 'word_target')
    block = dict(zip(cols, row))
    block['is_special'] = (block.get('block_type') or 'chapter') == 'special'
    block['heading'] = _heading_label(block.get('numbering') or '', block.get('title'))
    return block


def get_block_or_403(cur, block_id, allow_advisor=True):
    """Blok + loyiha + rol; egalik loyiha orqali tekshiriladi."""
    block = _fetch_block(cur, block_id)
    project, role = get_dissertation_or_403(cur, block['dissertation_id'],
                                            allow_advisor=allow_advisor)
    return block, project, role


def get_block_access(cur, block_id):
    """Blok + loyiha + caps (nozik ruxsatlar: tahrir/izoh/holat uchun)."""
    block = _fetch_block(cur, block_id)
    project, caps = get_access(cur, block['dissertation_id'])
    return block, project, caps


# ── HTML sanitizatsiya + plain text ──────────────────────────────────────────

_ALLOWED_TAGS = ['p', 'h1', 'h2', 'h3', 'h4', 'b', 'strong', 'i', 'em', 'u',
                 's', 'ul', 'ol', 'li', 'table', 'thead', 'tbody', 'tr', 'td',
                 'th', 'img', 'a', 'blockquote', 'sub', 'sup', 'br', 'span',
                 'pre', 'code']
# style atributi butunlay olib tashlanadi (CSS sanitizer dependensiyasisiz
# xavfsizlik); Quill hizalamani class orqali beradi (ql-align-*), class xavfsiz.


def _class_attr_ok(tag, name, value):
    """class ruxsat, LEKIN vizual annotation-highlight spanlari kontent emas —
    ular bazaga yozilsa, mijozda "o'zgardi" ko'rinib autosave halqasi va
    kontent ifloslanishini keltiradi (JS ham saqlashdan oldin ularni olib
    tashlaydi; bu server tomondagi ikkinchi to'siq)."""
    return name == 'class' and 'annotation-highlight' not in value


_ALLOWED_ATTRS = {
    '*': _class_attr_ok,
    'img': ['src', 'alt', 'width', 'height'],
    'a': ['href', 'title', 'target', 'rel'],
    'td': ['colspan', 'rowspan'], 'th': ['colspan', 'rowspan'],
}


def sanitize_html(raw):
    import bleach
    return bleach.clean(raw or '', tags=_ALLOWED_TAGS, attributes=_ALLOWED_ATTRS,
                        protocols=['http', 'https', 'mailto'], strip=True)


def html_to_plain(html_str):
    """Anchoring/qidiruv uchun barqaror plain text: teglar → bo'sh joy,
    entity'lar ochiladi, whitespace yig'iladi."""
    txt = re.sub(r'<[^>]+>', ' ', html_str or '')
    txt = html_mod.unescape(txt)
    return re.sub(r'\s+', ' ', txt).strip()


# ── TOC raqamlash ────────────────────────────────────────────────────────────

def _fetch_tree(cur, diss_id):
    """Bitta so'rov — daraxt Python'da quriladi (N+1 yo'q)."""
    cur.execute("""
        SELECT b.id, b.parent_id, b.title, b.numbering, b.sort_order, b.depth,
               b.word_count, b.review_status, b.block_type,
               (SELECT COUNT(*) FROM block_annotations a
                 WHERE a.block_id = b.id AND a.status = 'open')
        FROM dissertation_blocks b
        WHERE b.dissertation_id = %s
        ORDER BY b.sort_order, b.id
    """, (diss_id,))
    nodes = {}
    for r in cur.fetchall():
        bt = r[8] or 'chapter'
        nodes[r[0]] = {'id': r[0], 'parent_id': r[1], 'title': r[2],
                       'numbering': r[3] or '', 'sort_order': r[4], 'depth': r[5],
                       'word_count': r[6] or 0, 'review_status': r[7],
                       'block_type': bt, 'is_special': bt == 'special',
                       'heading': _heading_label(r[3] or '', r[2]),
                       'open_annotations_count': r[9] or 0, 'children': []}
    roots = []
    for n in nodes.values():
        if n['parent_id'] and n['parent_id'] in nodes:
            nodes[n['parent_id']]['children'].append(n)
        else:
            roots.append(n)
    def sort_rec(lst):
        lst.sort(key=lambda x: (x['sort_order'], x['id']))
        for c in lst:
            sort_rec(c['children'])
    sort_rec(roots)
    return roots


def _recompute_numbering(cur, diss_id):
    """Tuzilma o'zgargach: DFS bilan "1", "1.1", "1.1.1" raqamlarini qayta
    hisoblab, faqat o'zgarganlarini yangilaydi (bitta batched pass).
    'special' bloklar (Kirish, Xulosa...) RAQAMLANMAYDI (numbering = NULL) va
    sanoqqa kirmaydi; special blokning bolalari ham raqamlanmaydi."""
    tree = _fetch_tree(cur, diss_id)
    updates = []
    def walk(items, prefix, unnumbered):
        i = 0
        for n in items:
            if unnumbered or n['is_special']:
                if n['numbering']:                    # eski raqamni tozalash
                    updates.append((None, n['id']))
                walk(n['children'], '', True)
            else:
                i += 1
                num = f'{prefix}.{i}' if prefix else f'{i}'
                if n['numbering'] != num:
                    updates.append((num, n['id']))
                walk(n['children'], num, False)
    walk(tree, '', False)
    for num, bid in updates:
        cur.execute("UPDATE dissertation_blocks SET numbering = %s WHERE id = %s",
                    (num, bid))


# ── annotatsiyalarni qayta bog'lash (re-anchor) ──────────────────────────────

def _reanchor_annotations(cur, block_id, new_plain):
    """Kontent o'zgargach: har ochiq annotatsiyaning anchor_text'i yangi
    matnda qidiriladi. Topilsa offset yangilanadi, topilmasa 'orphaned'."""
    cur.execute("""
        SELECT id, anchor_text FROM block_annotations
        WHERE block_id = %s AND status = 'open'
    """, (block_id,))
    for aid, anchor in cur.fetchall():
        pos = new_plain.find(anchor or '')
        if pos >= 0:
            cur.execute("UPDATE block_annotations SET anchor_offset = %s WHERE id = %s",
                        (pos, aid))
        else:
            cur.execute("UPDATE block_annotations SET status = 'orphaned' WHERE id = %s",
                        (aid,))


# ── fayl saqlash (lokal; Supabase'ga o'tish shu bitta funksiyada) ────────────

def _store_upload(file_storage, subdir, filename):
    updir = os.path.join(current_app.static_folder, 'uploads', subdir)
    os.makedirs(updir, exist_ok=True)
    file_storage.save(os.path.join(updir, filename))
    return f'/static/uploads/{subdir}/{filename}'


# ── advisor_links qabul qilish + avtomatik chat yoqish ───────────────────────

def _upsert_accepted_link(cur, student_id, advisor_id, inviter_id):
    """advisor_links juftligini 'accepted' holatiga keltiradi (yo'q bo'lsa
    yaratadi). UNIQUE(student_id, advisor_id) va CHECK(student<>advisor) hurmat."""
    if student_id == advisor_id:
        return
    cur.execute("SELECT id FROM advisor_links WHERE student_id = %s AND advisor_id = %s",
                (student_id, advisor_id))
    ex = cur.fetchone()
    if ex:
        cur.execute("""UPDATE advisor_links SET status = 'accepted', invited_by = %s,
                       responded_at = NOW() WHERE id = %s""", (inviter_id, ex[0]))
    else:
        cur.execute("""INSERT INTO advisor_links
                       (student_id, advisor_id, status, invited_by, responded_at)
                       VALUES (%s, %s, 'accepted', %s, NOW())""",
                    (student_id, advisor_id, inviter_id))


def _enable_pair_messaging(cur, uid_a, uid_b, welcome_from_id, welcome_text):
    """Aloqa qabul qilinganda ikki tomon orasidagi direct suhbatni ochadi va
    bir marta (yangi suhbat bo'lsa) xush kelibsiz xabari yozadi. Chat xatosi
    qabul jarayonini BUZMAYDI (try/except)."""
    try:
        from blueprints.messages import ensure_direct_conversation, post_system_message
        cid, created = ensure_direct_conversation(cur, uid_a, uid_b)
        if cid and created and welcome_text:
            post_system_message(cur, cid, welcome_from_id, welcome_text)
        return cid
    except Exception:
        return None


MAX_COLLABORATORS = 3


def _collab_count(cur, diss_id):
    """pending + accepted hamkorlar soni (3 chegarasi uchun)."""
    cur.execute("""SELECT COUNT(*) FROM diss_collaborators
                   WHERE dissertation_id = %s AND status IN ('pending', 'accepted')""",
                (diss_id,))
    return cur.fetchone()[0] or 0


def _accept_collaborator_token(cur, diss_id, created_by, can_comment, can_edit, can_review):
    """Hamkorlik taklif havolasini qabul: joriy foydalanuvchini
    diss_collaborators ga 'accepted' qilib qo'shadi (max 3 accepted), messaging
    yoqadi, egaga xabar beradi. Muvaffaqiyatda None, aks holda xato matni."""
    if not diss_id:
        return "Havola loyihaga bog'lanmagan."
    cur.execute("SELECT owner_id FROM diss_projects WHERE id = %s", (diss_id,))
    pr = cur.fetchone()
    if not pr:
        return 'Loyiha topilmadi.'
    owner_id = pr[0]
    if owner_id == current_user.id:
        return 'Bu sizning loyihangiz.'
    cur.execute("SELECT id, status FROM diss_collaborators "
                "WHERE dissertation_id = %s AND user_id = %s", (diss_id, current_user.id))
    ex = cur.fetchone()
    if ex and ex[1] == 'accepted':
        return None                       # allaqachon hamkor — jim o'tkazamiz
    cur.execute("""SELECT COUNT(*) FROM diss_collaborators
                   WHERE dissertation_id = %s AND status = 'accepted'""", (diss_id,))
    if (cur.fetchone()[0] or 0) >= MAX_COLLABORATORS:
        return f"Bu loyihada hamkorlar soni to'lgan (maksimal {MAX_COLLABORATORS} ta)."
    if ex:
        cur.execute("""UPDATE diss_collaborators SET status = 'accepted', invited_by = %s,
                       can_comment = %s, can_edit = %s, can_review_status = %s
                       WHERE id = %s""",
                    (created_by, can_comment, can_edit, can_review, ex[0]))
    else:
        cur.execute("""INSERT INTO diss_collaborators
                       (dissertation_id, user_id, invited_by, status,
                        can_comment, can_edit, can_review_status)
                       VALUES (%s, %s, %s, 'accepted', %s, %s, %s)""",
                    (diss_id, current_user.id, created_by, can_comment, can_edit, can_review))
    _notify(cur, owner_id, 'invite_accepted', diss_id, payload={'by': current_user.username})
    _enable_pair_messaging(cur, owner_id, current_user.id, owner_id,
                           f"{current_user.username} hamkorlik havolasi orqali qo'shildi. ✍️")
    return None


# ═════════════════════════════ ADVISOR LINKING ══════════════════════════════

@dissertation_bp.route('/api/advisor/invite', methods=['POST'])
@csrf.exempt
@login_required
def advisor_invite():
    data = request.get_json(silent=True) or {}
    ident = (data.get('username_or_email') or '').strip()
    role = data.get('role') or 'advisor'
    message = (data.get('message') or '').strip()[:1000]
    if not ident or role not in ('advisor', 'student'):
        return jsonify({'success': False, 'error': "Noto'g'ri so'rov"}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("SELECT id, username FROM users "
                        "WHERE LOWER(username) = LOWER(%s) OR LOWER(email) = LOWER(%s)",
                        (ident, ident))
            target = cur.fetchone()
            if not target:
                return jsonify({'success': False,
                                'error': 'Foydalanuvchi topilmadi. Username yoki emailni tekshiring.'}), 404
            tid = target[0]
            if tid == current_user.id:
                return jsonify({'success': False, 'error': "O'zingizni taklif qila olmaysiz"}), 400
            student_id, advisor_id = ((current_user.id, tid) if role == 'advisor'
                                      else (tid, current_user.id))
            cur.execute("""
                SELECT id, status FROM advisor_links
                WHERE student_id = %s AND advisor_id = %s
            """, (student_id, advisor_id))
            existing = cur.fetchone()
            if existing and existing[1] in ('pending', 'accepted'):
                return jsonify({'success': False,
                                'error': 'Bu foydalanuvchi bilan aloqa allaqachon mavjud'}), 409
            if existing:
                cur.execute("""
                    UPDATE advisor_links SET status = 'pending', invited_by = %s,
                        invite_message = %s, created_at = NOW(), responded_at = NULL
                    WHERE id = %s RETURNING id
                """, (current_user.id, message or None, existing[0]))
            else:
                cur.execute("""
                    INSERT INTO advisor_links (student_id, advisor_id, status,
                                               invited_by, invite_message)
                    VALUES (%s, %s, 'pending', %s, %s) RETURNING id
                """, (student_id, advisor_id, current_user.id, message or None))
            link_id = cur.fetchone()[0]
            _notify(cur, tid, 'advisor_invite',
                    payload={'link_id': link_id, 'role': role,
                             'from': current_user.username, 'message': message})
        conn.commit()
        return jsonify({'success': True, 'link_id': link_id})
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        conn.close()


@dissertation_bp.route('/api/advisor/respond', methods=['POST'])
@csrf.exempt
@login_required
def advisor_respond():
    data = request.get_json(silent=True) or {}
    action = data.get('action')
    try:
        link_id = int(data.get('link_id'))
    except (TypeError, ValueError):
        return jsonify({'success': False, 'error': "Noto'g'ri so'rov"}), 400
    if action not in ('accept', 'decline'):
        return jsonify({'success': False, 'error': "Noto'g'ri amal"}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("""
                SELECT student_id, advisor_id, invited_by, status
                FROM advisor_links WHERE id = %s
            """, (link_id,))
            link = cur.fetchone()
            if not link:
                abort(404)
            student_id, advisor_id, invited_by, status = link
            # faqat taklif QILINGAN tomon javob bera oladi
            invited_party = advisor_id if invited_by == student_id else student_id
            if current_user.id != invited_party:
                abort(403)
            if status != 'pending':
                return jsonify({'success': False, 'error': 'Taklif allaqachon javoblangan'}), 409
            new_status = 'accepted' if action == 'accept' else 'declined'
            cur.execute("UPDATE advisor_links SET status = %s, responded_at = NOW() "
                        "WHERE id = %s", (new_status, link_id))
            if action == 'accept':
                _notify(cur, invited_by, 'invite_accepted',
                        payload={'link_id': link_id, 'by': current_user.username})
                # aloqa qabul qilindi → messaging avtomatik yoqiladi
                _enable_pair_messaging(
                    cur, student_id, advisor_id, current_user.id,
                    f"{current_user.username} taklifni qabul qildi — endi shu yerda "
                    f"bevosita muloqot qilishingiz mumkin. ✍️")
        conn.commit()
        return jsonify({'success': True, 'status': new_status})
    except Exception as e:
        conn.rollback()
        if getattr(e, 'code', None) in (403, 404):
            raise
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        conn.close()


@dissertation_bp.route('/api/advisor/my-links')
@login_required
def advisor_my_links():
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("""
                SELECT l.id, l.student_id, l.advisor_id, l.status, l.invited_by,
                       l.invite_message, l.created_at,
                       su.username, au.username
                FROM advisor_links l
                JOIN users su ON su.id = l.student_id
                JOIN users au ON au.id = l.advisor_id
                WHERE (l.student_id = %s OR l.advisor_id = %s)
                  AND l.status IN ('pending', 'accepted')
                ORDER BY l.created_at DESC LIMIT 100
            """, (current_user.id, current_user.id))
            advisors, students = [], []
            for r in cur.fetchall():
                item = {'link_id': r[0], 'status': r[3],
                        'invited_by_me': r[4] == current_user.id,
                        'message': r[5] or '', 'created_at': str(r[6])[:16]}
                if r[1] == current_user.id:   # men talabaman → bu mening rahbarim
                    advisors.append(dict(item, user_id=r[2], username=r[8]))
                else:                          # men rahbarman → bu shogirdim
                    students.append(dict(item, user_id=r[1], username=r[7]))
        conn.commit()
        return jsonify({'success': True, 'advisors': advisors, 'students': students})
    finally:
        conn.close()


@dissertation_bp.route('/api/advisor/remove/<int:link_id>', methods=['POST'])
@csrf.exempt
@login_required
def advisor_remove(link_id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("SELECT student_id, advisor_id FROM advisor_links WHERE id = %s",
                        (link_id,))
            link = cur.fetchone()
            if not link:
                abort(404)
            if current_user.id not in link:
                abort(403)
            cur.execute("UPDATE advisor_links SET status = 'removed', responded_at = NOW() "
                        "WHERE id = %s", (link_id,))
            # bu juftlikka bog'langan loyihalardan rahbarni uzish
            cur.execute("""
                UPDATE diss_projects SET advisor_id = NULL
                WHERE owner_id = %s AND advisor_id = %s
            """, (link[0], link[1]))
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


# ═════════════════════════ INVITE VIA SHAREABLE LINK ════════════════════════

@dissertation_bp.route('/api/advisor/invite-link', methods=['POST'])
@csrf.exempt
@login_required
def advisor_invite_link():
    """Bir martalik taklif havolasini yaratadi (advisor/student). 7 kun amal."""
    data = request.get_json(silent=True) or {}
    role = data.get('role') if data.get('role') in ('advisor', 'student') else 'advisor'
    token = secrets.token_urlsafe(32)[:64]
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("""INSERT INTO advisor_invite_tokens (token, created_by, role)
                           VALUES (%s, %s, %s)""", (token, current_user.id, role))
        conn.commit()
        url = request.url_root.rstrip('/') + '/invite/' + token
        return jsonify({'success': True, 'url': url, 'token': token})
    finally:
        conn.close()


def _invite_error(message, code):
    return render_template('dissertation/invite_error.html', message=message), code


@dissertation_bp.route('/invite/<token>')
def invite_landing(token):
    """Taklif havolasi sahifasi. Login bo'lmasa — login (next bilan)."""
    if not current_user.is_authenticated:
        return redirect(url_for('auth.login', next=request.path))
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("""
                SELECT t.created_by, t.role, t.dissertation_id, t.expires_at,
                       t.used_by, u.username
                FROM advisor_invite_tokens t JOIN users u ON u.id = t.created_by
                WHERE t.token = %s
            """, (token,))
            row = cur.fetchone()
        conn.commit()
    finally:
        conn.close()
    if not row:
        return _invite_error('Taklif havolasi topilmadi yoki bekor qilingan.', 404)
    created_by, role, diss_id, expires_at, used_by, inviter = row
    if used_by:
        return _invite_error('Bu taklif havolasi allaqachon ishlatilgan.', 410)
    if expires_at and expires_at < datetime.now():
        return _invite_error('Taklif havolasi muddati tugagan (havola 7 kun amal qiladi).', 410)
    if created_by == current_user.id:
        return _invite_error("Bu sizning taklif havolangiz — o'zingizni qo'sha olmaysiz.", 400)
    return render_template('dissertation/invite_confirm.html',
                           token=token, role=role, inviter=inviter)


@dissertation_bp.route('/invite/<token>/respond', methods=['POST'])
@csrf.exempt
@login_required
def invite_respond(token):
    action = (request.get_json(silent=True) or {}).get('action')
    if action not in ('accept', 'decline'):
        return jsonify({'success': False, 'error': "Noto'g'ri amal"}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("""
                SELECT id, created_by, role, dissertation_id, expires_at, used_by,
                       can_comment, can_edit, can_review_status
                FROM advisor_invite_tokens WHERE token = %s FOR UPDATE
            """, (token,))
            row = cur.fetchone()
            if not row:
                return jsonify({'success': False, 'error': 'Havola topilmadi'}), 404
            (tid, created_by, role, diss_id, expires_at, used_by,
             can_comment, can_edit, can_review) = row
            if used_by:
                return jsonify({'success': False, 'error': 'Havola allaqachon ishlatilgan'}), 409
            if expires_at and expires_at < datetime.now():
                return jsonify({'success': False, 'error': 'Havola muddati tugagan'}), 410
            if created_by == current_user.id:
                return jsonify({'success': False, 'error': "O'zingizni qo'sha olmaysiz"}), 400
            if action == 'decline':
                cur.execute("UPDATE advisor_invite_tokens SET used_by = %s, used_at = NOW() "
                            "WHERE id = %s", (current_user.id, tid))
                conn.commit()
                return jsonify({'success': True, 'status': 'declined', 'redirect': '/workspace'})
            # ── accept ──
            if role in ('advisor', 'student'):
                if role == 'advisor':        # yaratgan = shogird, men = rahbar
                    student_id, advisor_id = created_by, current_user.id
                else:                        # yaratgan = rahbar, men = shogird
                    student_id, advisor_id = current_user.id, created_by
                if student_id == advisor_id:
                    return jsonify({'success': False, 'error': "O'zingizni qo'sha olmaysiz"}), 400
                _upsert_accepted_link(cur, student_id, advisor_id, created_by)
                _notify(cur, created_by, 'invite_accepted',
                        payload={'by': current_user.username})
                _enable_pair_messaging(
                    cur, student_id, advisor_id, created_by,
                    f"{current_user.username} taklif havolasi orqali aloqani qabul qildi. ✍️")
            elif role == 'collaborator':
                err = _accept_collaborator_token(cur, diss_id, created_by,
                                                 can_comment, can_edit, can_review)
                if err:
                    return jsonify({'success': False, 'error': err}), 400
            cur.execute("UPDATE advisor_invite_tokens SET used_by = %s, used_at = NOW() "
                        "WHERE id = %s", (current_user.id, tid))
        conn.commit()
        return jsonify({'success': True, 'status': 'accepted', 'redirect': '/workspace'})
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        conn.close()


# ═══════════════════════════ COLLABORATORS (max 3) ══════════════════════════

@dissertation_bp.route('/api/dissertation/<int:id>/collaborators')
@login_required
def collaborator_list(id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            get_dissertation_or_403(cur, id, allow_advisor=False)   # faqat ega
            cur.execute("""
                SELECT c.id, c.user_id, u.username, c.status,
                       c.can_comment, c.can_edit, c.can_review_status
                FROM diss_collaborators c JOIN users u ON u.id = c.user_id
                WHERE c.dissertation_id = %s AND c.status IN ('pending', 'accepted')
                ORDER BY c.created_at
            """, (id,))
            items = [{'id': r[0], 'user_id': r[1], 'username': r[2], 'status': r[3],
                      'can_comment': r[4], 'can_edit': r[5], 'can_review_status': r[6]}
                     for r in cur.fetchall()]
        conn.commit()
        return jsonify({'success': True, 'collaborators': items, 'max': MAX_COLLABORATORS})
    finally:
        conn.close()


@dissertation_bp.route('/api/dissertation/<int:id>/collaborators/invite', methods=['POST'])
@csrf.exempt
@login_required
def collaborator_invite(id):
    data = request.get_json(silent=True) or {}
    ident = (data.get('username_or_email') or '').strip()
    if not ident:
        return jsonify({'success': False, 'error': 'Username yoki email kiriting'}), 400
    can_comment = bool(data.get('can_comment', True))
    can_edit = bool(data.get('can_edit', False))
    can_review = bool(data.get('can_review_status', False))
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            p, role = get_dissertation_or_403(cur, id, allow_advisor=False)  # faqat ega
            cur.execute("SELECT id FROM users WHERE LOWER(username) = LOWER(%s) "
                        "OR LOWER(email) = LOWER(%s)", (ident, ident))
            target = cur.fetchone()
            if not target:
                return jsonify({'success': False, 'error': 'Foydalanuvchi topilmadi'}), 404
            tid = target[0]
            if tid == current_user.id:
                return jsonify({'success': False, 'error': "O'zingizni qo'sha olmaysiz"}), 400
            if tid == p['advisor_id']:
                return jsonify({'success': False,
                                'error': 'Bu foydalanuvchi allaqachon ilmiy rahbar'}), 400
            cur.execute("SELECT id, status FROM diss_collaborators "
                        "WHERE dissertation_id = %s AND user_id = %s", (id, tid))
            ex = cur.fetchone()
            if ex and ex[1] in ('pending', 'accepted'):
                return jsonify({'success': False, 'error': 'Bu foydalanuvchi allaqachon hamkor'}), 409
            if _collab_count(cur, id) >= MAX_COLLABORATORS:
                return jsonify({'success': False,
                                'error': f'Maksimal {MAX_COLLABORATORS} ta hamkor qo\'shish mumkin'}), 400
            if ex:
                cur.execute("""UPDATE diss_collaborators SET status = 'pending', invited_by = %s,
                               can_comment = %s, can_edit = %s, can_review_status = %s
                               WHERE id = %s RETURNING id""",
                            (current_user.id, can_comment, can_edit, can_review, ex[0]))
            else:
                cur.execute("""INSERT INTO diss_collaborators
                               (dissertation_id, user_id, invited_by, status,
                                can_comment, can_edit, can_review_status)
                               VALUES (%s, %s, %s, 'pending', %s, %s, %s) RETURNING id""",
                            (id, tid, current_user.id, can_comment, can_edit, can_review))
            collab_id = cur.fetchone()[0]
            _notify(cur, tid, 'collaborator_invite', id,
                    payload={'collab_id': collab_id, 'from': current_user.username,
                             'diss_title': p['title']})
        conn.commit()
        return jsonify({'success': True, 'collab_id': collab_id})
    except Exception as e:
        conn.rollback()
        if getattr(e, 'code', None) in (403, 404):
            raise
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        conn.close()


@dissertation_bp.route('/api/dissertation/<int:id>/collaborators/invite-link', methods=['POST'])
@csrf.exempt
@login_required
def collaborator_invite_link(id):
    data = request.get_json(silent=True) or {}
    can_comment = bool(data.get('can_comment', True))
    can_edit = bool(data.get('can_edit', False))
    can_review = bool(data.get('can_review_status', False))
    token = secrets.token_urlsafe(32)[:64]
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            get_dissertation_or_403(cur, id, allow_advisor=False)   # faqat ega
            if _collab_count(cur, id) >= MAX_COLLABORATORS:
                return jsonify({'success': False,
                                'error': f'Maksimal {MAX_COLLABORATORS} ta hamkor qo\'shish mumkin'}), 400
            cur.execute("""INSERT INTO advisor_invite_tokens
                           (token, created_by, role, dissertation_id,
                            can_comment, can_edit, can_review_status)
                           VALUES (%s, %s, 'collaborator', %s, %s, %s, %s)""",
                        (token, current_user.id, id, can_comment, can_edit, can_review))
        conn.commit()
        url = request.url_root.rstrip('/') + '/invite/' + token
        return jsonify({'success': True, 'url': url})
    finally:
        conn.close()


@dissertation_bp.route('/api/collaborators/<int:cid>/permissions', methods=['POST'])
@csrf.exempt
@login_required
def collaborator_permissions(cid):
    data = request.get_json(silent=True) or {}
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("SELECT dissertation_id FROM diss_collaborators WHERE id = %s", (cid,))
            r = cur.fetchone()
            if not r:
                abort(404)
            get_dissertation_or_403(cur, r[0], allow_advisor=False)   # faqat ega
            fields, params = [], []
            for key in ('can_comment', 'can_edit', 'can_review_status'):   # oq ro'yxat
                if key in data:
                    fields.append(f'{key} = %s')
                    params.append(bool(data[key]))
            if not fields:
                return jsonify({'success': False, 'error': "O'zgartirish yo'q"}), 400
            params.append(cid)
            cur.execute(f"UPDATE diss_collaborators SET {', '.join(fields)} WHERE id = %s",
                        tuple(params))
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


@dissertation_bp.route('/api/collaborators/<int:cid>/remove', methods=['POST'])
@csrf.exempt
@login_required
def collaborator_remove(cid):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("SELECT dissertation_id FROM diss_collaborators WHERE id = %s", (cid,))
            r = cur.fetchone()
            if not r:
                abort(404)
            get_dissertation_or_403(cur, r[0], allow_advisor=False)   # faqat ega
            cur.execute("UPDATE diss_collaborators SET status = 'removed' WHERE id = %s", (cid,))
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


@dissertation_bp.route('/api/collaborators/respond', methods=['POST'])
@csrf.exempt
@login_required
def collaborator_respond():
    data = request.get_json(silent=True) or {}
    action = data.get('action')
    try:
        collab_id = int(data.get('collab_id'))
    except (TypeError, ValueError):
        return jsonify({'success': False, 'error': "Noto'g'ri so'rov"}), 400
    if action not in ('accept', 'decline'):
        return jsonify({'success': False, 'error': "Noto'g'ri amal"}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("""SELECT dissertation_id, user_id, status
                           FROM diss_collaborators WHERE id = %s FOR UPDATE""", (collab_id,))
            r = cur.fetchone()
            if not r:
                abort(404)
            diss_id, user_id, status = r
            if user_id != current_user.id:
                abort(403)
            if status != 'pending':
                return jsonify({'success': False, 'error': 'Taklif allaqachon javoblangan'}), 409
            if action == 'decline':
                cur.execute("UPDATE diss_collaborators SET status = 'declined' WHERE id = %s",
                            (collab_id,))
                conn.commit()
                return jsonify({'success': True, 'status': 'declined'})
            cur.execute("""SELECT COUNT(*) FROM diss_collaborators
                           WHERE dissertation_id = %s AND status = 'accepted'""", (diss_id,))
            if (cur.fetchone()[0] or 0) >= MAX_COLLABORATORS:
                return jsonify({'success': False,
                                'error': f'Bu loyihada hamkorlar soni to\'lgan (maksimal {MAX_COLLABORATORS} ta)'}), 400
            cur.execute("UPDATE diss_collaborators SET status = 'accepted' WHERE id = %s",
                        (collab_id,))
            cur.execute("SELECT owner_id FROM diss_projects WHERE id = %s", (diss_id,))
            pr = cur.fetchone()
            if pr:
                _notify(cur, pr[0], 'invite_accepted', diss_id,
                        payload={'by': current_user.username})
                _enable_pair_messaging(cur, pr[0], current_user.id, pr[0],
                                       f"{current_user.username} hamkorlik taklifini qabul qildi. ✍️")
        conn.commit()
        return jsonify({'success': True, 'status': 'accepted'})
    except Exception as e:
        conn.rollback()
        if getattr(e, 'code', None) in (403, 404):
            raise
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        conn.close()


# ═════════════════════════════ PROJECT CRUD ═════════════════════════════════

@dissertation_bp.route('/workspace')
@login_required
def workspace_list():
    conn = _conn()
    mine, students_work, pending_invites, collab_projects = [], {}, [], []
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            # mening loyihalarim + progress (1 so'rov)
            cur.execute(f"""
                SELECT {_project_cols('d')},
                       (SELECT COUNT(*) FROM dissertation_blocks b WHERE b.dissertation_id = d.id),
                       (SELECT COUNT(*) FROM dissertation_blocks b WHERE b.dissertation_id = d.id
                          AND b.review_status = 'approved')
                FROM diss_projects d
                WHERE d.owner_id = %s AND d.status <> 'archived'
                ORDER BY d.updated_at DESC LIMIT 50
            """, (current_user.id,))
            adv_ids = set()
            for r in cur.fetchall():
                p = _project_dict(r[:-2])
                p['total_blocks'], p['approved_blocks'] = r[-2] or 0, r[-1] or 0
                p['progress'] = round(p['approved_blocks'] * 100 / p['total_blocks']) if p['total_blocks'] else 0
                mine.append(p)
                if p['advisor_id']:
                    adv_ids.add(p['advisor_id'])
            # shogirdlarim loyihalari (rahbar sifatida)
            cur.execute(f"""
                SELECT {_project_cols('d')},
                       (SELECT COUNT(*) FROM dissertation_blocks b WHERE b.dissertation_id = d.id),
                       (SELECT COUNT(*) FROM dissertation_blocks b WHERE b.dissertation_id = d.id
                          AND b.review_status = 'approved'),
                       u.username
                FROM diss_projects d
                JOIN advisor_links l ON l.student_id = d.owner_id
                     AND l.advisor_id = %s AND l.status = 'accepted'
                JOIN users u ON u.id = d.owner_id
                WHERE d.advisor_id = %s AND d.status <> 'archived'
                ORDER BY d.last_submitted_at DESC NULLS LAST LIMIT 100
            """, (current_user.id, current_user.id))
            for r in cur.fetchall():
                p = _project_dict(r[:-3])
                p['total_blocks'], p['approved_blocks'] = r[-3] or 0, r[-2] or 0
                p['progress'] = round(p['approved_blocks'] * 100 / p['total_blocks']) if p['total_blocks'] else 0
                students_work.setdefault(r[-1], []).append(p)
            names = _usernames(cur, adv_ids)
            for p in mine:
                p['advisor_name'] = names.get(p['advisor_id'], '')
            # hamkorlikdagi loyihalar (men qabul qilingan hamkorman)
            cur.execute(f"""
                SELECT {_project_cols('d')},
                       (SELECT COUNT(*) FROM dissertation_blocks b WHERE b.dissertation_id = d.id),
                       (SELECT COUNT(*) FROM dissertation_blocks b WHERE b.dissertation_id = d.id
                          AND b.review_status = 'approved'),
                       u.username, c.can_comment, c.can_edit, c.can_review_status
                FROM diss_projects d
                JOIN diss_collaborators c ON c.dissertation_id = d.id
                     AND c.user_id = %s AND c.status = 'accepted'
                JOIN users u ON u.id = d.owner_id
                WHERE d.status <> 'archived'
                ORDER BY d.updated_at DESC LIMIT 50
            """, (current_user.id,))
            for r in cur.fetchall():
                p = _project_dict(r[:-6])
                p['total_blocks'], p['approved_blocks'] = r[-6] or 0, r[-5] or 0
                p['progress'] = round(p['approved_blocks'] * 100 / p['total_blocks']) if p['total_blocks'] else 0
                p['owner_name'] = r[-4]
                p['can_comment'], p['can_edit'], p['can_review_status'] = r[-3], r[-2], r[-1]
                collab_projects.append(p)
            # menga kelgan pending takliflar
            cur.execute("""
                SELECT l.id, l.invited_by, l.invite_message, u.username,
                       (l.invited_by = l.student_id)
                FROM advisor_links l JOIN users u ON u.id = l.invited_by
                WHERE ((l.student_id = %s AND l.invited_by = l.advisor_id)
                    OR (l.advisor_id = %s AND l.invited_by = l.student_id))
                  AND l.status = 'pending'
                ORDER BY l.created_at DESC LIMIT 20
            """, (current_user.id, current_user.id))
            pending_invites = [{'link_id': r[0], 'from': r[3],
                                'message': r[2] or '',
                                'as_advisor': bool(r[4])} for r in cur.fetchall()]
        conn.commit()
    finally:
        conn.close()
    return render_template('dissertation/list.html', mine=mine,
                           students_work=students_work,
                           pending_invites=pending_invites,
                           collab_projects=collab_projects,
                           degree_labels=DEGREE_LABELS, status_labels=STATUS_LABELS)


@dissertation_bp.route('/api/dissertation/create', methods=['POST'])
@csrf.exempt
@login_required
def project_create():
    data = request.get_json(silent=True) or {}
    title = (data.get('title') or '').strip()[:600]
    degree = data.get('degree_type') if data.get('degree_type') in DEGREE_LABELS else 'phd'
    specialty = (data.get('specialty_code') or '').strip()[:30]
    if not title:
        return jsonify({'success': False, 'error': 'Sarlavha kiritilishi shart'}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("SELECT COALESCE(is_premium, FALSE) FROM users WHERE id = %s",
                        (current_user.id,))
            premium = bool((cur.fetchone() or [False])[0])
            if not premium:
                cur.execute("SELECT COUNT(*) FROM diss_projects "
                            "WHERE owner_id = %s AND status <> 'archived'",
                            (current_user.id,))
                if (cur.fetchone()[0] or 0) >= 1:
                    return jsonify({'success': False,
                                    'error': 'Bepul tarifda 1 ta loyiha yaratish mumkin'}), 403
            cur.execute("""
                INSERT INTO diss_projects (owner_id, title, degree_type, specialty_code)
                VALUES (%s, %s, %s, %s) RETURNING id
            """, (current_user.id, title, degree, specialty or None))
            pid = cur.fetchone()[0]
            # boshlang'ich skelet: Kirish (special) / I bob (chapter) / Xulosa (special)
            skeleton = [('Kirish', 'special'), ('I bob', 'chapter'),
                        ('Xulosa, taklif va tavsiyalar', 'special')]
            for i, (t, bt) in enumerate(skeleton):
                cur.execute("""
                    INSERT INTO dissertation_blocks
                        (dissertation_id, title, sort_order, depth, block_type)
                    VALUES (%s, %s, %s, 0, %s)
                """, (pid, t, i, bt))
            _recompute_numbering(cur, pid)
        conn.commit()
        return jsonify({'success': True, 'id': pid})
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        conn.close()


@dissertation_bp.route('/api/dissertation/<int:id>/update', methods=['POST'])
@csrf.exempt
@login_required
def project_update(id):
    data = request.get_json(silent=True) or {}
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            p, role = get_dissertation_or_403(cur, id, allow_advisor=False)
            title = (data.get('title') or p['title']).strip()[:600]
            degree = data.get('degree_type') if data.get('degree_type') in DEGREE_LABELS else p['degree_type']
            specialty = (data.get('specialty_code') or '').strip()[:30] or p['specialty_code']
            cur.execute("""
                UPDATE diss_projects SET title = %s, degree_type = %s,
                    specialty_code = %s, updated_at = NOW() WHERE id = %s
            """, (title, degree, specialty, id))
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


@dissertation_bp.route('/api/dissertation/<int:id>/assign-advisor', methods=['POST'])
@csrf.exempt
@login_required
def project_assign_advisor(id):
    data = request.get_json(silent=True) or {}
    try:
        advisor_id = int(data.get('advisor_id'))
    except (TypeError, ValueError):
        return jsonify({'success': False, 'error': "Noto'g'ri so'rov"}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            get_dissertation_or_403(cur, id, allow_advisor=False)
            cur.execute("""
                SELECT 1 FROM advisor_links
                WHERE student_id = %s AND advisor_id = %s AND status = 'accepted'
            """, (current_user.id, advisor_id))
            if not cur.fetchone():
                return jsonify({'success': False,
                                'error': 'Avval rahbar taklifni qabul qilishi kerak'}), 400
            cur.execute("UPDATE diss_projects SET advisor_id = %s, updated_at = NOW() "
                        "WHERE id = %s", (advisor_id, id))
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


@dissertation_bp.route('/api/dissertation/<int:id>/archive', methods=['POST'])
@csrf.exempt
@login_required
def project_archive(id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            get_dissertation_or_403(cur, id, allow_advisor=False)
            cur.execute("UPDATE diss_projects SET status = 'archived', updated_at = NOW() "
                        "WHERE id = %s", (id,))
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


# ═════════════════════════════ BLOCKS (TOC) ═════════════════════════════════

@dissertation_bp.route('/api/dissertation/<int:id>/blocks')
@login_required
def blocks_tree(id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            get_dissertation_or_403(cur, id)
            tree = _fetch_tree(cur, id)
        conn.commit()
        return jsonify({'success': True, 'blocks': tree})
    finally:
        conn.close()


@dissertation_bp.route('/api/dissertation/<int:id>/blocks/create', methods=['POST'])
@csrf.exempt
@login_required
def block_create(id):
    data = request.get_json(silent=True) or {}
    title = (data.get('title') or '').strip()[:MAX_TITLE]
    parent_id = data.get('parent_id')
    if not title:
        return jsonify({'success': False, 'error': 'Sarlavha kiritilishi shart'}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            get_dissertation_or_403(cur, id, allow_advisor=False)
            depth = 0
            if parent_id:
                cur.execute("SELECT depth, dissertation_id FROM dissertation_blocks WHERE id = %s",
                            (parent_id,))
                parent = cur.fetchone()
                if not parent or parent[1] != id:
                    return jsonify({'success': False, 'error': 'Ota qism topilmadi'}), 404
                depth = parent[0] + 1
                if depth > MAX_DEPTH:
                    return jsonify({'success': False,
                                    'error': "Maksimal 3 daraja ichma-ich qism yaratish mumkin"}), 400
            cur.execute("""
                SELECT COALESCE(MAX(sort_order), -1) + 1 FROM dissertation_blocks
                WHERE dissertation_id = %s AND parent_id IS NOT DISTINCT FROM %s
            """, (id, parent_id))
            order = cur.fetchone()[0]
            btype = 'special' if _is_special_title(title) else 'chapter'
            cur.execute("""
                INSERT INTO dissertation_blocks
                    (dissertation_id, parent_id, title, sort_order, depth, block_type)
                VALUES (%s, %s, %s, %s, %s, %s) RETURNING id
            """, (id, parent_id, title, order, depth, btype))
            bid = cur.fetchone()[0]
            _recompute_numbering(cur, id)
        conn.commit()
        return jsonify({'success': True, 'id': bid})
    except Exception as e:
        conn.rollback()
        if getattr(e, 'code', None) in (403, 404):
            raise
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        conn.close()


@dissertation_bp.route('/api/blocks/<int:block_id>/rename', methods=['POST'])
@csrf.exempt
@login_required
def block_rename(block_id):
    data = request.get_json(silent=True) or {}
    title = (data.get('title') or '').strip()[:MAX_TITLE]
    if not title:
        return jsonify({'success': False, 'error': 'Sarlavha kiritilishi shart'}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            block, project, role = get_block_or_403(cur, block_id, allow_advisor=False)
            # "Raqamlanmasin" toggle'i berilgan bo'lsa — undan; aks holda sarlavhadan
            # avto-aniqlash (masalan "I bob" → "Xulosa" special bo'ladi).
            if 'is_special' in data:
                btype = 'special' if data.get('is_special') else 'chapter'
            else:
                btype = 'special' if _is_special_title(title) else 'chapter'
            cur.execute("UPDATE dissertation_blocks SET title = %s, block_type = %s, "
                        "updated_at = NOW() WHERE id = %s", (title, btype, block_id))
            _recompute_numbering(cur, project['id'])
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


@dissertation_bp.route('/api/blocks/<int:block_id>/delete', methods=['POST'])
@csrf.exempt
@login_required
def block_delete(block_id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            block, project, role = get_block_or_403(cur, block_id, allow_advisor=False)
            cur.execute("DELETE FROM dissertation_blocks WHERE id = %s", (block_id,))
            _recompute_numbering(cur, project['id'])
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


def _subtree_height(node):
    if not node['children']:
        return 0
    return 1 + max(_subtree_height(c) for c in node['children'])


@dissertation_bp.route('/api/blocks/<int:block_id>/reorder', methods=['POST'])
@csrf.exempt
@login_required
def block_reorder(block_id):
    data = request.get_json(silent=True) or {}
    new_parent = data.get('new_parent_id')
    try:
        new_order = int(data.get('new_sort_order', 0))
    except (TypeError, ValueError):
        new_order = 0
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            block, project, role = get_block_or_403(cur, block_id, allow_advisor=False)
            new_depth = 0
            if new_parent:
                if int(new_parent) == block_id:
                    return jsonify({'success': False, 'error': "O'zining ichiga ko'chirib bo'lmaydi"}), 400
                cur.execute("SELECT depth, dissertation_id FROM dissertation_blocks WHERE id = %s",
                            (new_parent,))
                parent = cur.fetchone()
                if not parent or parent[1] != project['id']:
                    return jsonify({'success': False, 'error': 'Ota qism topilmadi'}), 404
                new_depth = parent[0] + 1
            # chuqurlik cheklovi: ko'chirilayotgan poyaning balandligi bilan birga
            tree = _fetch_tree(cur, project['id'])
            def find(items):
                for n in items:
                    if n['id'] == block_id:
                        return n
                    f = find(n['children'])
                    if f:
                        return f
            node = find(tree)
            height = _subtree_height(node) if node else 0
            if new_depth + height > MAX_DEPTH:
                return jsonify({'success': False,
                                'error': "Ko'chirish ichki qismlar bilan birga 3 daraja chegarasidan oshadi"}), 400
            # ota ostidagi aylanish (parentni o'z avlodiga ko'chirish) tekshiruvi
            if new_parent and node:
                def contains(n, target):
                    return any(c['id'] == target or contains(c, target) for c in n['children'])
                if contains(node, int(new_parent)):
                    return jsonify({'success': False, 'error': "O'z ichki qismiga ko'chirib bo'lmaydi"}), 400
            cur.execute("""
                UPDATE dissertation_blocks
                SET parent_id = %s, sort_order = %s, updated_at = NOW() WHERE id = %s
            """, (new_parent, new_order, block_id))
            # depth'ni butun ko'chirilgan poya bo'ylab yangilash
            def set_depths(n, d):
                cur.execute("UPDATE dissertation_blocks SET depth = %s WHERE id = %s", (d, n['id']))
                for c in n['children']:
                    set_depths(c, d + 1)
            if node:
                set_depths(node, new_depth)
            # siblinglar orasiga joylashtirish: tartibni zichlash
            cur.execute("""
                SELECT id FROM dissertation_blocks
                WHERE dissertation_id = %s AND parent_id IS NOT DISTINCT FROM %s
                ORDER BY sort_order, id
            """, (project['id'], new_parent))
            for i, (sid,) in enumerate(cur.fetchall()):
                cur.execute("UPDATE dissertation_blocks SET sort_order = %s WHERE id = %s",
                            (i, sid))
            _recompute_numbering(cur, project['id'])
        conn.commit()
        return jsonify({'success': True})
    except Exception as e:
        conn.rollback()
        if getattr(e, 'code', None) in (403, 404):
            raise
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        conn.close()


# ═════════════════════════════ EDITOR + SAVE ════════════════════════════════

@dissertation_bp.route('/workspace/<int:id>/first')
@login_required
def workspace_first(id):
    """Loyihani ochish — birinchi blok muharririga yo'naltiradi."""
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            get_dissertation_or_403(cur, id)
            cur.execute("""
                SELECT id FROM dissertation_blocks
                WHERE dissertation_id = %s
                ORDER BY sort_order, id LIMIT 1
            """, (id,))
            r = cur.fetchone()
        conn.commit()
    finally:
        conn.close()
    if not r:
        return redirect('/workspace')
    return redirect(f'/workspace/{id}/edit/{r[0]}')


@dissertation_bp.route('/workspace/<int:id>/edit/<int:block_id>')
@login_required
def editor_page(id, block_id):
    conn = _conn()
    defense_days_left = None
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            block, project, caps = get_block_access(cur, block_id)
            if block['dissertation_id'] != id:
                abort(404)
            tree = _fetch_tree(cur, id)
            # Akademik Reja ulanmasi: himoya sanasi countdown (jadval hali
            # yaratilmagan bo'lishi mumkin — SAVEPOINT bilan himoyalangan)
            cur.execute("SAVEPOINT rj_plan")
            try:
                cur.execute("""SELECT target_defense_date FROM roadmap_plans
                               WHERE diss_project_id = %s AND is_active LIMIT 1""",
                            (id,))
                r = cur.fetchone()
                if r and r[0]:
                    from datetime import date as _date
                    defense_days_left = (r[0] - _date.today()).days
            except Exception:
                cur.execute("ROLLBACK TO SAVEPOINT rj_plan")
        conn.commit()
    finally:
        conn.close()
    # naive DB timestamp (Neon/VPS — UTC) → epoch ms: mijoz qoralama
    # solishtiruvi timezone'ga bog'liq bo'lmasin
    u = block.get('updated_at')
    try:
        updated_at_ms = int(u.replace(tzinfo=timezone.utc).timestamp() * 1000) if u else 0
    except (TypeError, AttributeError):
        updated_at_ms = 0
    return render_template('dissertation/editor.html', project=project,
                           block=block, tree=tree, role=caps['role'],
                           read_only=not caps['can_edit'],
                           can_review=caps['can_review_status'],
                           can_comment=caps['can_comment'],
                           defense_days_left=defense_days_left,
                           updated_at_ms=updated_at_ms,
                           review_labels=REVIEW_LABELS, annot_types=ANNOT_TYPES,
                           status_labels=STATUS_LABELS)


@dissertation_bp.route('/api/blocks/<int:block_id>/save', methods=['POST'])
@csrf.exempt
@login_required
def block_save(block_id):
    data = request.get_json(silent=True) or {}
    raw = data.get('content') or ''
    save_type = 'autosave' if data.get('save_type') == 'autosave' else 'manual'
    if len(raw.encode('utf-8', 'ignore')) > MAX_CONTENT_BYTES:
        return jsonify({'success': False,
                        'error': 'Qism hajmi 2MB dan oshdi — rasmlarni kichraytiring'}), 413
    # rate limit (per-block, per-worker): autosave 5 soniyada 1 marta
    now = time.time()
    if save_type == 'autosave' and now - _last_save.get(block_id, 0) < SAVE_MIN_INTERVAL:
        return jsonify({'success': False, 'rate_limited': True}), 429
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            block, project, caps = get_block_access(cur, block_id)
            if not caps['can_edit']:          # ega yoki can_edit hamkor
                abort(403)
            # soft-lock: boshqa foydalanuvchining yangi (3 daqiqadan yosh) locki
            cur.execute("""
                SELECT is_locked_by FROM dissertation_blocks
                WHERE id = %s AND is_locked_by IS NOT NULL AND is_locked_by <> %s
                  AND locked_at > NOW() - INTERVAL '3 minutes'
            """, (block_id, current_user.id))
            if cur.fetchone():
                return jsonify({'success': False,
                                'error': 'Bu qismni hozir boshqa foydalanuvchi tahrirlamoqda'}), 409
            clean = sanitize_html(raw)
            plain = html_to_plain(clean)
            wc = len(plain.split())
            if clean == (block['content'] or ''):
                _last_save[block_id] = now
                return jsonify({'success': True, 'unchanged': True, 'word_count': wc})
            # avvalgi kontent versiyaga (keyin 20 talik retensiya)
            cur.execute("""
                INSERT INTO block_versions (block_id, content, word_count, saved_by, save_type)
                VALUES (%s, %s, %s, %s, %s)
            """, (block_id, block['content'] or '', block['word_count'] or 0,
                  current_user.id, save_type))
            cur.execute("""
                DELETE FROM block_versions WHERE id IN (
                    SELECT id FROM block_versions WHERE block_id = %s
                    ORDER BY created_at DESC OFFSET %s)
            """, (block_id, VERSIONS_KEPT))
            cur.execute("""
                UPDATE dissertation_blocks
                SET content = %s, content_plain = %s, word_count = %s,
                    is_locked_by = %s, locked_at = NOW(), updated_at = NOW()
                WHERE id = %s
            """, (clean, plain, wc, current_user.id, block_id))
            _reanchor_annotations(cur, block_id, plain)
        conn.commit()
        _last_save[block_id] = now
        return jsonify({'success': True, 'word_count': wc,
                        'saved_at': datetime.now().strftime('%H:%M')})
    except Exception as e:
        conn.rollback()
        if getattr(e, 'code', None) in (403, 404):
            raise
        current_app.logger.exception('block save failed: block=%s user=%s',
                                     block_id, current_user.id)
        return jsonify({'success': False,
                        'error': "Serverda xatolik — qayta urinib ko'ring"}), 500
    finally:
        conn.close()


@dissertation_bp.route('/api/blocks/<int:block_id>/lock', methods=['POST'])
@csrf.exempt
@login_required
def block_lock(block_id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            block, project, caps = get_block_access(cur, block_id)
            if not caps['can_edit']:
                abort(403)
            if (block['is_locked_by'] and block['is_locked_by'] != current_user.id
                    and block['locked_at']
                    and (datetime.now() - block['locked_at']).total_seconds() < LOCK_TTL):
                return jsonify({'success': False,
                                'error': 'Bu qismni hozir boshqa foydalanuvchi tahrirlamoqda'}), 409
            cur.execute("UPDATE dissertation_blocks SET is_locked_by = %s, locked_at = NOW() "
                        "WHERE id = %s", (current_user.id, block_id))
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


@dissertation_bp.route('/api/blocks/<int:block_id>/unlock', methods=['POST'])
@csrf.exempt
@login_required
def block_unlock(block_id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("""
                UPDATE dissertation_blocks SET is_locked_by = NULL, locked_at = NULL
                WHERE id = %s AND is_locked_by = %s
            """, (block_id, current_user.id))
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


# ═════════════════════════════ VERSIONS ═════════════════════════════════════

@dissertation_bp.route('/api/blocks/<int:block_id>/versions')
@login_required
def versions_list(block_id):
    page = max(1, request.args.get('page', 1, type=int))
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            get_block_or_403(cur, block_id)
            cur.execute("""
                SELECT v.id, v.created_at, v.save_type, v.word_count, u.username
                FROM block_versions v LEFT JOIN users u ON u.id = v.saved_by
                WHERE v.block_id = %s
                ORDER BY v.created_at DESC LIMIT 20 OFFSET %s
            """, (block_id, (page - 1) * 20))
            items = [{'id': r[0], 'created_at': str(r[1])[:16],
                      'save_type': r[2], 'word_count': r[3] or 0,
                      'saved_by': r[4] or ''} for r in cur.fetchall()]
        conn.commit()
        return jsonify({'success': True, 'versions': items, 'page': page})
    finally:
        conn.close()


@dissertation_bp.route('/api/blocks/<int:block_id>/versions/<int:version_id>')
@login_required
def version_content(block_id, version_id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            get_block_or_403(cur, block_id)
            cur.execute("SELECT content, word_count, created_at FROM block_versions "
                        "WHERE id = %s AND block_id = %s", (version_id, block_id))
            r = cur.fetchone()
            if not r:
                abort(404)
        conn.commit()
        return jsonify({'success': True, 'content': r[0],
                        'word_count': r[1] or 0, 'created_at': str(r[2])[:16]})
    finally:
        conn.close()


@dissertation_bp.route('/api/blocks/<int:block_id>/restore/<int:version_id>', methods=['POST'])
@csrf.exempt
@login_required
def version_restore(block_id, version_id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            block, project, caps = get_block_access(cur, block_id)
            if not caps['can_edit']:
                abort(403)
            cur.execute("SELECT content FROM block_versions WHERE id = %s AND block_id = %s",
                        (version_id, block_id))
            v = cur.fetchone()
            if not v:
                abort(404)
            # tiklash ham qaytariladigan bo'lsin: avval joriy holat versiyaga
            cur.execute("""
                INSERT INTO block_versions (block_id, content, word_count, saved_by, save_type)
                VALUES (%s, %s, %s, %s, 'pre_restore')
            """, (block_id, block['content'] or '', block['word_count'] or 0, current_user.id))
            clean = sanitize_html(v[0])
            plain = html_to_plain(clean)
            cur.execute("""
                UPDATE dissertation_blocks
                SET content = %s, content_plain = %s, word_count = %s, updated_at = NOW()
                WHERE id = %s
            """, (clean, plain, len(plain.split()), block_id))
            _reanchor_annotations(cur, block_id, plain)
            cur.execute("""
                DELETE FROM block_versions WHERE id IN (
                    SELECT id FROM block_versions WHERE block_id = %s
                    ORDER BY created_at DESC OFFSET %s)
            """, (block_id, VERSIONS_KEPT))
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


# ═════════════════════════════ IMAGE UPLOAD ═════════════════════════════════

_IMG_TYPES = {'image/jpeg': '.jpg', 'image/png': '.png', 'image/webp': '.webp'}


@dissertation_bp.route('/api/blocks/<int:block_id>/upload-image', methods=['POST'])
@csrf.exempt
@login_required
def block_upload_image(block_id):
    f = request.files.get('image')
    if not f or not f.filename:
        return jsonify({'success': False, 'error': 'Fayl tanlanmagan'}), 400
    from werkzeug.utils import secure_filename
    ext = os.path.splitext(secure_filename(f.filename))[1].lower()
    mime = (f.mimetype or '').lower()
    if mime not in _IMG_TYPES or ext not in ('.jpg', '.jpeg', '.png', '.webp'):
        return jsonify({'success': False, 'error': 'Faqat JPG, PNG, WEBP qabul qilinadi'}), 400
    f.seek(0, os.SEEK_END)
    size = f.tell()
    f.seek(0)
    if size > IMG_MAX_BYTES:
        return jsonify({'success': False, 'error': 'Rasm hajmi 5MB dan oshmasligi kerak'}), 413
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            block, project, caps = get_block_access(cur, block_id)
            if not caps['can_edit']:
                abort(403)
        conn.commit()
    finally:
        conn.close()
    try:
        filename = f'{uuid.uuid4().hex}{_IMG_TYPES[mime]}'
        url = _store_upload(f, f'dissertation/diss_{project["id"]}/block_{block_id}', filename)
        return jsonify({'success': True, 'url': url})
    except Exception:
        return jsonify({'success': False,
                        'error': "Fayl yuklashda xatolik. Qayta urinib ko'ring."}), 500


# ═════════════════════════════ ANNOTATIONS ══════════════════════════════════

@dissertation_bp.route('/api/blocks/<int:block_id>/annotations')
@login_required
def annotations_list(block_id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            get_block_or_403(cur, block_id)
            cur.execute("""
                SELECT a.id, a.annotation_type, a.anchor_text, a.anchor_offset,
                       a.body, a.status, a.created_at, a.author_id, u.username,
                       a.resolved_at
                FROM block_annotations a JOIN users u ON u.id = a.author_id
                WHERE a.block_id = %s
                  AND (a.status IN ('open', 'orphaned')
                       OR (a.status = 'resolved' AND a.resolved_at > NOW() - INTERVAL '90 days'))
                ORDER BY a.created_at DESC LIMIT 200
            """, (block_id,))
            rows = cur.fetchall()
            ann_ids = [r[0] for r in rows]
            replies = {}
            if ann_ids:
                cur.execute("""
                    SELECT r.annotation_id, r.id, r.body, r.created_at, u.username
                    FROM annotation_replies r JOIN users u ON u.id = r.author_id
                    WHERE r.annotation_id = ANY(%s)
                    ORDER BY r.created_at
                """, (ann_ids,))
                for r in cur.fetchall():
                    replies.setdefault(r[0], []).append(
                        {'id': r[1], 'body': r[2], 'created_at': str(r[3])[:16],
                         'author': r[4]})
            open_, resolved, orphaned = [], [], []
            for r in rows:
                item = {'id': r[0], 'type': r[1], 'type_label': ANNOT_TYPES.get(r[1], r[1]),
                        'anchor_text': r[2], 'anchor_offset': r[3],
                        'body': r[4], 'status': r[5], 'created_at': str(r[6])[:16],
                        'author_id': r[7], 'author': r[8],
                        'mine': r[7] == current_user.id,
                        'replies': replies.get(r[0], [])}
                (open_ if r[5] == 'open' else
                 orphaned if r[5] == 'orphaned' else resolved).append(item)
        conn.commit()
        return jsonify({'success': True, 'open': open_,
                        'resolved': resolved[:10], 'orphaned': orphaned})
    finally:
        conn.close()


@dissertation_bp.route('/api/blocks/<int:block_id>/annotations/create', methods=['POST'])
@csrf.exempt
@login_required
def annotation_create(block_id):
    data = request.get_json(silent=True) or {}
    atype = data.get('annotation_type') if data.get('annotation_type') in ANNOT_TYPES else 'comment'
    anchor = (data.get('anchor_text') or '').strip()
    body = (data.get('body') or '').strip()
    if not anchor or len(anchor) > 500:
        return jsonify({'success': False, 'error': "Belgilangan matn 1-500 belgi bo'lishi kerak"}), 400
    if not body or len(body) > MAX_ANNOT_BODY:
        return jsonify({'success': False, 'error': 'Izoh matni 1-5000 belgi oralig\'ida bo\'lishi kerak'}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            block, project, caps = get_block_access(cur, block_id)
            if not caps['can_comment']:
                abort(403)
            offset = data.get('anchor_offset')
            pos = (block['content_plain'] or '').find(anchor)
            if pos >= 0:
                offset = pos
            cur.execute("""
                INSERT INTO block_annotations
                    (block_id, author_id, annotation_type, anchor_text,
                     anchor_prefix, anchor_suffix, anchor_offset, body)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s) RETURNING id
            """, (block_id, current_user.id, atype, anchor,
                  (data.get('anchor_prefix') or '')[:100],
                  (data.get('anchor_suffix') or '')[:100],
                  offset, body))
            aid = cur.fetchone()[0]
            # asosiy tomonlarga (ega + rahbar) xabar — o'zidan boshqasiga
            for target in {project['owner_id'], project['advisor_id']}:
                if target and target != current_user.id:
                    _notify(cur, target, 'annotation_added', project['id'], block_id,
                            {'annotation_id': aid, 'type': atype,
                             'block_title': block['title']})
        conn.commit()
        return jsonify({'success': True, 'id': aid})
    finally:
        conn.close()


def _get_annotation_or_403(cur, aid):
    cur.execute("""
        SELECT a.id, a.block_id, a.author_id, a.status,
               (SELECT COUNT(*) FROM annotation_replies r WHERE r.annotation_id = a.id)
        FROM block_annotations a WHERE a.id = %s
    """, (aid,))
    r = cur.fetchone()
    if not r:
        abort(404)
    block, project, caps = get_block_access(cur, r[1])
    return {'id': r[0], 'block_id': r[1], 'author_id': r[2], 'status': r[3],
            'reply_count': r[4]}, project, caps


@dissertation_bp.route('/api/annotations/<int:id>/reply', methods=['POST'])
@csrf.exempt
@login_required
def annotation_reply(id):
    body = ((request.get_json(silent=True) or {}).get('body') or '').strip()
    if not body or len(body) > MAX_ANNOT_BODY:
        return jsonify({'success': False, 'error': 'Javob matni 1-5000 belgi oralig\'ida'}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            ann, project, caps = _get_annotation_or_403(cur, id)
            if not caps['can_comment']:
                abort(403)
            cur.execute("""
                INSERT INTO annotation_replies (annotation_id, author_id, body)
                VALUES (%s, %s, %s) RETURNING id
            """, (id, current_user.id, body))
            rid = cur.fetchone()[0]
        conn.commit()
        return jsonify({'success': True, 'id': rid})
    finally:
        conn.close()


@dissertation_bp.route('/api/annotations/<int:id>/resolve', methods=['POST'])
@csrf.exempt
@login_required
def annotation_resolve(id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            ann, project, caps = _get_annotation_or_403(cur, id)
            if not caps['can_comment']:
                abort(403)
            # hamkor faqat O'Z izohini hal qila oladi; ega/rahbar — istalganini
            if caps['role'] == 'collaborator' and ann['author_id'] != current_user.id:
                abort(403)
            cur.execute("""
                UPDATE block_annotations
                SET status = 'resolved', resolved_by = %s, resolved_at = NOW()
                WHERE id = %s
            """, (current_user.id, id))
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


@dissertation_bp.route('/api/annotations/<int:id>/reopen', methods=['POST'])
@csrf.exempt
@login_required
def annotation_reopen(id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            ann, project, caps = _get_annotation_or_403(cur, id)
            if not caps['can_review_status']:      # rahbar yoki holat-huquqli hamkor
                abort(403)
            cur.execute("UPDATE block_annotations SET status = 'open', "
                        "resolved_by = NULL, resolved_at = NULL WHERE id = %s", (id,))
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


@dissertation_bp.route('/api/annotations/<int:id>/delete', methods=['POST'])
@csrf.exempt
@login_required
def annotation_delete(id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            ann, project, caps = _get_annotation_or_403(cur, id)
            if ann['author_id'] != current_user.id:
                abort(403)
            if ann['reply_count']:
                return jsonify({'success': False,
                                'error': "Javoblari bor izohni o'chirib bo'lmaydi"}), 400
            cur.execute("DELETE FROM block_annotations WHERE id = %s", (id,))
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


# ═════════════════════════════ REVIEW LOOP ══════════════════════════════════

@dissertation_bp.route('/api/blocks/<int:block_id>/review-status', methods=['POST'])
@csrf.exempt
@login_required
def block_review_status(block_id):
    status = (request.get_json(silent=True) or {}).get('status')
    if status not in ('deficiencies', 'task_assigned', 'approved'):
        return jsonify({'success': False, 'error': "Noto'g'ri holat"}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            block, project, caps = get_block_access(cur, block_id)
            if not caps['can_review_status']:     # rahbar yoki holat-huquqli hamkor
                abort(403)
            cur.execute("""
                UPDATE dissertation_blocks
                SET review_status = %s, review_status_by = %s, review_status_at = NOW()
                WHERE id = %s
            """, (status, current_user.id, block_id))
            _notify(cur, project['owner_id'], 'status_changed', project['id'], block_id,
                    {'status': status, 'label': REVIEW_LABELS.get(status),
                     'block_title': block['title']})
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


@dissertation_bp.route('/api/dissertation/<int:id>/submit-for-review', methods=['POST'])
@csrf.exempt
@login_required
def submit_for_review(id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            p, role = get_dissertation_or_403(cur, id, allow_advisor=False)
            if not p['advisor_id']:
                return jsonify({'success': False,
                                'error': 'Avval ilmiy rahbar biriktiring'}), 400
            cur.execute("""
                UPDATE diss_projects SET status = 'in_review',
                    last_submitted_at = NOW(), updated_at = NOW() WHERE id = %s
            """, (id,))
            _notify(cur, p['advisor_id'], 'student_submitted', id,
                    payload={'title': p['title'],
                             'message': "Shogirdingiz berilgan tuzatishlarni kiritdi, "
                                        "qayta ko'rib chiqasizmi?"})
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


@dissertation_bp.route('/api/dissertation/<int:id>/finish-review', methods=['POST'])
@csrf.exempt
@login_required
def finish_review(id):
    """Rahbar sessiyani yakunlaganda: BITTA jamlangan xabar (spam o'rniga)."""
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            p, role = get_dissertation_or_403(cur, id)
            if role != 'advisor':
                abort(403)
            since = p['last_submitted_at']
            cur.execute("""
                SELECT
                  COUNT(*) FILTER (WHERE a.annotation_type IN ('comment', 'correction')),
                  COUNT(*) FILTER (WHERE a.annotation_type = 'task')
                FROM block_annotations a
                JOIN dissertation_blocks b ON b.id = a.block_id
                WHERE b.dissertation_id = %s AND a.author_id = %s
                  AND (%s::timestamp IS NULL OR a.created_at > %s)
            """, (id, current_user.id, since, since))
            comments, tasks = cur.fetchone()
            cur.execute("""
                SELECT COUNT(*) FROM dissertation_blocks
                WHERE dissertation_id = %s AND review_status = 'approved'
                  AND (%s::timestamp IS NULL OR review_status_at > %s)
            """, (id, since, since))
            approved = cur.fetchone()[0]
            cur.execute("UPDATE diss_projects SET status = 'revision', updated_at = NOW() "
                        "WHERE id = %s", (id,))
            msg = (f"Ilmiy rahbaringiz ishingizni ko'rib chiqdi: {comments or 0} ta izoh, "
                   f"{tasks or 0} ta topshiriq, {approved or 0} ta qism ma'qullandi")
            _notify(cur, p['owner_id'], 'advisor_reviewed', id,
                    payload={'message': msg, 'comments': comments or 0,
                             'tasks': tasks or 0, 'approved': approved or 0})
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


# ═════════════════════════════ NOTIFICATIONS ════════════════════════════════

@dissertation_bp.route('/api/diss-notifications')
@login_required
def diss_notifications():
    unread_only = request.args.get('unread') == '1'
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute(f"""
                SELECT id, event_type, dissertation_id, block_id, payload,
                       is_read, created_at
                FROM diss_notifications
                WHERE user_id = %s {"AND is_read = FALSE" if unread_only else ""}
                ORDER BY created_at DESC LIMIT 20
            """, (current_user.id,))
            items = [{'id': r[0], 'event_type': r[1], 'dissertation_id': r[2],
                      'block_id': r[3], 'payload': r[4] or {},
                      'is_read': r[5], 'created_at': str(r[6])[:16]}
                     for r in cur.fetchall()]
        conn.commit()
        return jsonify({'success': True, 'notifications': items})
    finally:
        conn.close()


@dissertation_bp.route('/api/diss-notifications/mark-read', methods=['POST'])
@csrf.exempt
@login_required
def diss_notifications_mark_read():
    data = request.get_json(silent=True) or {}
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            if data.get('all'):
                cur.execute("UPDATE diss_notifications SET is_read = TRUE "
                            "WHERE user_id = %s AND is_read = FALSE", (current_user.id,))
            else:
                ids = [int(i) for i in (data.get('ids') or []) if str(i).isdigit()][:100]
                if ids:
                    cur.execute("UPDATE diss_notifications SET is_read = TRUE "
                                "WHERE user_id = %s AND id = ANY(%s)",
                                (current_user.id, ids))
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


# ═════════════════════════════ PREVIEW + EXPORT ═════════════════════════════

def _ordered_blocks(cur, diss_id):
    """Daraxt tartibida tekislangan bloklar (preview/eksport uchun)."""
    tree = _fetch_tree(cur, diss_id)
    ids = []
    def walk(items):
        for n in items:
            ids.append(n['id'])
            walk(n['children'])
    walk(tree)
    if not ids:
        return []
    cur.execute("""
        SELECT id, title, numbering, depth, content, block_type
        FROM dissertation_blocks WHERE id = ANY(%s)
    """, (ids,))
    by_id = {r[0]: {'id': r[0], 'title': r[1], 'numbering': r[2] or '',
                    'depth': r[3], 'content': r[4] or '',
                    'block_type': r[5] or 'chapter',
                    'is_special': (r[5] or 'chapter') == 'special',
                    'heading': _heading_label(r[2] or '', r[1])}
             for r in cur.fetchall()}
    return [by_id[i] for i in ids if i in by_id]


@dissertation_bp.route('/workspace/<int:id>/preview')
@login_required
def preview_page(id):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            p, role = get_dissertation_or_403(cur, id)
            blocks = _ordered_blocks(cur, id)
            owner = _usernames(cur, [p['owner_id']]).get(p['owner_id'], '')
        conn.commit()
    finally:
        conn.close()
    return render_template('dissertation/preview.html', project=p,
                           blocks=blocks, owner_name=owner,
                           year=datetime.now().year)


@dissertation_bp.route('/workspace/<int:id>/export-docx')
@login_required
def export_docx(id):
    # bir foydalanuvchi 30 soniyada 1 marta (og'ir operatsiya)
    now = time.time()
    if now - _last_export.get(current_user.id, 0) < 30:
        return render_template('dissertation/preview_error.html',
                               message='Eksport tayyorlanmoqda, biroz kuting'), 429
    _last_export[current_user.id] = now
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            p, role = get_dissertation_or_403(cur, id)
            blocks = _ordered_blocks(cur, id)
            owner = _usernames(cur, [p['owner_id']]).get(p['owner_id'], '')
        conn.commit()
    finally:
        conn.close()
    try:
        buf = _build_docx(p, blocks, owner)
    except Exception:
        return render_template('dissertation/preview_error.html',
                               message="Word fayl tayyorlashda xatolik yuz berdi. "
                                       "Qayta urinib ko'ring."), 500
    from institutions import transliterate
    slug = re.sub(r'[^a-z0-9]+', '-', transliterate(p['title'].lower()))[:60].strip('-') or 'dissertatsiya'
    return send_file(buf, as_attachment=True, download_name=f'{slug}.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


def _build_docx(project, blocks, owner_name):
    """OAK standartida .docx: Times New Roman 14pt, 1.5 interval,
    chap 3sm / o'ng 1.5sm / tepa-past 2sm."""
    import io
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.left_margin, section.right_margin = Cm(3), Cm(1.5)
        section.top_margin = section.bottom_margin = Cm(2)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5

    # titul varag'i
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(project['title'])
    run.bold = True
    run.font.size = Pt(16)
    for text in [owner_name, project.get('specialty_code') or '',
                 str(datetime.now().year)]:
        if text:
            pr = doc.add_paragraph()
            pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pr.add_run(str(text))
    doc.add_page_break()

    # mundarija
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run('MUNDARIJA').bold = True
    for b in blocks:
        doc.add_paragraph(b.get('heading') or b['title'])
    doc.add_page_break()

    # bo'limlar
    for b in blocks:
        heading = doc.add_heading(level=min(b['depth'] + 1, 3))
        hr = heading.add_run(b.get('heading') or b['title'])
        hr.font.name = 'Times New Roman'
        _html_to_docx(doc, b['content'])
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _html_to_docx(doc, html_str):
    """Soddalashtirilgan HTML→docx: p/h/b/i/u/ro'yxat/jadval/rasm.
    Rasm yuklab bo'lmasa '[Rasm yuklanmadi]' o'rinbosari qo'yiladi."""
    from docx.shared import Cm
    if not (html_str or '').strip():
        return
    # bloklarga bo'lish
    parts = re.split(r'(<img[^>]*>|<table.*?</table>|<[uo]l>.*?</[uo]l>|<h[1-4][^>]*>.*?</h[1-4]>|<p[^>]*>.*?</p>|<blockquote[^>]*>.*?</blockquote>)',
                     html_str, flags=re.S | re.I)
    for part in parts:
        part = (part or '').strip()
        if not part:
            continue
        low = part.lower()
        if low.startswith('<img'):
            m = re.search(r'src=["\']([^"\']+)["\']', part)
            _docx_image(doc, m.group(1) if m else '')
        elif low.startswith('<table'):
            rows = re.findall(r'<tr.*?</tr>', part, flags=re.S | re.I)
            if rows:
                ncols = max(len(re.findall(r'<t[dh]', r2, flags=re.I)) for r2 in rows)
                table = doc.add_table(rows=0, cols=max(1, ncols))
                table.style = 'Table Grid'
                for r2 in rows:
                    cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', r2, flags=re.S | re.I)
                    row = table.add_row()
                    for i, c in enumerate(cells[:ncols]):
                        row.cells[i].text = html_to_plain(c)
        elif low.startswith('<ul') or low.startswith('<ol'):
            items = re.findall(r'<li[^>]*>(.*?)</li>', part, flags=re.S | re.I)
            style_name = 'List Bullet' if low.startswith('<ul') else 'List Number'
            for it in items:
                try:
                    doc.add_paragraph(html_to_plain(it), style=style_name)
                except Exception:
                    doc.add_paragraph('• ' + html_to_plain(it))
        else:
            # p / h / blockquote / yalang'och matn — runlar bilan (b/i/u)
            inner = re.sub(r'^<[^>]+>|</[^>]+>$', '', part, flags=re.S)
            para = doc.add_paragraph()
            _runs_from_html(para, inner)
            # ichidagi rasmlar (p ichida bo'lsa)
            for m in re.finditer(r'<img[^>]*src=["\']([^"\']+)["\']', part, flags=re.I):
                _docx_image(doc, m.group(1))


def _runs_from_html(para, inner_html):
    """b/i/u formatlash bilan runlar. Murakkab holatlar plain matnga tushadi."""
    tokens = re.split(r'(</?(?:b|strong|i|em|u)>)', inner_html, flags=re.I)
    bold = italic = underline = False
    for tok in tokens:
        low = tok.lower()
        if low in ('<b>', '<strong>'):
            bold = True
        elif low in ('</b>', '</strong>'):
            bold = False
        elif low in ('<i>', '<em>'):
            italic = True
        elif low in ('</i>', '</em>'):
            italic = False
        elif low == '<u>':
            underline = True
        elif low == '</u>':
            underline = False
        else:
            text = html_to_plain(tok)
            if text:
                run = para.add_run(text + ' ')
                run.bold, run.italic, run.underline = bold, italic, underline


def _docx_image(doc, src):
    from docx.shared import Cm
    try:
        if src.startswith('/static/'):
            path = os.path.join(current_app.root_path, src.lstrip('/'))
            doc.add_picture(path, width=Cm(14))
        elif src.startswith('http'):
            import io
            import requests
            r = requests.get(src, timeout=10)
            r.raise_for_status()
            doc.add_picture(io.BytesIO(r.content), width=Cm(14))
        else:
            raise ValueError('unsupported src')
    except Exception:
        doc.add_paragraph('[Rasm yuklanmadi]')
