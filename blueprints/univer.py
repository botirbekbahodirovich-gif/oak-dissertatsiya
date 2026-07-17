"""Universitet B2B paneli (univer_bp) — Part 1: kirish modeli, workspace,
monitoring ko'rinishlari.

Universitetning vakolatli xodimi tizimga kirib, platformadagi o'z
universitetiga tegishli HAMMA narsani bitta boshqaruv markazida ko'radi:
olimlar, dissertatsiyalar, faol doktorantlar progressi, statistika.

Manba jadvallar (hech biri yangi emas):
  dissertations        — himoyalar korpusi (muassasa bo'yicha skoplanadi)
  institution_map      — canonical_name ↔ cyrillic_name variantlari (kengaytma
                         data.py dashboard Muassasa faceti bilan bir xil)
  olim_profiles        — claimed profillar (institution ustuni — universitetga
                         bog'lash zanjirining boshi)
  cabinet_users/users  — olim_profiles.cabinet_user_id → cabinet_users.email →
                         users (olimlar_catalog naqshidagi zanjir)
  roadmap_* / diss_*   — doktorant progress METRIKALARI (faqat sonlar!)

Yangi jadvallar (migrations/add_university_panel.sql aksi):
  university_staff, university_licenses, university_invite_tokens
  + users.hide_from_university (maxfiylik opt-out)

MAXFIYLIK CHEGARASI (kritik invariant):
  Universitet xodimi HECH QACHON dissertatsiya MATNINI
  (dissertation_blocks.content), chat xabarlarini, annotatsiyalarni yoki
  shaxsiy yozuvlarni ko'rmaydi. Faqat METRIKALAR: so'z soni, progress %,
  holatlar, nashr sonlari. Har endpoint server tomonda shuni ta'minlaydi —
  bu modul content/content_plain ustunlariga umuman murojaat qilmaydi,
  block_id/diss_id identifikatorlari ham javoblardan olib tashlanadi.

Xavfsizlik: barcha endpointlar @login_required + universitet-rol tekshiruvi
(get_university_access_or_403) + parametrlangan SQL. Boshqa universitet
ma'lumotiga urinish → 403. Litsenziya tugagan → do'stona sahifa.
"""
import re
import secrets
import time
from datetime import date, datetime

from flask import (Blueprint, jsonify, request, render_template, redirect,
                   abort, make_response, Response, g, url_for, flash)
from flask_login import login_required, current_user

from app import csrf

univer_bp = Blueprint('univer', __name__)

_schema_ready = False

_CACHE_TTL = 900            # universitet-skopli agregat keshi — 15 daqiqa
_NAV_TTL = 120              # navbar bayrog'i keshi — 2 daqiqa
_univ_cache = {}            # canonical → {'data': dict, 'ts': float}
_nav_cache = {}             # user_id → (ts, canonical yoki '')
_export_last = {}           # user_id → ts (CSV eksport rate-limit, 30s)

PLAN_LABELS = {'pilot': 'Pilot', 'standard': 'Standard', 'premium': 'Premium'}
ROLE_LABELS = {'owner': 'Egasi', 'staff': 'Xodim', 'viewer': "Ko'ruvchi"}
DEFAULT_MAX_STAFF = 5

# data.py bilan bir xil yil ajratish ifodasi (sana — erkin matn)
_YEAR_EXPR = r"(regexp_match(TRIM(d.sana), '(19|20)\d{2}'))[1]"

_PHD_LIKE = "UPPER(TRIM(d.daraja)) LIKE '%%PHD%%'"
_DSC_LIKE = "UPPER(TRIM(d.daraja)) LIKE '%%DSC%%'"


def _conn():
    from data import get_connection
    return get_connection()


# ── sxema (lazy, idempotent — migrations/add_university_panel.sql aksi) ──────

def _ensure_schema(cur):
    global _schema_ready
    if _schema_ready:
        return
    cur.execute("""
        CREATE TABLE IF NOT EXISTS university_staff (
            id SERIAL PRIMARY KEY,
            user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
            canonical_institution VARCHAR(500) NOT NULL,
            role VARCHAR(30) NOT NULL DEFAULT 'staff'
                CHECK (role IN ('owner', 'staff', 'viewer')),
            title VARCHAR(200),
            invited_by INTEGER REFERENCES users(id),
            status VARCHAR(20) DEFAULT 'active'
                CHECK (status IN ('active', 'suspended')),
            created_at TIMESTAMP DEFAULT NOW(),
            UNIQUE(user_id, canonical_institution)
        )""")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_university_staff_user "
                "ON university_staff(user_id) WHERE status = 'active'")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_university_staff_inst "
                "ON university_staff(canonical_institution)")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS university_licenses (
            id SERIAL PRIMARY KEY,
            canonical_institution VARCHAR(500) UNIQUE NOT NULL,
            plan VARCHAR(20) DEFAULT 'pilot'
                CHECK (plan IN ('pilot', 'standard', 'premium')),
            valid_until DATE,
            max_staff INTEGER DEFAULT 5,
            created_at TIMESTAMP DEFAULT NOW(),
            notes TEXT
        )""")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS university_invite_tokens (
            id SERIAL PRIMARY KEY,
            token VARCHAR(64) UNIQUE NOT NULL,
            license_id INTEGER NOT NULL
                REFERENCES university_licenses(id) ON DELETE CASCADE,
            created_by INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
            expires_at TIMESTAMP DEFAULT (NOW() + INTERVAL '7 days'),
            used_by INTEGER REFERENCES users(id),
            used_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT NOW()
        )""")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_university_invite_token "
                "ON university_invite_tokens(token)")
    cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS "
                "hide_from_university BOOLEAN DEFAULT FALSE")
    # ── Part 2: boshqaruv, hisobotlar, engagement ────────────────────────
    # jurnallar: universitetga bog'lash + so'rov oqimi (journals jadvali
    # app.py init'da yaratiladi — bu ALTERlar xavfsiz)
    try:
        for _col, _typ in (('canonical_institution', 'VARCHAR(500)'),
                           ('submitted_by', 'INTEGER'),
                           ('moderation_status', "VARCHAR(20) DEFAULT 'approved'")):
            cur.execute(f"ALTER TABLE journals ADD COLUMN IF NOT EXISTS {_col} {_typ}")
    except Exception:
        pass  # journals hali yo'q muhit (toza test bazasi) — keyingi so'rovda
    # taklif tokenlari: doktorant roli + email logi
    cur.execute("ALTER TABLE university_invite_tokens "
                "ADD COLUMN IF NOT EXISTS role VARCHAR(30) DEFAULT 'staff'")
    cur.execute("ALTER TABLE university_invite_tokens "
                "ADD COLUMN IF NOT EXISTS email VARCHAR(255)")
    # doktorant ↔ universitet to'g'ridan-to'g'ri bog'lash (taklif qabulida)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS university_doctorant_links (
            id SERIAL PRIMARY KEY,
            user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
            canonical_institution VARCHAR(500) NOT NULL,
            invited_by INTEGER REFERENCES users(id),
            created_at TIMESTAMP DEFAULT NOW(),
            UNIQUE(user_id, canonical_institution)
        )""")
    # kunlik xodim digest'i: xodim darajasida ON/OFF + kunlik dedup logi
    cur.execute("ALTER TABLE university_staff "
                "ADD COLUMN IF NOT EXISTS digest_enabled BOOLEAN DEFAULT TRUE")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS university_digest_log (
            id SERIAL PRIMARY KEY,
            canonical_institution VARCHAR(500) NOT NULL,
            digest_date DATE NOT NULL DEFAULT CURRENT_DATE,
            events_count INTEGER DEFAULT 0,
            sent_at TIMESTAMP DEFAULT NOW(),
            UNIQUE(canonical_institution, digest_date)
        )""")
    # ommaviy profil (litsenziya perki) — canonical nom bilan kalitlangan
    cur.execute("""
        CREATE TABLE IF NOT EXISTS university_public_profiles (
            id SERIAL PRIMARY KEY,
            canonical_institution VARCHAR(500) UNIQUE NOT NULL,
            description TEXT,
            website VARCHAR(500),
            logo_url VARCHAR(500),
            contact_email VARCHAR(255),
            updated_by INTEGER REFERENCES users(id),
            updated_at TIMESTAMP DEFAULT NOW()
        )""")
    # audit jurnali — har bir boshqaruv amali uchun bitta qator
    cur.execute("""
        CREATE TABLE IF NOT EXISTS university_audit_log (
            id SERIAL PRIMARY KEY,
            canonical_institution VARCHAR(500) NOT NULL,
            user_id INTEGER REFERENCES users(id),
            action VARCHAR(200) NOT NULL,
            created_at TIMESTAMP DEFAULT NOW()
        )""")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_univ_audit "
                "ON university_audit_log(canonical_institution, created_at DESC)")
    _schema_ready = True


def _audit(cur, canonical, action):
    """Boshqaruv amalini audit jurnaliga yozadi (hech qachon amalni buzmaydi)."""
    try:
        uid = getattr(current_user, 'id', None)
        cur.execute("INSERT INTO university_audit_log "
                    "(canonical_institution, user_id, action) VALUES (%s, %s, %s)",
                    (canonical, uid, (action or '')[:200]))
    except Exception:
        pass


# ── kirish modeli ────────────────────────────────────────────────────────────

_LIC_COLS = ('id', 'canonical_institution', 'plan', 'valid_until', 'max_staff',
             'created_at', 'notes')


def _license_row(cur, canonical):
    cur.execute(f"SELECT {', '.join(_LIC_COLS)} FROM university_licenses "
                "WHERE canonical_institution = %s", (canonical,))
    r = cur.fetchone()
    if not r:
        return None
    lic = dict(zip(_LIC_COLS, r))
    lic['plan_label'] = PLAN_LABELS.get(lic['plan'], lic['plan'])
    lic['days_left'] = ((lic['valid_until'] - date.today()).days
                        if lic['valid_until'] else None)
    lic['expired'] = lic['days_left'] is not None and lic['days_left'] < 0
    return lic


def _resolve_access():
    """Joriy foydalanuvchining universitet aloqasi:
    dict(canonical, role, title, license) yoki None (xodim emas).
    So'rov ichida g'da keshlanadi."""
    if not getattr(current_user, 'is_authenticated', False):
        return None
    if hasattr(g, '_univer_access'):
        return g._univer_access
    acc = None
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                _ensure_schema(cur)
                cur.execute(
                    "SELECT canonical_institution, role, title "
                    "FROM university_staff "
                    "WHERE user_id = %s AND status = 'active' "
                    "ORDER BY id LIMIT 1", (current_user.id,))
                r = cur.fetchone()
                if r:
                    acc = {'canonical': r[0], 'role': r[1], 'title': r[2] or '',
                           'license': _license_row(cur, r[0])}
            conn.commit()
        finally:
            conn.close()
    except Exception:
        acc = None
    g._univer_access = acc
    return acc


def _is_api():
    return request.path.startswith('/univer/api/')


def get_university_access_or_403():
    """(canonical_institution, role) — joriy foydalanuvchi uchun.
    Faol xodim bo'lmasa 403 (sahifada — /univer landing'ga redirect).
    Litsenziya yo'q/muddati tugagan → do'stona 'Litsenziya muddati tugagan'
    sahifasi (API'da 403 JSON). HAR b2b endpointda chaqiriladi."""
    acc = _resolve_access()
    if not acc:
        if _is_api():
            abort(make_response(jsonify({'success': False, 'error': 'forbidden'}), 403))
        abort(make_response(redirect('/univer')))
    lic = acc['license']
    if lic is None or lic['expired']:
        if _is_api():
            abort(make_response(jsonify(
                {'success': False, 'error': 'license_expired'}), 403))
        abort(make_response(render_template(
            'univer/expired.html', canonical=acc['canonical'],
            latin_name=_latin(acc['canonical']), license=lic), 403))
    return acc['canonical'], acc['role']


# ── muassasa variantlari / moslashtirish to'plami ────────────────────────────

def _latin(name):
    from institutions import transliterate_display
    try:
        return transliterate_display(name or '')
    except Exception:
        return name or ''


def _variants(cur, canonical):
    """Canonical nom → dissertations.muassasa xom (kirill) variantlari.
    Dashboard Muassasa faceti ishlatadigan kengaytma bilan bir xil."""
    cur.execute(
        "SELECT cyrillic_name FROM institution_map "
        "WHERE COALESCE(canonical_name, cyrillic_name) = %s AND is_active = TRUE",
        (canonical,))
    v = [r[0] for r in cur.fetchall() if r[0]]
    if canonical not in v:
        v.append(canonical)
    return v


def _match_set(cur, canonical):
    """olim_profiles.institution ni moslashtirish uchun lowercase to'plam:
    canonical + lotin displayi + barcha xom variantlar (+ lotinlari).
    (Kabinet autocomplete lotin nomni saqlaydi, autofill esa xom kirillni.)"""
    out = set()
    for v in _variants(cur, canonical):
        out.add(v.strip().lower())
        out.add(_latin(v).strip().lower())
    out.add(canonical.strip().lower())
    out.add(_latin(canonical).strip().lower())
    return [s for s in out if s]


# ── universitet-skopli agregat kesh (dashboard) ──────────────────────────────

def _fetchall(cur, sql, params=None):
    cur.execute(sql, params or ())
    return cur.fetchall()


def _doctoral_user_ids(cur, canonical):
    """Universitetga bog'langan (olim_profiles.institution mos) va faol Roadmap
    rejasi YOKI Konstruktor loyihasi bor foydalanuvchi id'lari + profil
    ma'lumotlari. hide_from_university=TRUE bo'lganlar CHIQARILADI."""
    ms = _match_set(cur, canonical)
    rows = _fetchall(cur, """
        SELECT u.id,
               COALESCE(NULLIF(TRIM(op.olim_name), ''), u.username),
               op.photo_url, op.academic_degree, op.ixtisoslik,
               op.advisor_name, op.slug
        FROM olim_profiles op
        JOIN cabinet_users cu ON cu.id = op.cabinet_user_id
        JOIN users u ON LOWER(u.email) = LOWER(cu.email)
        WHERE LOWER(TRIM(COALESCE(op.institution, ''))) = ANY(%s)
          AND COALESCE(u.hide_from_university, FALSE) = FALSE
    """, (ms,))
    users = {}
    for uid, name, photo, degree, ixt, advisor, slug in rows:
        users.setdefault(uid, {
            'user_id': uid, 'name': name, 'photo_url': photo or '',
            'profile_degree': degree or '', 'profile_ixtisoslik': ixt or '',
            'profile_advisor': advisor or '', 'slug': slug or '',
        })
    # Part 2: taklif havolasi orqali to'g'ridan-to'g'ri bog'langan doktorantlar
    # (olim_profiles zanjiri bo'lmasa ham monitoring ko'radi)
    try:
        for uid, name in _fetchall(cur, """
                SELECT u.id, u.username FROM university_doctorant_links l
                JOIN users u ON u.id = l.user_id
                WHERE l.canonical_institution = %s
                  AND COALESCE(u.hide_from_university, FALSE) = FALSE
                """, (canonical,)):
            users.setdefault(uid, {
                'user_id': uid, 'name': name or f'user{uid}', 'photo_url': '',
                'profile_degree': '', 'profile_ixtisoslik': '',
                'profile_advisor': '', 'slug': '',
            })
    except Exception:
        pass
    return users


def _build_univ_stats(canonical):
    """Universitet dashboard agregatlari (15 daqiqa TTL, modul-daraja kesh)."""
    now = time.time()
    hit = _univ_cache.get(canonical)
    if hit and (now - hit['ts']) < _CACHE_TTL:
        return hit['data']

    data = {
        'total_diss': 0, 'active_advisors': 0, 'this_year': 0,
        'claimed_scholars': 0, 'active_doctoral': 0,
        'yearly': [], 'top_specs': [], 'degree_split': {'phd': 0, 'dsc': 0},
        'feed': [],
    }
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                _ensure_schema(cur)
                variants = _variants(cur, canonical)
                this_year = date.today().year
                W = "TRIM(d.muassasa) = ANY(%s)"

                cur.execute(f"SELECT COUNT(*) FROM dissertations d WHERE {W}",
                            (variants,))
                data['total_diss'] = cur.fetchone()[0] or 0

                cur.execute(
                    f"SELECT COUNT(DISTINCT TRIM(d.ilmiy_rahbar)) "
                    f"FROM dissertations d WHERE {W} "
                    f"AND d.ilmiy_rahbar IS NOT NULL AND TRIM(d.ilmiy_rahbar) <> '' "
                    f"AND d.sana ~ '(19|20)\\d{{2}}' "
                    f"AND ({_YEAR_EXPR})::int >= %s",
                    (variants, this_year - 4))
                data['active_advisors'] = cur.fetchone()[0] or 0

                cur.execute(
                    f"SELECT COUNT(*) FROM dissertations d WHERE {W} "
                    f"AND d.sana ~ '(19|20)\\d{{2}}' AND ({_YEAR_EXPR})::int = %s",
                    (variants, this_year))
                data['this_year'] = cur.fetchone()[0] or 0

                ms = _match_set(cur, canonical)
                cur.execute(
                    "SELECT COUNT(*) FROM olim_profiles "
                    "WHERE cabinet_user_id IS NOT NULL "
                    "AND LOWER(TRIM(COALESCE(institution, ''))) = ANY(%s)", (ms,))
                data['claimed_scholars'] = cur.fetchone()[0] or 0

                # faol doktorantlar: bog'langan userlar ∩ (faol reja ∪ loyiha)
                candidates = _doctoral_user_ids(cur, canonical)
                if candidates:
                    ids = list(candidates)
                    cur.execute(
                        "SELECT DISTINCT user_id FROM roadmap_plans "
                        "WHERE is_active AND user_id = ANY(%s)", (ids,))
                    active = {r[0] for r in cur.fetchall()}
                    cur.execute(
                        "SELECT DISTINCT owner_id FROM diss_projects "
                        "WHERE status <> 'archived' AND owner_id = ANY(%s)", (ids,))
                    active |= {r[0] for r in cur.fetchall()}
                    data['active_doctoral'] = len(active)

                # yillik himoyalar (oxirgi 10 yil)
                for yr, cnt in _fetchall(cur,
                        f"SELECT ({_YEAR_EXPR})::int AS yr, COUNT(*) "
                        f"FROM dissertations d WHERE {W} "
                        f"AND d.sana ~ '(19|20)\\d{{2}}' "
                        f"GROUP BY 1 ORDER BY 1", (variants,)):
                    if yr and yr >= this_year - 9:
                        data['yearly'].append({'year': yr, 'count': cnt})

                # top-10 ixtisoslik
                for code, nm, cnt in _fetchall(cur,
                        f"SELECT TRIM(d.ixtisoslik), MAX(d.ixtisoslik_nomi), COUNT(*) "
                        f"FROM dissertations d WHERE {W} "
                        f"AND d.ixtisoslik IS NOT NULL AND TRIM(d.ixtisoslik) <> '' "
                        f"GROUP BY 1 ORDER BY 3 DESC LIMIT 10", (variants,)):
                    data['top_specs'].append(
                        {'code': code, 'name': (nm or '').strip(), 'count': cnt})

                cur.execute(
                    f"SELECT COUNT(*) FILTER (WHERE {_PHD_LIKE}), "
                    f"       COUNT(*) FILTER (WHERE {_DSC_LIKE}) "
                    f"FROM dissertations d WHERE {W}", (variants,))
                r = cur.fetchone() or (0, 0)
                data['degree_split'] = {'phd': r[0] or 0, 'dsc': r[1] or 0}

                # ── jonli lenta ──
                feed = []
                for olim, sana, daraja in _fetchall(cur,
                        f"SELECT d.olim, d.sana, d.daraja FROM dissertations d "
                        f"WHERE {W} ORDER BY d.id DESC LIMIT 5", (variants,)):
                    from data import clean_olim_name
                    feed.append({'icon': '🎓', 'ts': None,
                                 'text': f"Yangi himoya: {clean_olim_name(olim or '')} "
                                         f"({(daraja or '').strip()}, {(sana or '').strip()})"})
                try:
                    for name, ts in _fetchall(cur,
                            "SELECT TRIM(olim_name), created_at FROM olim_profiles "
                            "WHERE cabinet_user_id IS NOT NULL "
                            "AND LOWER(TRIM(COALESCE(institution, ''))) = ANY(%s) "
                            "ORDER BY created_at DESC NULLS LAST LIMIT 3", (ms,)):
                        from data import clean_olim_name
                        feed.append({'icon': '👤', 'ts': str(ts)[:10] if ts else None,
                                     'text': f"Yangi tasdiqlangan profil: {clean_olim_name(name or '')}"})
                except Exception:
                    pass
                if candidates:
                    try:
                        for uid, ts in _fetchall(cur, """
                                SELECT rp.user_id, pub.created_at
                                FROM roadmap_publications pub
                                JOIN roadmap_plans rp ON rp.id = pub.plan_id
                                WHERE rp.user_id = ANY(%s)
                                ORDER BY pub.created_at DESC LIMIT 3
                                """, (list(candidates),)):
                            nm = candidates.get(uid, {}).get('name', 'Doktorant')
                            feed.append({'icon': '📄', 'ts': str(ts)[:10] if ts else None,
                                         'text': f"{nm} yangi nashr qo'shdi"})
                    except Exception:
                        pass
                data['feed'] = feed[:10]
            conn.commit()
        finally:
            conn.close()
    except Exception:
        if hit:
            return hit['data']
    _univ_cache[canonical] = {'data': data, 'ts': now}
    return data


# ── workspace umumiy kontekst ────────────────────────────────────────────────

def _uni_logo(cur, canonical):
    """199-universitet admin jadvalidan logo (nom mosligi bo'yicha) yoki ''."""
    try:
        latin = _latin(canonical)
        cur.execute(
            "SELECT logo_url FROM universities "
            "WHERE LOWER(name) IN (%s, %s) AND logo_url IS NOT NULL "
            "AND TRIM(logo_url) <> '' LIMIT 1",
            (canonical.lower(), latin.lower()))
        r = cur.fetchone()
        return r[0] if r else ''
    except Exception:
        return ''


def _workspace_ctx(section):
    """Har workspace sahifasi uchun umumiy kontekst (guard + header)."""
    canonical, role = get_university_access_or_403()
    acc = g._univer_access
    logo = ''
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                logo = _uni_logo(cur, canonical)
            conn.commit()
        finally:
            conn.close()
    except Exception:
        pass
    return {
        'canonical': canonical,
        'latin_name': _latin(canonical),
        'role': role,
        'role_label': ROLE_LABELS.get(role, role),
        'license': acc['license'],
        'logo_url': logo,
        'section': section,
    }


# ═════════════════════════════ SAHIFALAR ════════════════════════════════════

@univer_bp.route('/univer')
def univer_home():
    """Faol xodim → workspace (dashboard); litsenziya tugagan → do'stona
    sahifa; boshqalar (mehmon ham) → B2B marketing/landing sahifasi."""
    acc = _resolve_access()
    if acc:
        lic = acc['license']
        if lic is None or lic['expired']:
            return render_template('univer/expired.html',
                                   canonical=acc['canonical'],
                                   latin_name=_latin(acc['canonical']),
                                   license=lic), 403
        ctx = _workspace_ctx('dashboard')
        stats = _build_univ_stats(ctx['canonical'])
        return render_template('univer/dashboard.html', stats=stats, **ctx)
    return render_template('univer/landing.html')


@univer_bp.route('/univer/olimlar')
@login_required
def univer_olimlar():
    ctx = _workspace_ctx('olimlar')
    # facetlar sahifa renderida — universitet olimlaridan jonli hisoblanadi
    from blueprints.olimlar_catalog import _build_cache, _compute_facets
    cache = _build_cache()
    items = [s for s in cache['data'] if ctx['canonical'] in s['institutions']]
    facets = _compute_facets(items)
    return render_template('univer/olimlar.html', total_scholars=len(items),
                           facets=facets, **ctx)


@univer_bp.route('/univer/dissertatsiyalar')
@login_required
def univer_dissertations():
    ctx = _workspace_ctx('dissertatsiyalar')
    return render_template('univer/dissertations.html', **ctx)


@univer_bp.route('/univer/doktorantlar')
@login_required
def univer_doktorantlar():
    ctx = _workspace_ctx('doktorantlar')
    rows = _doctoral_rows(ctx['canonical'])
    advisors = _advisor_load(ctx['canonical'])
    invites = _doctorant_invites(ctx['canonical'])
    return render_template('univer/doktorantlar.html', rows=rows,
                           advisors=advisors, invites=invites, **ctx)


@univer_bp.route('/univer/sozlamalar')
@login_required
def univer_sozlamalar():
    """Part 2: barcha xodimlarga ochiq — o'z digest sozlamasi + litsenziya
    ma'lumoti. Xodimlar ro'yxati va audit jurnali faqat owner ko'rinishida
    (server tomonda ham: boshqaruv API'lari owner-only)."""
    ctx = _workspace_ctx('sozlamalar')
    staff, audit, digest_on = [], [], True
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                _ensure_schema(cur)
                cur.execute("SELECT COALESCE(digest_enabled, TRUE) "
                            "FROM university_staff WHERE user_id = %s "
                            "AND canonical_institution = %s",
                            (current_user.id, ctx['canonical']))
                r = cur.fetchone()
                digest_on = bool(r[0]) if r else True
                if ctx['role'] == 'owner':
                    cur.execute("""
                        SELECT s.id, s.user_id, u.username, u.email, s.role,
                               s.title, s.status, s.created_at
                        FROM university_staff s JOIN users u ON u.id = s.user_id
                        WHERE s.canonical_institution = %s
                        ORDER BY (s.role = 'owner') DESC, s.id
                    """, (ctx['canonical'],))
                    staff = [dict(zip(('id', 'user_id', 'username', 'email', 'role',
                                       'title', 'status', 'created_at'), r))
                             for r in cur.fetchall()]
                    cur.execute("""
                        SELECT a.action, COALESCE(u.username, '—'), a.created_at
                        FROM university_audit_log a
                        LEFT JOIN users u ON u.id = a.user_id
                        WHERE a.canonical_institution = %s
                        ORDER BY a.created_at DESC LIMIT 50
                    """, (ctx['canonical'],))
                    audit = [{'action': r[0], 'username': r[1],
                              'created_at': str(r[2])[:16] if r[2] else ''}
                             for r in cur.fetchall()]
            conn.commit()
        finally:
            conn.close()
    except Exception:
        staff, audit = [], []
    for s in staff:
        s['role_label'] = ROLE_LABELS.get(s['role'], s['role'])
    lic = ctx['license']
    return render_template('univer/sozlamalar.html', staff=staff, audit=audit,
                           digest_on=digest_on,
                           max_staff=(lic['max_staff'] if lic else DEFAULT_MAX_STAFF),
                           **ctx)


# ═════════════════════ API: OLIMLARIMIZ (katalog qayta ishlatiladi) ══════════

@univer_bp.route('/univer/api/olimlar')
@login_required
def api_univer_olimlar():
    """Olimlar katalogining universitet-skopli varianti: keshlangan agregat
    ro'yxat Python'da institutsiya bo'yicha filtrlab beriladi (og'ir so'rovlar
    TAKRORLANMAYDI). Viloyat faceti yo'q (universitet bitta joyda)."""
    canonical, _role = get_university_access_or_403()
    from blueprints.olimlar_catalog import (
        _build_cache, _apply_filters, _sort_items, _scholar_public,
        _compute_facets, _user_follows, PER_PAGE_CARDS)
    cache = _build_cache()
    items = [s for s in cache['data'] if canonical in s['institutions']]
    f = request.args.to_dict()  # MultiDict → flat dict (birinchi qiymatlar)
    f.pop('viloyat', None)      # region qulfi: universitet manzili o'zgarmas
    f.pop('muassasa', None)     # muassasa qulfi: faqat shu universitet
    items = _apply_filters(items, f)
    sort = (request.args.get('sort') or 'students').strip()
    items = _sort_items(items, sort)
    total = len(items)
    follows = _user_follows()
    page = max(1, request.args.get('page', 1, type=int))
    pages = max(1, (total + PER_PAGE_CARDS - 1) // PER_PAGE_CARDS)
    page = min(page, pages)
    start = (page - 1) * PER_PAGE_CARDS
    return jsonify({
        'ok': True, 'total': total, 'sort': sort, 'page': page, 'pages': pages,
        'scholars': [_scholar_public(s, follows)
                     for s in items[start:start + PER_PAGE_CARDS]],
        'facets': _compute_facets(items),
    })


# ═════════════════ API: DISSERTATSIYALAR (skoplangan jadval) ═════════════════

def _uni_diss_where(cur, canonical, a):
    """Universitetga qulflangan WHERE: mavjud _build_filter_clause (qidiruv
    pillari + translit) + muassasa variantlari. (clause, params) qaytaradi."""
    from data import _build_filter_clause
    scope = (a.get('scope') or 'all').strip()
    if scope not in ('all', 'olim', 'rahbar', 'mavzu', 'opponent'):
        scope = 'all'
    daraja = (a.get('daraja') or '').strip()
    if daraja.upper() not in ('PHD', 'DSC'):
        daraja = ''
    yil = (a.get('yil') or '').strip()
    if not re.match(r'^\d{4}$', yil):
        yil = ''
    clause, params = _build_filter_clause(
        (a.get('q') or '').strip(), daraja, '', '',
        sana_yil=yil, scope=scope)
    variants = _variants(cur, canonical)
    uni = "TRIM(d.muassasa) = ANY(%s)"
    if clause:
        clause = clause + " AND " + uni
    else:
        clause = " WHERE " + uni
    return clause, params + [variants]


@univer_bp.route('/univer/api/dissertations')
@login_required
def api_univer_dissertations():
    canonical, _role = get_university_access_or_403()
    from data import _query_rows, _SANA_ORDER_DESC
    a = request.args
    try:
        page = max(1, int(a.get('page', 1)))
    except ValueError:
        page = 1
    per_page = 25
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                clause, params = _uni_diss_where(cur, canonical, a)
                cur.execute(f"SELECT COUNT(*) FROM dissertations d{clause}", params)
                total = cur.fetchone()[0] or 0
            conn.commit()
        finally:
            conn.close()
        pages = max(1, (total + per_page - 1) // per_page)
        page = min(page, pages)
        rows = _query_rows(
            'SELECT d.id, d.oak_id, d.sana AS "Sana", d.daraja AS "Daraja", '
            'd.olim AS "Olim", d.mavzu AS "Mavzu", d.ixtisoslik AS "Ixtisoslik", '
            'd.muassasa AS "Muassasa", d.ilmiy_rahbar AS "Ilmiy_rahbar", '
            'd.link AS "Link" '
            f'FROM dissertations d{clause} ORDER BY {_SANA_ORDER_DESC} '
            'LIMIT %s OFFSET %s',
            params + [per_page, (page - 1) * per_page])
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500
    return jsonify({'ok': True, 'records': rows, 'total': total,
                    'page': page, 'pages': pages, 'per_page': per_page})


@univer_bp.route('/univer/api/dissertations/export.csv')
@login_required
def api_univer_export():
    """Universitetning O'Z dissertatsiyalari CSV eksporti — TO'LIQ, 50 qatorlik
    cap YO'Q (bu ularning ma'lumoti; sotuv nuqtasi). 30s rate-limit."""
    canonical, _role = get_university_access_or_403()
    now = time.time()
    if now - _export_last.get(current_user.id, 0) < 30:
        return Response("Eksport tayyorlanmoqda, biroz kuting", status=429,
                        mimetype='text/plain; charset=utf-8')
    _export_last[current_user.id] = now
    from data import _SANA_ORDER_DESC
    conn = _conn()
    try:
        with conn.cursor() as cur:
            clause, params = _uni_diss_where(cur, canonical, request.args)
    finally:
        conn.close()

    def generate():
        import csv
        import io
        buf = io.StringIO()
        writer = csv.writer(buf)
        yield '﻿'  # UTF-8 BOM — Excel
        writer.writerow(['sana', 'olim', 'mavzu', 'daraja', 'ixtisoslik',
                         'ixtisoslik_nomi', 'fan_tarmoqi', 'ilmiy_rahbar',
                         'muassasa', 'ilmiy_kengash', 'link'])
        yield buf.getvalue()
        buf.seek(0); buf.truncate(0)
        conn2 = _conn()
        try:
            with conn2.cursor() as cur:
                cur.execute(
                    "SELECT d.sana, d.olim, d.mavzu, d.daraja, d.ixtisoslik, "
                    "COALESCE(d.ixtisoslik_nomi, ''), COALESCE(d.fan_tarmoqi, ''), "
                    "d.ilmiy_rahbar, d.muassasa, COALESCE(d.ilmiy_kengash, ''), "
                    f"COALESCE(d.link, '') FROM dissertations d{clause} "
                    f"ORDER BY {_SANA_ORDER_DESC}", params)
                for row in cur.fetchall():
                    writer.writerow([(x or '').strip() if isinstance(x, str)
                                     else (x or '') for x in row])
                    yield buf.getvalue()
                    buf.seek(0); buf.truncate(0)
        finally:
            conn2.close()
    return Response(generate(), mimetype='text/csv; charset=utf-8',
                    headers={'Content-Disposition':
                             'attachment; filename="universitet-dissertatsiyalar.csv"'})


# ═══════════════ DOKTORANTLAR MONITORINGI (faqat metrikalar!) ════════════════

def _oak_req(degree_type):
    from blueprints.roadmap import OAK_REQUIREMENTS
    return OAK_REQUIREMENTS.get(degree_type, OAK_REQUIREMENTS['phd'])


def _default_word_target():
    from blueprints.roadmap import OAK_STRUCTURE
    return sum(t for _n, _b, t in OAK_STRUCTURE)


def _doctoral_rows(canonical):
    """Doktorantlar jadvali qatorlari — FAQAT METRIKALAR (matn yo'q, havola
    yo'q). Universitetga bog'langan foydalanuvchilardan faol reja yoki
    Konstruktor loyihasi borlari."""
    rows = []
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                _ensure_schema(cur)
                users = _doctoral_user_ids(cur, canonical)
                if not users:
                    return []
                ids = list(users)

                plans = {}
                for r in _fetchall(cur, """
                        SELECT user_id, id, degree_type, specialty_code,
                               target_defense_date, diss_project_id
                        FROM roadmap_plans
                        WHERE is_active AND user_id = ANY(%s)""", (ids,)):
                    plans[r[0]] = {'plan_id': r[1], 'degree_type': r[2],
                                   'specialty': r[3] or '',
                                   'defense_date': r[4],
                                   'diss_project_id': r[5]}

                # nashr sonlari (chop etilganlari) — reja orqali
                pub_done = {}
                for uid, ptype, cnt in _fetchall(cur, """
                        SELECT rp.user_id, pub.pub_type, COUNT(*)
                        FROM roadmap_publications pub
                        JOIN roadmap_plans rp ON rp.id = pub.plan_id
                        WHERE rp.user_id = ANY(%s) AND rp.is_active
                          AND pub.status = 'chop_etilgan'
                        GROUP BY 1, 2""", (ids,)):
                    pub_done.setdefault(uid, {})[ptype] = cnt

                # loyiha so'z metrikalari (faqat sonlar: SUM(word_count))
                proj = {}   # owner_id → (project_id, words, target, updated_at)
                for pid, owner, upd, words, target in _fetchall(cur, """
                        SELECT p.id, p.owner_id, p.updated_at,
                               COALESCE(SUM(b.word_count), 0),
                               COALESCE(SUM(b.word_target), 0)
                        FROM diss_projects p
                        LEFT JOIN dissertation_blocks b ON b.dissertation_id = p.id
                        WHERE p.owner_id = ANY(%s) AND p.status <> 'archived'
                        GROUP BY p.id, p.owner_id, p.updated_at""", (ids,)):
                    linked = plans.get(owner, {}).get('diss_project_id')
                    cur_best = proj.get(owner)
                    # rejaga ulangan loyiha ustuvor; aks holda eng yangisi
                    if pid == linked or cur_best is None or (
                            cur_best[0] != linked and upd and
                            (cur_best[3] is None or upd > cur_best[3])):
                        proj[owner] = (pid, words or 0, target or 0, upd)

                # rahbar (advisor_links accepted) → username
                advisors = {}
                for sid, aname in _fetchall(cur, """
                        SELECT al.student_id, MAX(u2.username)
                        FROM advisor_links al JOIN users u2 ON u2.id = al.advisor_id
                        WHERE al.student_id = ANY(%s) AND al.status = 'accepted'
                        GROUP BY al.student_id""", (ids,)):
                    advisors[sid] = aname or ''

                # oxirgi faollik (risk bayrog'i uchun): blok saqlashlari +
                # nashr qo'shishlari — faqat vaqt belgilari, kontent EMAS
                last_act = {}
                for uid_, ts in _fetchall(cur, """
                        SELECT p.owner_id, MAX(b.updated_at)
                        FROM diss_projects p
                        JOIN dissertation_blocks b ON b.dissertation_id = p.id
                        WHERE p.owner_id = ANY(%s) AND p.status <> 'archived'
                        GROUP BY p.owner_id""", (ids,)):
                    if ts:
                        last_act[uid_] = ts
                for uid_, ts in _fetchall(cur, """
                        SELECT rp.user_id, MAX(pub.created_at)
                        FROM roadmap_publications pub
                        JOIN roadmap_plans rp ON rp.id = pub.plan_id
                        WHERE rp.user_id = ANY(%s)
                        GROUP BY rp.user_id""", (ids,)):
                    if ts and (uid_ not in last_act or ts > last_act[uid_]):
                        last_act[uid_] = ts
            conn.commit()
        finally:
            conn.close()
    except Exception:
        return []

    default_target = _default_word_target()
    for uid, u in users.items():
        plan = plans.get(uid)
        p = proj.get(uid)
        if not plan and not p:
            continue        # platformada faol emas — jadvalga kirmaydi
        degree = (plan or {}).get('degree_type') or (
            'dsc' if 'dsc' in (u['profile_degree'] or '').lower() else 'phd')
        req = _oak_req(degree)
        done = pub_done.get(uid, {})
        oak_req = req.get('maqola_milliy', 0)
        int_req = req.get('maqola_xalqaro', 0)
        words = p[1] if p else 0
        target = (p[2] if p and p[2] else default_target)
        word_pct = min(round(words * 100 / target), 100) if target else 0
        defense = (plan or {}).get('defense_date')
        days_left = (defense - date.today()).days if defense else None
        pub_pct_parts = []
        if oak_req:
            pub_pct_parts.append(min(done.get('maqola_milliy', 0) / oak_req, 1))
        if int_req:
            pub_pct_parts.append(min(done.get('maqola_xalqaro', 0) / int_req, 1))
        pub_pct = round(sum(pub_pct_parts) * 100 / len(pub_pct_parts)) \
            if pub_pct_parts else 0
        overall = round((word_pct + pub_pct) / 2)
        if overall <= 0:
            holat, holat_class = 'Boshlanmagan', 'muted'
        elif overall >= 70:
            holat, holat_class = 'Yakuniy bosqich', 'good'
        else:
            holat, holat_class = 'Faol', 'active'
        # ── risk bayroqlari (Part 2) ──
        # 🔴 himoya < 90 kun VA nashrlar to'liq emas; 🟡 30+ kun faollik yo'q
        pubs_incomplete = (done.get('maqola_milliy', 0) < oak_req
                           or done.get('maqola_xalqaro', 0) < int_req)
        act = last_act.get(uid)
        act_date = act.date() if hasattr(act, 'date') else act
        idle_days = (date.today() - act_date).days if act_date else None
        risk, risk_note = '', ''
        if days_left is not None and days_left < 90 and pubs_incomplete:
            risk = 'red'
            risk_note = "Himoyagacha 90 kundan kam, nashrlar to'liq emas"
        elif idle_days is None or idle_days >= 30:
            risk = 'yellow'
            risk_note = ("30+ kun faollik yo'q" if idle_days is not None
                         else "Platformada faollik qayd etilmagan")
        rows.append({
            'user_id': uid, 'name': u['name'], 'photo_url': u['photo_url'],
            'degree': degree.upper() if degree != 'magistr' else 'Magistr',
            'specialty': (plan or {}).get('specialty') or u['profile_ixtisoslik'],
            'advisor': advisors.get(uid) or u['profile_advisor'] or '—',
            'oak_done': done.get('maqola_milliy', 0), 'oak_req': oak_req,
            'int_done': done.get('maqola_xalqaro', 0), 'int_req': int_req,
            'words': words, 'word_target': target, 'word_pct': word_pct,
            'defense_date': defense.isoformat() if defense else '',
            'days_left': days_left,
            'holat': holat, 'holat_class': holat_class, 'overall': overall,
            'risk': risk, 'risk_note': risk_note, 'idle_days': idle_days,
        })
    rows.sort(key=lambda r: -r['overall'])
    return rows


def _advisor_load(canonical):
    """'Rahbarlar yuki' sub-tab: universitet rahbarlari — platformadagi joriy
    shogirdlar soni (advisor_links accepted) + tarixiy (dissertations,
    olimlar_catalog keshidan). >5 joriy shogird — ortiqcha yuklangan."""
    current_load = {}
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                _ensure_schema(cur)
                users = _doctoral_user_ids(cur, canonical)
                if users:
                    for name, cnt in _fetchall(cur, """
                            SELECT u2.username, COUNT(*)
                            FROM advisor_links al
                            JOIN users u2 ON u2.id = al.advisor_id
                            WHERE al.student_id = ANY(%s) AND al.status = 'accepted'
                            GROUP BY u2.username""", (list(users),)):
                        current_load[(name or '').strip().lower()] = \
                            {'name': name, 'active': cnt}
            conn.commit()
        finally:
            conn.close()
    except Exception:
        pass
    out = []
    try:
        from blueprints.olimlar_catalog import _build_cache
        cache = _build_cache()
        for s in cache['data']:
            if canonical not in s['institutions']:
                continue
            key = s['display'].strip().lower()
            cur_row = current_load.pop(key, None) or \
                current_load.pop(s['name'].strip().lower(), None)
            active = cur_row['active'] if cur_row else 0
            out.append({'name': s['display'], 'full_name': s['name'],
                        'historical': s['total_students'],
                        'last_year': s['last_year'], 'active': active,
                        'overloaded': active > 5})
    except Exception:
        pass
    # faqat platformada joriy shogirdi bor, lekin tarixiy bazada yo'q rahbarlar
    for row in current_load.values():
        out.append({'name': row['name'], 'full_name': row['name'],
                    'historical': 0, 'last_year': None,
                    'active': row['active'], 'overloaded': row['active'] > 5})
    out.sort(key=lambda r: (-r['active'], -r['historical']))
    return out[:100]


@univer_bp.route('/univer/api/doktorant/<int:uid>')
@login_required
def api_doktorant_detail(uid):
    """Doktorant detal draweri — FAQAT METRIKALAR: progress halqasi, bob
    so'z barlari (sonlar), nashrlar ro'yxati (sarlavha+jurnal+yil — CV-ochiq
    ma'lumot), jadval bosqichlari. Konstruktor muharririga HAVOLA YO'Q,
    block_id/diss_id javobdan olib tashlanadi, kontent o'qilmaydi."""
    canonical, _role = get_university_access_or_403()
    from blueprints.roadmap import (_fetch_plan, _pub_progress, _diss_progress,
                                    _timeline, _chapter_states, _readiness,
                                    PUB_TYPE_LABELS, PUB_STATUS_LABELS)
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                _ensure_schema(cur)
                users = _doctoral_user_ids(cur, canonical)
                if uid not in users:
                    # boshqa universitet doktoranti yoki yashiringan — 403
                    return jsonify({'success': False, 'error': 'forbidden'}), 403
                u = users[uid]
                plan = _fetch_plan(cur, uid)
                pub_rows, pub_pct = ([], 0)
                pubs, timeline, diss = [], [], None
                if plan:
                    pub_rows, pub_pct = _pub_progress(cur, plan)
                    cur.execute(
                        "SELECT pub_type, title, venue, status, year "
                        "FROM roadmap_publications WHERE plan_id = %s "
                        "ORDER BY created_at DESC LIMIT 50", (plan['id'],))
                    pubs = [{'type': PUB_TYPE_LABELS.get(r[0], r[0]),
                             'title': r[1], 'venue': r[2] or '',
                             'status': PUB_STATUS_LABELS.get(r[3], r[3]),
                             'year': r[4]} for r in cur.fetchall()]
                    # o'qish-faqat: reja loyihaga ulanmagan bo'lsa xotirada
                    # bog'laymiz (_link_diss_project YOZADI — ishlatilmaydi)
                    if not plan.get('diss_project_id'):
                        cur.execute(
                            "SELECT id FROM diss_projects WHERE owner_id = %s "
                            "AND status <> 'archived' "
                            "ORDER BY updated_at DESC LIMIT 1", (uid,))
                        r = cur.fetchone()
                        if r:
                            plan['diss_project_id'] = r[0]
                    diss = _diss_progress(cur, plan)
                    cur.execute("SELECT phase_key, due_date, is_done "
                                "FROM roadmap_milestones WHERE plan_id = %s",
                                (plan['id'],))
                    overrides = {r[0]: (r[1], r[2]) for r in cur.fetchall()}
                    timeline = _timeline(plan, overrides, _chapter_states(diss))
            conn.commit()
        finally:
            conn.close()
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

    # MAXFIYLIK: bo'limlardan faqat sonlar; block_id/diss_id chiqarilmaydi
    sections = []
    if diss:
        sections = [{'title': s['title'], 'words': s['words'],
                     'target': s['target'], 'pct': s['pct']}
                    for s in diss['sections']]
    tl = [{'key': t['key'], 'label': t['label'], 'start': t['start'],
           'end': t['end'], 'state': t['state'], 'days_left': t['days_left']}
          for t in timeline]
    diss_pct = diss['pct'] if diss else None
    return jsonify({
        'success': True,
        'name': u['name'], 'photo_url': u['photo_url'],
        'degree': (plan['degree_label'] if plan else u['profile_degree']),
        'specialty': (plan.get('specialty_code') if plan else '') or u['profile_ixtisoslik'],
        'defense_date': (plan['target_defense_date'].isoformat()
                         if plan and plan['target_defense_date'] else ''),
        'readiness': _readiness(pub_pct, diss_pct),
        'pub_pct': pub_pct,
        'pub_requirements': pub_rows,
        'publications': pubs,
        'sections': sections,
        'total_words': diss['total_words'] if diss else 0,
        'total_target': diss['total_target'] if diss else 0,
        'diss_pct': diss_pct or 0,
        'timeline': tl,
    })


# ═══════════════════ SOZLAMALAR API (owner roli) ═════════════════════════════

def _owner_or_403():
    canonical, role = get_university_access_or_403()
    if role != 'owner':
        abort(make_response(jsonify({'success': False,
                                     'error': "Faqat egasi (owner) uchun"}), 403))
    return canonical


def _manager_or_403():
    """Boshqaruv amallari (Part 2): owner yoki staff. Viewer → 403 (barcha
    POST'larda server tomonda ham tekshiriladi)."""
    canonical, role = get_university_access_or_403()
    if role not in ('owner', 'staff'):
        abort(make_response(jsonify(
            {'success': False,
             'error': "Ko'ruvchi (viewer) roli o'zgartirish kirita olmaydi"}), 403))
    return canonical


def _staff_count(cur, canonical):
    cur.execute("SELECT COUNT(*) FROM university_staff "
                "WHERE canonical_institution = %s AND status = 'active'",
                (canonical,))
    return cur.fetchone()[0] or 0


def _add_staff(cur, canonical, ident, role, title, invited_by, max_staff):
    """Foydalanuvchini xodim qilib qo'shadi. Muvaffaqiyatda (True, msg),
    aks holda (False, xato)."""
    if role not in ('staff', 'viewer'):
        role = 'staff'
    cur.execute("SELECT id, username FROM users "
                "WHERE LOWER(username) = LOWER(%s) OR LOWER(email) = LOWER(%s)",
                (ident, ident))
    target = cur.fetchone()
    if not target:
        return False, "Foydalanuvchi topilmadi. Username yoki emailni tekshiring."
    tid = target[0]
    cur.execute("SELECT id, status FROM university_staff "
                "WHERE user_id = %s AND canonical_institution = %s",
                (tid, canonical))
    ex = cur.fetchone()
    if ex and ex[1] == 'active':
        return False, "Bu foydalanuvchi allaqachon xodimlar ro'yxatida."
    if _staff_count(cur, canonical) >= max_staff:
        return False, (f"Xodimlar soni to'lgan (maksimal {max_staff} ta). "
                       "Tarifni oshirish uchun biz bilan bog'laning.")
    if ex:
        cur.execute("UPDATE university_staff SET status = 'active', role = %s, "
                    "title = %s, invited_by = %s WHERE id = %s",
                    (role, title or None, invited_by, ex[0]))
    else:
        cur.execute("""
            INSERT INTO university_staff
                (user_id, canonical_institution, role, title, invited_by)
            VALUES (%s, %s, %s, %s, %s)""",
                    (tid, canonical, role, title or None, invited_by))
    _nav_cache.pop(tid, None)
    return True, f"{target[1]} xodimlar ro'yxatiga qo'shildi."


@univer_bp.route('/univer/api/staff/add', methods=['POST'])
@csrf.exempt
@login_required
def api_staff_add():
    canonical = _owner_or_403()
    data = request.get_json(silent=True) or {}
    ident = (data.get('username_or_email') or '').strip()
    if not ident:
        return jsonify({'success': False, 'error': "Username yoki email kiriting"}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            lic = _license_row(cur, canonical)
            max_staff = lic['max_staff'] if lic else DEFAULT_MAX_STAFF
            ok, msg = _add_staff(cur, canonical, ident,
                                 (data.get('role') or 'staff').strip(),
                                 (data.get('title') or '').strip()[:200],
                                 current_user.id, max_staff)
            if ok:
                _audit(cur, canonical, f"Xodim qo'shildi: {ident}")
        conn.commit()
        return (jsonify({'success': True, 'message': msg}) if ok
                else (jsonify({'success': False, 'error': msg}), 400))
    finally:
        conn.close()


def _staff_row_for_update(cur, sid, canonical):
    cur.execute("SELECT id, user_id, role FROM university_staff "
                "WHERE id = %s AND canonical_institution = %s", (sid, canonical))
    r = cur.fetchone()
    if not r:
        abort(make_response(jsonify({'success': False, 'error': 'topilmadi'}), 404))
    if r[2] == 'owner':
        # owner qatorini faqat sayt admini (litsenziya sahifasi) boshqaradi
        abort(make_response(jsonify({'success': False,
                                     'error': "Owner qatorini o'zgartirib bo'lmaydi"}), 400))
    return r


@univer_bp.route('/univer/api/staff/<int:sid>/role', methods=['POST'])
@csrf.exempt
@login_required
def api_staff_role(sid):
    canonical = _owner_or_403()
    role = ((request.get_json(silent=True) or {}).get('role') or '').strip()
    if role not in ('staff', 'viewer'):
        return jsonify({'success': False, 'error': "Noto'g'ri rol"}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            _staff_row_for_update(cur, sid, canonical)
            cur.execute("UPDATE university_staff SET role = %s WHERE id = %s",
                        (role, sid))
            _audit(cur, canonical, f"Xodim roli o'zgartirildi (#{sid} → {role})")
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


@univer_bp.route('/univer/api/staff/<int:sid>/suspend', methods=['POST'])
@csrf.exempt
@login_required
def api_staff_suspend(sid):
    canonical = _owner_or_403()
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            row = _staff_row_for_update(cur, sid, canonical)
            cur.execute("UPDATE university_staff SET status = "
                        "CASE WHEN status = 'active' THEN 'suspended' "
                        "ELSE 'active' END WHERE id = %s RETURNING status", (sid,))
            new_status = cur.fetchone()[0]
            _nav_cache.pop(row[1], None)
            _audit(cur, canonical, f"Xodim holati o'zgartirildi (#{sid} → {new_status})")
        conn.commit()
        return jsonify({'success': True, 'status': new_status})
    finally:
        conn.close()


# ═══════════ MAXFIYLIK: "Universitetim faolligimni ko'rmasin" ═══════════════

@univer_bp.route('/univer/api/privacy', methods=['GET', 'POST'])
@csrf.exempt
@login_required
def api_privacy():
    """Doktorant o'zini universitet monitoringidan yashirishi mumkin
    (users.hide_from_university). Kabinet sozlamalaridagi checkbox shu yerga
    yozadi. Bu HAR QANDAY login foydalanuvchi uchun (xodim roli shart emas)."""
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            if request.method == 'POST':
                hidden = bool((request.get_json(silent=True) or {}).get('hidden'))
                cur.execute("UPDATE users SET hide_from_university = %s "
                            "WHERE id = %s", (hidden, current_user.id))
                conn.commit()
                return jsonify({'success': True, 'hidden': hidden})
            cur.execute("SELECT COALESCE(hide_from_university, FALSE) "
                        "FROM users WHERE id = %s", (current_user.id,))
            r = cur.fetchone()
        conn.commit()
        return jsonify({'success': True, 'hidden': bool(r and r[0])})
    finally:
        conn.close()


# ═══════════════ XODIM TAKLIF HAVOLASI (pilot onboarding) ════════════════════

@univer_bp.route('/univer/invite/<token>')
def univer_invite_landing(token):
    """Taklif havolasi: login bo'lmagan → login (next bilan); login bo'lgan →
    tasdiqlash sahifasi. advisor_invite_tokens oqimi naqshi."""
    if not current_user.is_authenticated:
        return redirect(url_for('auth.login', next=request.path))
    row = None
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                _ensure_schema(cur)
                cur.execute("""
                    SELECT t.id, t.license_id, t.expires_at, t.used_by,
                           l.canonical_institution, COALESCE(t.role, 'staff')
                    FROM university_invite_tokens t
                    JOIN university_licenses l ON l.id = t.license_id
                    WHERE t.token = %s""", (token,))
                row = cur.fetchone()
            conn.commit()
        finally:
            conn.close()
    except Exception:
        row = None
    if not row:
        return render_template('univer/invite.html', error='Taklif havolasi '
                               'topilmadi yoki bekor qilingan.'), 404
    _tid, _lid, expires_at, used_by, canonical, role = row
    if used_by:
        return render_template('univer/invite.html', error='Bu taklif havolasi '
                               'allaqachon ishlatilgan.'), 410
    if expires_at and expires_at < datetime.now():
        return render_template('univer/invite.html', error='Taklif havolasi '
                               'muddati tugagan (havola 7 kun amal qiladi).'), 410
    return render_template('univer/invite.html', token=token, role=role,
                           canonical=canonical, latin_name=_latin(canonical))


@univer_bp.route('/univer/invite/<token>/respond', methods=['POST'])
@csrf.exempt
@login_required
def univer_invite_respond(token):
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("""
                SELECT t.id, t.expires_at, t.used_by, t.created_by,
                       l.canonical_institution, l.max_staff, l.valid_until,
                       COALESCE(t.role, 'staff')
                FROM university_invite_tokens t
                JOIN university_licenses l ON l.id = t.license_id
                WHERE t.token = %s FOR UPDATE""", (token,))
            row = cur.fetchone()
            if not row:
                return jsonify({'success': False, 'error': 'Havola topilmadi'}), 404
            (tid, expires_at, used_by, created_by, canonical, max_staff,
             valid_until, role) = row
            if used_by:
                return jsonify({'success': False,
                                'error': 'Havola allaqachon ishlatilgan'}), 409
            if expires_at and expires_at < datetime.now():
                return jsonify({'success': False,
                                'error': 'Havola muddati tugagan'}), 410
            if valid_until and valid_until < date.today():
                return jsonify({'success': False,
                                'error': 'Universitet litsenziyasi muddati tugagan'}), 410
            if role == 'doctorant':
                # doktorant: xodim EMAS — monitoring uchun universitetga
                # bog'lanadi (university_doctorant_links) + profil institusiyasi
                # bo'sh bo'lsa to'ldiriladi
                cur.execute("""
                    INSERT INTO university_doctorant_links
                        (user_id, canonical_institution, invited_by)
                    VALUES (%s, %s, %s)
                    ON CONFLICT (user_id, canonical_institution) DO NOTHING
                """, (current_user.id, canonical, created_by))
                try:
                    email = (current_user.email or '').strip().lower()
                    if email:
                        cur.execute("""
                            UPDATE olim_profiles op SET institution = %s
                            FROM cabinet_users cu
                            WHERE cu.id = op.cabinet_user_id
                              AND LOWER(cu.email) = %s
                              AND COALESCE(TRIM(op.institution), '') = ''
                        """, (canonical, email))
                except Exception:
                    pass
                _audit(cur, canonical,
                       f"Doktorant taklifni qabul qildi: {current_user.username}")
            else:
                ok, msg = _add_staff(cur, canonical, current_user.email or
                                     current_user.username, 'staff', '',
                                     created_by, max_staff or DEFAULT_MAX_STAFF)
                if not ok:
                    return jsonify({'success': False, 'error': msg}), 400
                _audit(cur, canonical,
                       f"Xodim taklif havola orqali qo'shildi: {current_user.username}")
            cur.execute("UPDATE university_invite_tokens "
                        "SET used_by = %s, used_at = NOW() WHERE id = %s",
                        (current_user.id, tid))
        conn.commit()
        _nav_cache.pop(current_user.id, None)
        return jsonify({'success': True,
                        'redirect': '/reja' if role == 'doctorant' else '/univer'})
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        conn.close()


# ═══════════════ SAYT-ADMIN: LITSENZIYALARNI BOSHQARISH ══════════════════════
# Universitet bilan shartnoma imzolangach, sayt egasi shu yerdan onboard
# qiladi: litsenziya yaratadi, owner xodim qo'shadi yoki taklif havolasi beradi.

def _admin_guard():
    from app import _require_admin
    _require_admin()


@univer_bp.route('/admin/universities/licenses')
@login_required
def admin_licenses():
    _admin_guard()
    items, canon_names = [], []
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                _ensure_schema(cur)
                cur.execute("""
                    SELECT l.id, l.canonical_institution, l.plan, l.valid_until,
                           l.max_staff, l.created_at, l.notes,
                           (SELECT COUNT(*) FROM university_staff s
                             WHERE s.canonical_institution = l.canonical_institution
                               AND s.status = 'active') AS staff_count
                    FROM university_licenses l ORDER BY l.created_at DESC""")
                for r in cur.fetchall():
                    it = dict(zip(_LIC_COLS + ('staff_count',), r))
                    it['plan_label'] = PLAN_LABELS.get(it['plan'], it['plan'])
                    it['latin_name'] = _latin(it['canonical_institution'])
                    it['days_left'] = ((it['valid_until'] - date.today()).days
                                       if it['valid_until'] else None)
                    items.append(it)
                # institusiya autocomplete (datalist) — canonical nomlar
                cur.execute("""
                    SELECT COALESCE(canonical_name, cyrillic_name) AS canon,
                           COUNT(*) AS n
                    FROM institution_map WHERE is_active = TRUE
                    GROUP BY canon ORDER BY n DESC LIMIT 800""")
                canon_names = [r[0] for r in cur.fetchall()]
            conn.commit()
        finally:
            conn.close()
    except Exception:
        pass
    return render_template('admin_univer_licenses.html', items=items,
                           canon_names=canon_names, plan_labels=PLAN_LABELS)


@univer_bp.route('/admin/universities/licenses/create', methods=['POST'])
@login_required
def admin_license_create():
    _admin_guard()
    canonical = (request.form.get('canonical_institution') or '').strip()
    plan = request.form.get('plan') or 'pilot'
    if plan not in PLAN_LABELS:
        plan = 'pilot'
    valid_until = (request.form.get('valid_until') or '').strip() or None
    try:
        max_staff = max(1, min(100, int(request.form.get('max_staff') or DEFAULT_MAX_STAFF)))
    except ValueError:
        max_staff = DEFAULT_MAX_STAFF
    notes = (request.form.get('notes') or '').strip()[:2000] or None
    if not canonical:
        flash("Muassasa nomi majburiy.", "error")
        return redirect(url_for('univer.admin_licenses'))
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                _ensure_schema(cur)
                cur.execute("""
                    INSERT INTO university_licenses
                        (canonical_institution, plan, valid_until, max_staff, notes)
                    VALUES (%s, %s, %s, %s, %s)
                    ON CONFLICT (canonical_institution) DO UPDATE SET
                        plan = EXCLUDED.plan, valid_until = EXCLUDED.valid_until,
                        max_staff = EXCLUDED.max_staff, notes = EXCLUDED.notes
                    RETURNING id""", (canonical, plan, valid_until, max_staff, notes))
            conn.commit()
        finally:
            conn.close()
        flash("Litsenziya saqlandi.", "success")
    except Exception as e:
        flash(f"Xatolik: {e}", "error")
    return redirect(url_for('univer.admin_licenses'))


@univer_bp.route('/admin/universities/licenses/<int:lid>')
@login_required
def admin_license_detail(lid):
    _admin_guard()
    lic, staff = None, []
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                _ensure_schema(cur)
                cur.execute(f"SELECT {', '.join(_LIC_COLS)} FROM university_licenses "
                            "WHERE id = %s", (lid,))
                r = cur.fetchone()
                if r:
                    lic = dict(zip(_LIC_COLS, r))
                    lic['plan_label'] = PLAN_LABELS.get(lic['plan'], lic['plan'])
                    lic['latin_name'] = _latin(lic['canonical_institution'])
                    lic['days_left'] = ((lic['valid_until'] - date.today()).days
                                        if lic['valid_until'] else None)
                    cur.execute("""
                        SELECT s.id, u.username, u.email, s.role, s.title,
                               s.status, s.created_at
                        FROM university_staff s JOIN users u ON u.id = s.user_id
                        WHERE s.canonical_institution = %s
                        ORDER BY (s.role = 'owner') DESC, s.id""",
                                (lic['canonical_institution'],))
                    staff = [dict(zip(('id', 'username', 'email', 'role', 'title',
                                       'status', 'created_at'), sr))
                             for sr in cur.fetchall()]
            conn.commit()
        finally:
            conn.close()
    except Exception:
        pass
    if not lic:
        abort(404)
    for s in staff:
        s['role_label'] = ROLE_LABELS.get(s['role'], s['role'])
    return render_template('admin_univer_license_detail.html', lic=lic,
                           staff=staff, role_labels=ROLE_LABELS)


@univer_bp.route('/admin/universities/licenses/<int:lid>/staff/add',
                 methods=['POST'])
@login_required
def admin_license_staff_add(lid):
    _admin_guard()
    ident = (request.form.get('username_or_email') or '').strip()
    role = request.form.get('role') or 'staff'
    if role not in ('owner', 'staff', 'viewer'):
        role = 'staff'
    title = (request.form.get('title') or '').strip()[:200]
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("SELECT canonical_institution, max_staff "
                        "FROM university_licenses WHERE id = %s", (lid,))
            r = cur.fetchone()
            if not r:
                abort(404)
            canonical, max_staff = r
            if role == 'owner':
                # owner ni sayt admini qo'shadi — _add_staff faqat staff/viewer
                cur.execute("SELECT id, username FROM users "
                            "WHERE LOWER(username) = LOWER(%s) "
                            "OR LOWER(email) = LOWER(%s)", (ident, ident))
                t = cur.fetchone()
                if not t:
                    ok, msg = False, "Foydalanuvchi topilmadi."
                elif _staff_count(cur, canonical) >= (max_staff or DEFAULT_MAX_STAFF):
                    ok, msg = False, f"Xodimlar soni to'lgan (maksimal {max_staff} ta)."
                else:
                    cur.execute("""
                        INSERT INTO university_staff
                            (user_id, canonical_institution, role, title, invited_by)
                        VALUES (%s, %s, 'owner', %s, %s)
                        ON CONFLICT (user_id, canonical_institution) DO UPDATE SET
                            role = 'owner', status = 'active',
                            title = EXCLUDED.title""",
                                (t[0], canonical, title or None, current_user.id))
                    _nav_cache.pop(t[0], None)
                    ok, msg = True, f"{t[1]} owner sifatida qo'shildi."
            else:
                ok, msg = _add_staff(cur, canonical, ident, role, title,
                                     current_user.id,
                                     max_staff or DEFAULT_MAX_STAFF)
        conn.commit()
        flash(msg, "success" if ok else "error")
    finally:
        conn.close()
    return redirect(url_for('univer.admin_license_detail', lid=lid))


@univer_bp.route('/admin/universities/licenses/<int:lid>/staff/<int:sid>',
                 methods=['POST'])
@login_required
def admin_license_staff_action(lid, sid):
    _admin_guard()
    action = request.form.get('action') or ''
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("SELECT canonical_institution FROM university_licenses "
                        "WHERE id = %s", (lid,))
            r = cur.fetchone()
            if not r:
                abort(404)
            canonical = r[0]
            cur.execute("SELECT user_id FROM university_staff "
                        "WHERE id = %s AND canonical_institution = %s",
                        (sid, canonical))
            sr = cur.fetchone()
            if not sr:
                abort(404)
            if action == 'suspend':
                cur.execute("UPDATE university_staff SET status = "
                            "CASE WHEN status = 'active' THEN 'suspended' "
                            "ELSE 'active' END WHERE id = %s", (sid,))
            elif action == 'remove':
                cur.execute("DELETE FROM university_staff WHERE id = %s", (sid,))
            elif action in ('owner', 'staff', 'viewer'):
                cur.execute("UPDATE university_staff SET role = %s WHERE id = %s",
                            (action, sid))
            _nav_cache.pop(sr[0], None)
        conn.commit()
    finally:
        conn.close()
    return redirect(url_for('univer.admin_license_detail', lid=lid))


@univer_bp.route('/admin/universities/licenses/<int:lid>/invite-link',
                 methods=['POST'])
@csrf.exempt
@login_required
def admin_license_invite_link(lid):
    """7 kunlik bir martalik xodim taklif havolasi (pilotlar uchun).
    Havolani ochgan login foydalanuvchi 'staff' sifatida biriktiriladi."""
    _admin_guard()
    token = secrets.token_urlsafe(32)[:64]
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("SELECT id FROM university_licenses WHERE id = %s", (lid,))
            if not cur.fetchone():
                return jsonify({'success': False, 'error': 'Litsenziya topilmadi'}), 404
            cur.execute("""
                INSERT INTO university_invite_tokens (token, license_id, created_by)
                VALUES (%s, %s, %s)""", (token, lid, current_user.id))
        conn.commit()
    finally:
        conn.close()
    url = request.url_root.rstrip('/') + '/univer/invite/' + token
    return jsonify({'success': True, 'url': url})


# ═══════════════ KONFERENSIYALARIMIZ (Part 2 — boshqaruv) ════════════════════
# Universitet o'z konferensiyalarini boshqaradi: canonical_institution
# to'ldirilgan yozuvlar avtoritetli; organizer fuzzy-mos kelganlar ham
# ko'rsatiladi (faqat o'qish). Yangi yozuv 'pending' + is_active=FALSE bilan
# yaratiladi — admin tasdiqlamaguncha ommaviy sahifalarda KO'RINMAYDI.

_UNI_CONF_FIELDS = ('title', 'start_date', 'end_date', 'city', 'field',
                    'event_type', 'description', 'source_url')


def _organizer_match_sql(latin_name):
    """(sql, params) — organizer matnida universitet kalit so'zlari (AND)."""
    try:
        from app import _uni_keywords
        kws = _uni_keywords(latin_name)[:4]
    except Exception:
        kws = []
    if not kws:
        return "FALSE", []
    return (" AND ".join(["organizer ILIKE %s"] * len(kws)),
            [f'%{k}%' for k in kws])


def _uni_conferences(cur, canonical):
    """Universitet konferensiyalari: o'ziniki (canonical) + organizer mosi."""
    from blueprints.conferences import _ensure_schema as _ensure_conf
    _ensure_conf(cur)
    org_sql, org_params = _organizer_match_sql(_latin(canonical))
    cur.execute(f"""
        SELECT c.id, c.title, c.title_slug, c.start_date, c.end_date, c.city,
               c.field, c.event_type, c.source_url, c.is_active,
               COALESCE(c.moderation_status, 'approved'),
               (c.canonical_institution = %s) AS own,
               (SELECT COUNT(*) FROM conference_notifications_log l
                 WHERE l.conference_id = c.id) AS notified
        FROM conferences c
        WHERE c.canonical_institution = %s
           OR (c.canonical_institution IS NULL AND ({org_sql}))
        ORDER BY (c.canonical_institution = %s) DESC,
                 c.start_date DESC NULLS LAST, c.id DESC
        LIMIT 200
    """, [canonical, canonical] + org_params + [canonical])
    cols = ('id', 'title', 'slug', 'start_date', 'end_date', 'city', 'field',
            'event_type', 'source_url', 'is_active', 'moderation_status',
            'own', 'notified')
    return [dict(zip(cols, r)) for r in cur.fetchall()]


@univer_bp.route('/univer/konferensiyalar')
@login_required
def univer_conferences_page():
    ctx = _workspace_ctx('konferensiyalar')
    items, fields = [], []
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                _ensure_schema(cur)
                items = _uni_conferences(cur, ctx['canonical'])
                cur.execute("SELECT name FROM conference_fields ORDER BY name")
                fields = [r[0] for r in cur.fetchall()]
            conn.commit()
        finally:
            conn.close()
    except Exception:
        items = []
    from blueprints.conferences import EVENT_TYPES_LOCAL
    return render_template('univer/konferensiyalar.html', items=items,
                           fields=fields, event_types=EVENT_TYPES_LOCAL, **ctx)


def _conf_form_vals(data):
    from blueprints.conferences import _parse_date as _pd
    v = {k: (data.get(k) or '').strip() or None for k in _UNI_CONF_FIELDS}
    v['start_date'] = _pd(v['start_date'])
    v['end_date'] = _pd(v['end_date'])
    if v['title']:
        v['title'] = v['title'][:600]
    return v


@univer_bp.route('/univer/api/conferences/add', methods=['POST'])
@csrf.exempt
@login_required
def api_uni_conference_add():
    canonical = _manager_or_403()
    v = _conf_form_vals(request.get_json(silent=True) or {})
    if not v['title']:
        return jsonify({'success': False, 'error': 'Sarlavha kiritilishi shart'}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            from blueprints.conferences import (_ensure_schema as _ensure_conf,
                                                make_slug)
            _ensure_conf(cur)
            slug = make_slug(v['title'], cur)
            cur.execute("""
                INSERT INTO conferences
                    (title, title_slug, scope, organizer, field, city,
                     event_type, start_date, end_date, is_multiday, description,
                     source_url, country, source, canonical_institution,
                     submitted_by, moderation_status, is_active)
                VALUES (%s, %s, 'local', %s, %s, %s, %s, %s, %s, %s, %s, %s,
                        'O''zbekiston', 'universitet', %s, %s, 'pending', FALSE)
                RETURNING id
            """, (v['title'], slug, _latin(canonical), v['field'], v['city'],
                  v['event_type'], v['start_date'], v['end_date'],
                  bool(v['start_date'] and v['end_date']
                       and v['end_date'] != v['start_date']),
                  v['description'], v['source_url'], canonical,
                  current_user.id))
            new_id = cur.fetchone()[0]
            _audit(cur, canonical, f"Konferensiya yaratildi (moderatsiyaga): "
                                   f"{v['title'][:80]}")
        conn.commit()
        return jsonify({'success': True, 'id': new_id,
                        'message': "Yuborildi — sayt admini tasdiqlagach "
                                   "ommaviy e'lon qilinadi"})
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        conn.close()


def _own_conference_or_403(cur, canonical, cid):
    """Konferensiya AYNAN shu universitetniki ekanini tekshiradi (egalik
    guardi — boshqa universitet yozuviga urinish 403)."""
    cur.execute("SELECT id FROM conferences "
                "WHERE id = %s AND canonical_institution = %s", (cid, canonical))
    if not cur.fetchone():
        abort(make_response(jsonify({'success': False, 'error': 'forbidden'}), 403))


@univer_bp.route('/univer/api/conferences/<int:cid>/update', methods=['POST'])
@csrf.exempt
@login_required
def api_uni_conference_update(cid):
    canonical = _manager_or_403()
    v = _conf_form_vals(request.get_json(silent=True) or {})
    if not v['title']:
        return jsonify({'success': False, 'error': 'Sarlavha kiritilishi shart'}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            _own_conference_or_403(cur, canonical, cid)
            cur.execute("""
                UPDATE conferences SET title = %s, field = %s, city = %s,
                    event_type = %s, start_date = %s, end_date = %s,
                    is_multiday = %s, description = %s, source_url = %s,
                    updated_at = NOW()
                WHERE id = %s
            """, (v['title'], v['field'], v['city'], v['event_type'],
                  v['start_date'], v['end_date'],
                  bool(v['start_date'] and v['end_date']
                       and v['end_date'] != v['start_date']),
                  v['description'], v['source_url'], cid))
            _audit(cur, canonical, f"Konferensiya tahrirlandi (#{cid})")
        conn.commit()
        return jsonify({'success': True})
    finally:
        conn.close()


@univer_bp.route('/univer/api/conferences/<int:cid>/toggle', methods=['POST'])
@csrf.exempt
@login_required
def api_uni_conference_toggle(cid):
    """O'z (tasdiqlangan) konferensiyasini yashirish/qayta yoqish.
    'pending' yozuvni yoqib bo'lmaydi — faqat admin tasdiqlaydi."""
    canonical = _manager_or_403()
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            _own_conference_or_403(cur, canonical, cid)
            cur.execute("""
                UPDATE conferences SET is_active = NOT is_active, updated_at = NOW()
                WHERE id = %s AND COALESCE(moderation_status, 'approved') = 'approved'
                RETURNING is_active
            """, (cid,))
            r = cur.fetchone()
            if not r:
                return jsonify({'success': False,
                                'error': "Moderatsiyadagi yozuvni o'zgartirib "
                                         "bo'lmaydi"}), 400
            _audit(cur, canonical, f"Konferensiya holati o'zgartirildi (#{cid})")
        conn.commit()
        return jsonify({'success': True, 'is_active': bool(r[0])})
    finally:
        conn.close()


# ═══════════════════════ JURNALLARIMIZ (Part 2) ══════════════════════════════

@univer_bp.route('/univer/jurnallar')
@login_required
def univer_journals_page():
    ctx = _workspace_ctx('jurnallar')
    items = []
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                _ensure_schema(cur)
                org_sql, org_params = _organizer_match_sql(_latin(ctx['canonical']))
                pub_sql = org_sql.replace('organizer', 'publisher')
                cur.execute(f"""
                    SELECT id, name, issn, oak_approved, scopus_indexed,
                           wos_indexed, is_active,
                           COALESCE(moderation_status, 'approved'),
                           (canonical_institution = %s) AS own
                    FROM journals
                    WHERE canonical_institution = %s
                       OR (canonical_institution IS NULL AND ({pub_sql}))
                    ORDER BY (canonical_institution = %s) DESC, LOWER(name)
                    LIMIT 100
                """, [ctx['canonical'], ctx['canonical']] + org_params
                     + [ctx['canonical']])
                cols = ('id', 'name', 'issn', 'oak', 'scopus', 'wos',
                        'is_active', 'moderation_status', 'own')
                items = [dict(zip(cols, r)) for r in cur.fetchall()]
                # platformadagi iqtiboslar: Roadmap nashrlari shu jurnal nomini
                # venue sifatida ko'rsatgan soni (FAQAT SON — metrika)
                for it in items:
                    try:
                        cur.execute(
                            "SELECT COUNT(*) FROM roadmap_publications "
                            "WHERE venue ILIKE %s", (f"%{it['name'][:80]}%",))
                        it['cited'] = cur.fetchone()[0] or 0
                    except Exception:
                        it['cited'] = 0
            conn.commit()
        finally:
            conn.close()
    except Exception:
        items = []
    return render_template('univer/jurnallar.html', items=items, **ctx)


@univer_bp.route('/univer/api/journals/request', methods=['POST'])
@csrf.exempt
@login_required
def api_uni_journal_request():
    """"Jurnal qo'shish so'rovi" — pending journals qatori (is_active=FALSE,
    ommaviy ro'yxatlarda ko'rinmaydi) — sayt admini /admin/journals da
    tasdiqlaydi."""
    canonical = _manager_or_403()
    data = request.get_json(silent=True) or {}
    name = (data.get('name') or '').strip()[:500]
    if not name:
        return jsonify({'success': False, 'error': 'Jurnal nomi kiritilishi shart'}), 400
    issn = (data.get('issn') or '').strip()[:20] or None
    website = (data.get('website') or '').strip()[:500] or None
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            from blueprints.journal_check import normalize_title
            cur.execute("""
                INSERT INTO journals (name, name_normalized, issn, website,
                    publisher, country, canonical_institution, submitted_by,
                    moderation_status, is_active)
                VALUES (%s, %s, %s, %s, %s, 'O''zbekiston', %s, %s,
                        'pending', FALSE)
                ON CONFLICT (name) DO NOTHING RETURNING id
            """, (name, normalize_title(name), issn, website,
                  _latin(canonical), canonical, current_user.id))
            row = cur.fetchone()
            if not row:
                return jsonify({'success': False,
                                'error': 'Bu nomdagi jurnal allaqachon mavjud'}), 400
            _audit(cur, canonical, f"Jurnal so'rovi yuborildi: {name[:80]}")
        conn.commit()
        return jsonify({'success': True, 'id': row[0],
                        'message': "So'rov yuborildi — admin tasdiqlagach "
                                   "ro'yxatda ko'rinadi"})
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        conn.close()


@univer_bp.route('/admin/journals/approve/<int:jid>', methods=['POST'])
@login_required
def admin_journal_approve(jid):
    """Universitet yuborgan pending jurnalni tasdiqlash (sayt admini)."""
    _admin_guard()
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("UPDATE journals SET moderation_status = 'approved', "
                        "is_active = TRUE, updated_at = NOW() "
                        "WHERE id = %s AND moderation_status = 'pending' "
                        "RETURNING name", (jid,))
            r = cur.fetchone()
        conn.commit()
        flash(f"Jurnal tasdiqlandi: {r[0]}" if r else 'Tasdiqlanadigan yozuv topilmadi.',
              'success' if r else 'error')
    finally:
        conn.close()
    return redirect(url_for('admin.admin_journals'))


# ═══════════ DOKTORANT TAKLIFLARI (bulk invite) + DIGEST (Part 2) ════════════

MAX_BULK_INVITES = 50
_EMAIL_RE = re.compile(r'^[^@\s]+@[^@\s]+\.[^@\s]+$')


def _doctorant_invites(canonical):
    """Yuborilgan doktorant takliflari ro'yxati (holati bilan)."""
    out = []
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                _ensure_schema(cur)
                cur.execute("""
                    SELECT t.email, t.token, t.created_at, t.expires_at,
                           t.used_at, u.username
                    FROM university_invite_tokens t
                    JOIN university_licenses l ON l.id = t.license_id
                    LEFT JOIN users u ON u.id = t.used_by
                    WHERE l.canonical_institution = %s AND t.role = 'doctorant'
                    ORDER BY t.created_at DESC LIMIT 100
                """, (canonical,))
                now = datetime.now()
                for email, token, created, expires, used_at, uname in cur.fetchall():
                    if used_at:
                        status, cls = 'Qabul qilingan', 'good'
                    elif expires and expires < now:
                        status, cls = 'Muddati tugagan', 'muted'
                    else:
                        status, cls = 'Kutilmoqda', 'active'
                    out.append({'email': email or '—',
                                'url': f'/univer/invite/{token}',
                                'created_at': str(created)[:10] if created else '',
                                'status': status, 'status_class': cls,
                                'accepted_by': uname or ''})
            conn.commit()
        finally:
            conn.close()
    except Exception:
        pass
    return out


@univer_bp.route('/univer/api/doktorant/invites', methods=['POST'])
@csrf.exempt
@login_required
def api_doctorant_invites():
    """Bulk taklif: har qatorda bitta email (maks. 50). Har biriga 7 kunlik
    bir martalik havola yaratiladi — xodim ularni o'zi yuboradi (SMTP yo'q,
    havolalar ro'yxati qaytadi)."""
    canonical = _manager_or_403()
    raw = (request.get_json(silent=True) or {}).get('emails') or ''
    emails, bad = [], []
    for line in raw.splitlines():
        e = line.strip().lower()
        if not e:
            continue
        (emails if _EMAIL_RE.match(e) else bad).append(e)
    emails = list(dict.fromkeys(emails))          # takrorlarni olib tashlash
    if not emails:
        return jsonify({'success': False,
                        'error': "Kamida bitta to'g'ri email kiriting"}), 400
    if len(emails) > MAX_BULK_INVITES:
        return jsonify({'success': False,
                        'error': f"Bir so'rovda ko'pi bilan {MAX_BULK_INVITES} ta "
                                 f"email (siz {len(emails)} ta yubordingiz)"}), 400
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("SELECT id FROM university_licenses "
                        "WHERE canonical_institution = %s", (canonical,))
            lic = cur.fetchone()
            if not lic:
                return jsonify({'success': False, 'error': 'Litsenziya topilmadi'}), 404
            links = []
            for e in emails:
                token = secrets.token_urlsafe(32)[:64]
                cur.execute("""
                    INSERT INTO university_invite_tokens
                        (token, license_id, created_by, role, email)
                    VALUES (%s, %s, %s, 'doctorant', %s)
                """, (token, lic[0], current_user.id, e))
                links.append({'email': e,
                              'url': request.url_root.rstrip('/')
                                     + '/univer/invite/' + token})
            _audit(cur, canonical,
                   f"Doktorant takliflari yaratildi ({len(links)} ta)")
        conn.commit()
        return jsonify({'success': True, 'links': links, 'invalid': bad})
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        conn.close()


@univer_bp.route('/univer/api/digest-pref', methods=['POST'])
@csrf.exempt
@login_required
def api_digest_pref():
    """Xodimning o'z kunlik digest toggle'i (university_staff.digest_enabled)."""
    canonical, _role = get_university_access_or_403()
    enabled = bool((request.get_json(silent=True) or {}).get('enabled'))
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("UPDATE university_staff SET digest_enabled = %s "
                        "WHERE user_id = %s AND canonical_institution = %s",
                        (enabled, current_user.id, canonical))
        conn.commit()
        return jsonify({'success': True, 'enabled': enabled})
    finally:
        conn.close()


DIGEST_DEFENSE_MARKS = (60, 30)


def _digest_events(cur, canonical):
    """Oxirgi 24 soatdagi doktorant voqealari — FAQAT METRIKALAR:
    nashr qo'shildi, bob yakunlandi (word_target'ga yetdi), himoya 60/30 kun."""
    users = _doctoral_user_ids(cur, canonical)
    if not users:
        return []
    ids = list(users)
    events = []
    for uid, title in _fetchall(cur, """
            SELECT rp.user_id, pub.title FROM roadmap_publications pub
            JOIN roadmap_plans rp ON rp.id = pub.plan_id
            WHERE rp.user_id = ANY(%s)
              AND pub.created_at >= NOW() - INTERVAL '1 day'
            LIMIT 20""", (ids,)):
        nm = users.get(uid, {}).get('name', 'Doktorant')
        events.append(f"📄 {nm} yangi nashr qo'shdi")
    for uid, btitle in _fetchall(cur, """
            SELECT p.owner_id, b.title FROM dissertation_blocks b
            JOIN diss_projects p ON p.id = b.dissertation_id
            WHERE p.owner_id = ANY(%s) AND p.status <> 'archived'
              AND COALESCE(b.word_target, 0) > 0
              AND b.word_count >= b.word_target
              AND b.updated_at >= NOW() - INTERVAL '1 day'
            LIMIT 20""", (ids,)):
        nm = users.get(uid, {}).get('name', 'Doktorant')
        events.append(f"✅ {nm}: \"{(btitle or '')[:40]}\" bo'limi maqsadga yetdi")
    for uid, days in _fetchall(cur, """
            SELECT user_id, (target_defense_date - CURRENT_DATE)
            FROM roadmap_plans
            WHERE is_active AND user_id = ANY(%s)
              AND (target_defense_date - CURRENT_DATE) = ANY(%s)""",
            (ids, list(DIGEST_DEFENSE_MARKS))):
        nm = users.get(uid, {}).get('name', 'Doktorant')
        events.append(f"⏰ {nm} himoyasiga {days} kun qoldi")
    return events[:15]


@univer_bp.route('/api/v1/univer/dispatch-digest', methods=['POST'])
@csrf.exempt
def dispatch_univer_digest():
    """Kunlik cron (GitHub Actions, REMINDERS_API_KEY) — har faol litsenziyali
    universitet uchun voqealarni yig'ib, xodimlarga KUNIGA KO'PI BILAN BITTA
    user_alert yuboradi (university_digest_log dedup)."""
    import os as _os
    key = request.headers.get('X-Api-Key') or request.args.get('key') or ''
    expected = _os.environ.get('REMINDERS_API_KEY', '')
    if not expected or key != expected:
        return jsonify({'ok': False, 'error': 'forbidden'}), 403
    sent = skipped = 0
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            try:
                from blueprints.notifications import _ensure_schema as _ensure_notif
                _ensure_notif(cur)
            except Exception:
                pass
            cur.execute("""
                SELECT canonical_institution FROM university_licenses
                WHERE valid_until IS NULL OR valid_until >= CURRENT_DATE
            """)
            unis = [r[0] for r in cur.fetchall()]
            for canonical in unis:
                # kunlik dedup — INSERT muvaffaqiyatsiz bo'lsa bugun yuborilgan
                cur.execute("""
                    INSERT INTO university_digest_log
                        (canonical_institution, digest_date)
                    VALUES (%s, CURRENT_DATE)
                    ON CONFLICT (canonical_institution, digest_date) DO NOTHING
                    RETURNING id""", (canonical,))
                log = cur.fetchone()
                if not log:
                    skipped += 1
                    continue
                events = _digest_events(cur, canonical)
                if not events:
                    # voqea yo'q — log qoladi (qayta urinilmaydi), alert yo'q
                    continue
                cur.execute("UPDATE university_digest_log SET events_count = %s "
                            "WHERE id = %s", (len(events), log[0]))
                msg = ("Doktorantlaringiz bo'yicha bugungi xulosa:\n"
                       + "\n".join(events)
                       + "\n🔗 /univer/doktorantlar")
                cur.execute("""
                    SELECT user_id FROM university_staff
                    WHERE canonical_institution = %s AND status = 'active'
                      AND COALESCE(digest_enabled, TRUE)""", (canonical,))
                for (staff_uid,) in cur.fetchall():
                    cur.execute("""
                        INSERT INTO user_alerts (user_id, title, message, level)
                        VALUES (%s, %s, %s, 'info')
                    """, (staff_uid, f"🏛 {_latin(canonical)} — kunlik digest", msg))
                    sent += 1
        conn.commit()
        return jsonify({'ok': True, 'universities': len(unis),
                        'alerts_sent': sent, 'already_sent_today': skipped})
    except Exception as e:
        conn.rollback()
        return jsonify({'ok': False, 'error': str(e)}), 500
    finally:
        conn.close()


# ═══════════════════ HISOBOTLAR (Part 2 — OAK hisoboti) ═════════════════════

_report_last = {}       # user_id → ts (1 hisobot/daqiqa)
REPORT_SECTIONS = {
    'himoyalar': "Himoyalar ro'yxati",
    'ixtisosliklar': 'Ixtisosliklar kesimi',
    'rahbarlar': 'Rahbarlar faolligi',
    'doktorantlar': 'Doktorantlar progressi',
    'nashrlar': 'Nashrlar statistikasi',
}
_REPORT_ROW_CAP = 2000


def _report_period(a):
    """(y0, y1, m0, m1, label) — yil / chorak / oraliq."""
    this_year = date.today().year
    try:
        y0 = int(a.get('y0') or a.get('year') or this_year)
    except (TypeError, ValueError):
        y0 = this_year
    try:
        y1 = int(a.get('y1') or y0)
    except (TypeError, ValueError):
        y1 = y0
    if y0 > y1:
        y0, y1 = y1, y0
    y0, y1 = max(1950, y0), min(2100, y1)
    quarter = a.get('quarter') or ''
    m0, m1 = 1, 12
    label = f'{y0}' if y0 == y1 else f'{y0}–{y1}'
    if y0 == y1 and quarter in ('1', '2', '3', '4'):
        q = int(quarter)
        m0, m1 = (q - 1) * 3 + 1, q * 3
        label = f'{y0}, {q}-chorak'
    return y0, y1, m0, m1, label


def _report_data(canonical, y0, y1, m0, m1, sections):
    """Hisobot bo'limlari uchun ma'lumot — mavjud skoplangan so'rovlar ustida
    (yangi og'ir agregatsiya yo'q, faqat prezentatsiya)."""
    out = {}
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            variants = _variants(cur, canonical)
            W = "TRIM(d.muassasa) = ANY(%s)"
            yr = (f"d.sana ~ '(19|20)\\d{{2}}' "
                  f"AND ({_YEAR_EXPR})::int BETWEEN %s AND %s")
            mon = ""
            mparams = []
            if (m0, m1) != (1, 12):
                mon = (" AND TRIM(d.sana) ~ '^[0-9]{2}[.][0-9]{2}[.][0-9]{4}$' "
                       "AND substring(TRIM(d.sana) from 4 for 2)::int "
                       "BETWEEN %s AND %s")
                mparams = [m0, m1]
            base_params = [variants, y0, y1] + mparams

            if 'himoyalar' in sections:
                cur.execute(f"""
                    SELECT d.sana, d.olim, d.mavzu, d.daraja, d.ixtisoslik,
                           d.ilmiy_rahbar
                    FROM dissertations d WHERE {W} AND {yr}{mon}
                    ORDER BY ({_YEAR_EXPR})::int, d.id LIMIT %s
                """, base_params + [_REPORT_ROW_CAP])
                out['himoyalar'] = [
                    {'sana': (r[0] or '').strip(), 'olim': (r[1] or '').strip(),
                     'mavzu': (r[2] or '').strip(), 'daraja': (r[3] or '').strip(),
                     'ixtisoslik': (r[4] or '').strip(),
                     'rahbar': (r[5] or '').strip()} for r in cur.fetchall()]

            if 'ixtisosliklar' in sections:
                cur.execute(f"""
                    SELECT TRIM(d.ixtisoslik), MAX(d.ixtisoslik_nomi), COUNT(*)
                    FROM dissertations d WHERE {W} AND {yr}{mon}
                      AND d.ixtisoslik IS NOT NULL AND TRIM(d.ixtisoslik) <> ''
                    GROUP BY 1 ORDER BY 3 DESC LIMIT 50
                """, base_params)
                out['ixtisosliklar'] = [
                    {'code': r[0], 'name': (r[1] or '').strip(), 'count': r[2]}
                    for r in cur.fetchall()]

            if 'rahbarlar' in sections:
                cur.execute(f"""
                    SELECT TRIM(d.ilmiy_rahbar), COUNT(*),
                           COUNT(*) FILTER (WHERE {_PHD_LIKE}),
                           COUNT(*) FILTER (WHERE {_DSC_LIKE})
                    FROM dissertations d WHERE {W} AND {yr}{mon}
                      AND d.ilmiy_rahbar IS NOT NULL AND TRIM(d.ilmiy_rahbar) <> ''
                    GROUP BY 1 ORDER BY 2 DESC LIMIT 50
                """, base_params)
                out['rahbarlar'] = [
                    {'name': r[0], 'count': r[1], 'phd': r[2], 'dsc': r[3]}
                    for r in cur.fetchall()]

            if 'nashrlar' in sections:
                users = _doctoral_user_ids(cur, canonical)
                rows = []
                if users:
                    from blueprints.roadmap import (PUB_TYPE_LABELS,
                                                    PUB_STATUS_LABELS)
                    cur.execute("""
                        SELECT pub.pub_type, pub.status, COUNT(*)
                        FROM roadmap_publications pub
                        JOIN roadmap_plans rp ON rp.id = pub.plan_id
                        WHERE rp.user_id = ANY(%s)
                        GROUP BY 1, 2 ORDER BY 1, 2
                    """, (list(users),))
                    for ptype, status, cnt in cur.fetchall():
                        rows.append({
                            'type': PUB_TYPE_LABELS.get(ptype, ptype),
                            'status': PUB_STATUS_LABELS.get(status, status),
                            'count': cnt})
                out['nashrlar'] = rows
        conn.commit()
    finally:
        conn.close()
    if 'doktorantlar' in sections:
        out['doktorantlar'] = _doctoral_rows(canonical)
    return out


@univer_bp.route('/univer/hisobotlar')
@login_required
def univer_reports_page():
    ctx = _workspace_ctx('hisobotlar')
    return render_template('univer/hisobotlar.html',
                           sections=REPORT_SECTIONS,
                           this_year=date.today().year, **ctx)


@univer_bp.route('/univer/hisobot')
@login_required
def univer_report_generate():
    canonical, _role = get_university_access_or_403()
    now = time.time()
    if now - _report_last.get(current_user.id, 0) < 60:
        return Response("Hisobot tayyorlanmoqda — bir daqiqadan so'ng qayta "
                        "urinib ko'ring", status=429,
                        mimetype='text/plain; charset=utf-8')
    _report_last[current_user.id] = now
    a = request.args
    y0, y1, m0, m1, period_label = _report_period(a)
    sections = [s for s in (a.get('sections') or '').split(',')
                if s in REPORT_SECTIONS] or list(REPORT_SECTIONS)
    fmt = a.get('format') if a.get('format') in ('screen', 'xlsx', 'docx') \
        else 'screen'
    try:
        data = _report_data(canonical, y0, y1, m0, m1, sections)
    except Exception:
        _report_last.pop(current_user.id, None)   # muvaffaqiyatsiz urinish limitga kirmaydi
        return Response("Hisobot ma'lumotlarini olishda xatolik — birozdan "
                        "so'ng qayta urinib ko'ring", status=503,
                        mimetype='text/plain; charset=utf-8')
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                _ensure_schema(cur)
                _audit(cur, canonical,
                       f"Hisobot yaratildi ({period_label}, {fmt})")
            conn.commit()
        finally:
            conn.close()
    except Exception:
        pass
    latin = _latin(canonical)
    if fmt == 'xlsx':
        return _report_xlsx(latin, period_label, sections, data)
    if fmt == 'docx':
        return _report_docx(latin, period_label, sections, data)
    return render_template('univer/report.html', latin_name=latin,
                           period_label=period_label, sections=sections,
                           section_labels=REPORT_SECTIONS, data=data,
                           generated_at=datetime.now().strftime('%d.%m.%Y %H:%M'))


_REPORT_HEADERS = {
    'himoyalar': ['Sana', 'Olim', 'Mavzu', 'Daraja', 'Ixtisoslik', 'Ilmiy rahbar'],
    'ixtisosliklar': ['Shifr', 'Nomi', 'Himoyalar soni'],
    'rahbarlar': ['Ilmiy rahbar', 'Jami', 'PhD', 'DSc'],
    'doktorantlar': ['Doktorant', 'Daraja', 'Ixtisoslik', 'Rahbar',
                     'OAK nashrlari', 'Xalqaro nashrlar', "So'zlar", 'Progress %',
                     'Himoya sanasi', 'Holat'],
    'nashrlar': ['Nashr turi', 'Holati', 'Soni'],
}


def _report_rows(section, data):
    """Bo'lim ma'lumotini jadval qatorlariga aylantiradi (xlsx/docx uchun)."""
    rows = data.get(section) or []
    if section == 'himoyalar':
        return [[r['sana'], r['olim'], r['mavzu'], r['daraja'],
                 r['ixtisoslik'], r['rahbar']] for r in rows]
    if section == 'ixtisosliklar':
        return [[r['code'], r['name'], r['count']] for r in rows]
    if section == 'rahbarlar':
        return [[r['name'], r['count'], r['phd'], r['dsc']] for r in rows]
    if section == 'doktorantlar':
        return [[r['name'], r['degree'], r['specialty'], r['advisor'],
                 f"{r['oak_done']}/{r['oak_req']}",
                 f"{r['int_done']}/{r['int_req']}" if r['int_req'] else '—',
                 f"{r['words']}/{r['word_target']}", r['overall'],
                 r['defense_date'] or '—', r['holat']] for r in rows]
    if section == 'nashrlar':
        return [[r['type'], r['status'], r['count']] for r in rows]
    return []


def _report_xlsx(latin, period_label, sections, data):
    import io
    import openpyxl
    from openpyxl.styles import Font
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    for s in sections:
        ws = wb.create_sheet(REPORT_SECTIONS[s][:31])
        ws.append([f"{latin} — {REPORT_SECTIONS[s]} ({period_label})"])
        ws['A1'].font = Font(bold=True, size=13)
        ws.append([])
        ws.append(_REPORT_HEADERS[s])
        for c in ws[3]:
            c.font = Font(bold=True)
        for row in _report_rows(s, data):
            ws.append(row)
        for i, w in enumerate((14, 30, 60, 12, 14, 30)[:len(_REPORT_HEADERS[s])]):
            ws.column_dimensions[chr(65 + i)].width = w
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return Response(buf.read(),
                    mimetype='application/vnd.openxmlformats-officedocument'
                             '.spreadsheetml.sheet',
                    headers={'Content-Disposition':
                             'attachment; filename="hisobot.xlsx"'})


def _report_docx(latin, period_label, sections, data):
    import io
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading(f'«{latin}»', 0)
    doc.add_heading(f'Ilmiy faoliyat hisoboti, {period_label}', level=1)
    p = doc.add_paragraph(
        f"Hisobot olimlar.uz platformasi ma'lumotlari asosida "
        f"{datetime.now().strftime('%d.%m.%Y')} sanasida tayyorlandi.")
    p.runs[0].font.size = Pt(10)
    for s in sections:
        doc.add_heading(REPORT_SECTIONS[s], level=2)
        rows = _report_rows(s, data)
        if not rows:
            doc.add_paragraph("Ma'lumot topilmadi.")
            continue
        headers = _REPORT_HEADERS[s]
        # .docx jadvali katta ro'yxatlarda og'irlashadi — 300 qator kifoya,
        # to'liq ro'yxat uchun .xlsx bor
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row in rows[:300]:
            cells = table.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v if v is not None else '')
        if len(rows) > 300:
            doc.add_paragraph(f"... va yana {len(rows) - 300} ta qator "
                              f"(to'liq ro'yxat .xlsx formatida).")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(buf.read(),
                    mimetype='application/vnd.openxmlformats-officedocument'
                             '.wordprocessingml.document',
                    headers={'Content-Disposition':
                             'attachment; filename="hisobot.docx"'})


# ═══════════ OMMAVIY PROFIL TAHRIRI (litsenziya perki, Part 2) ═══════════════

@univer_bp.route('/univer/api/public-profile', methods=['POST'])
@csrf.exempt
@login_required
def api_public_profile():
    """Ommaviy /university/<name> sahifasidagi tavsif/veb-sayt/logo/email —
    owner+staff tahrirlaydi. body.canonical AYNAN xodimning universiteti
    bo'lishi shart (boshqa universitet sahifasidan urinish → 403)."""
    canonical = _manager_or_403()
    data = request.get_json(silent=True) or {}
    if (data.get('canonical') or '').strip() != canonical:
        return jsonify({'success': False, 'error': 'forbidden'}), 403
    desc = (data.get('description') or '').strip()[:3000] or None
    website = (data.get('website') or '').strip()[:500] or None
    logo = (data.get('logo_url') or '').strip()[:500] or None
    email = (data.get('contact_email') or '').strip()[:255] or None
    if website and not website.startswith(('http://', 'https://')):
        website = 'https://' + website
    conn = _conn()
    try:
        with conn.cursor() as cur:
            _ensure_schema(cur)
            cur.execute("""
                INSERT INTO university_public_profiles
                    (canonical_institution, description, website, logo_url,
                     contact_email, updated_by, updated_at)
                VALUES (%s, %s, %s, %s, %s, %s, NOW())
                ON CONFLICT (canonical_institution) DO UPDATE SET
                    description = EXCLUDED.description,
                    website = EXCLUDED.website,
                    logo_url = EXCLUDED.logo_url,
                    contact_email = EXCLUDED.contact_email,
                    updated_by = EXCLUDED.updated_by, updated_at = NOW()
            """, (canonical, desc, website, logo, email, current_user.id))
            _audit(cur, canonical, 'Ommaviy profil tahrirlandi')
        conn.commit()
        return jsonify({'success': True})
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        conn.close()


def get_public_profile_context(cur, canonical, user_id=None):
    """/university/<name> sahifasi uchun (content.py chaqiradi):
    {official, profile{...}, can_edit} — jadval yo'q/DB xatosida bo'sh dict."""
    out = {'official': False, 'profile': None, 'can_edit': False}
    if not canonical:
        return out
    try:
        cur.execute("SELECT 1 FROM university_licenses "
                    "WHERE canonical_institution = %s "
                    "AND (valid_until IS NULL OR valid_until >= CURRENT_DATE)",
                    (canonical,))
        out['official'] = cur.fetchone() is not None
        cur.execute("SELECT description, website, logo_url, contact_email "
                    "FROM university_public_profiles "
                    "WHERE canonical_institution = %s", (canonical,))
        r = cur.fetchone()
        if r:
            out['profile'] = {'description': r[0] or '', 'website': r[1] or '',
                              'logo_url': r[2] or '', 'contact_email': r[3] or ''}
        if user_id and out['official']:
            cur.execute("SELECT role FROM university_staff "
                        "WHERE user_id = %s AND canonical_institution = %s "
                        "AND status = 'active'", (user_id, canonical))
            sr = cur.fetchone()
            out['can_edit'] = bool(sr and sr[0] in ('owner', 'staff'))
    except Exception:
        pass
    return out


# ── navbar bayrog'i (har so'rovda yengil, 2 daqiqa modul keshi) ──────────────

@univer_bp.app_context_processor
def inject_univer_nav():
    """base.html: xodimlarga '🏛 Universitetim' navbar bandi. Jadval hali
    yaratilmagan/DB xatosida jimgina '' (sayt yiqilmaydi)."""
    uid = getattr(current_user, 'id', None) if \
        getattr(current_user, 'is_authenticated', False) else None
    if not uid:
        return {'univer_workspace': ''}
    now = time.time()
    hit = _nav_cache.get(uid)
    if hit and (now - hit[0]) < _NAV_TTL:
        return {'univer_workspace': hit[1]}
    canonical = ''
    try:
        conn = _conn()
        try:
            with conn.cursor() as cur:
                cur.execute("SELECT canonical_institution FROM university_staff "
                            "WHERE user_id = %s AND status = 'active' LIMIT 1",
                            (uid,))
                r = cur.fetchone()
                canonical = r[0] if r else ''
            conn.commit()
        finally:
            conn.close()
    except Exception:
        canonical = ''
    _nav_cache[uid] = (now, canonical)
    return {'univer_workspace': canonical}
