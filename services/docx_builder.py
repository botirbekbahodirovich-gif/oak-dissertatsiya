"""Kafedra hujjatlari — AI yaratgan JSON kontentni rasmiy .docx faylga aylantiradi.

Shakllantirish (loyiha talabiga ko'ra, `blueprints/dissertation.py:_build_docx` bilan bir xil
konvensiya): Times New Roman, asosiy matn 12pt / sarlavhalar 14pt, interval 1.15,
chekkalar yuqori/pastdan 2sm, chapdan 3sm, o'ngdan 1.5sm.

Muassasa va shaxslar maydonlari (universitet, tuzuvchi F.I.Sh. va h.k.) foydalanuvchi
tomonidan bo'sh qoldirilishi mumkin — bunday hollarda maydon nomi sariq highlight
(`add_placeholder_run`) bilan qo'yiladi, foydalanuvchi Word'da ochib darrov ko'radi.
"""
import io
import uuid

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Cm, Pt

from services.document_prompts import DOCUMENT_TYPES, RAHBAR_LABELS, SECTIONS

GENERATED_DOCS_DIR = 'static/generated_docs'


def _setup_document():
    doc = Document()
    for section in doc.sections:
        section.left_margin, section.right_margin = Cm(3), Cm(1.5)
        section.top_margin = section.bottom_margin = Cm(2)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    return doc


def add_placeholder_run(paragraph, text):
    """Bo'sh qolgan maydon o'rniga sariq highlight bilan maydon nomini yozadi —
    foydalanuvchi Word'da ochib darrov ko'rib, o'zi to'ldiradi."""
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return run


def _heading(doc, text, size=14):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(6)
    return h


def _para(doc, text, bullet=False):
    text = (text or '').strip()
    if not text:
        return
    p = doc.add_paragraph(text, style='List Bullet' if bullet else None)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)


def _table(doc, headers, rows):
    """rows — list of tuples/lists mos uzunlikda `headers` bilan."""
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = '' if val is None else str(val)
    doc.add_paragraph()


def _info_grid(doc, pairs):
    """pairs — [(label, value_or_None), ...]. Bo'sh value → sariq placeholder (label matni)."""
    if not pairs:
        return
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for label, value in pairs:
        row = table.add_row().cells
        lp = row[0].paragraphs[0]
        lrun = lp.add_run(label)
        lrun.bold = True
        lrun.font.name = 'Times New Roman'
        lrun.font.size = Pt(12)
        vp = row[1].paragraphs[0]
        value = (value or '').strip() if isinstance(value, str) else value
        if value:
            vrun = vp.add_run(str(value))
            vrun.font.name = 'Times New Roman'
            vrun.font.size = Pt(12)
        else:
            add_placeholder_run(vp, label)
    doc.add_paragraph()


def _muassasa_field_map(doc_type):
    """Shu hujjat turi qaysi muassasa bo'limidan foydalanadi (full/qisqa) — label lookup uchun."""
    cfg = DOCUMENT_TYPES[doc_type]
    for sect_key in ('muassasa_shaxslar_full', 'muassasa_shaxslar_qisqa'):
        if sect_key in cfg.get('sections', []):
            return sect_key, SECTIONS[sect_key]['fields']
    return None, {}


def _add_header_block(doc, metadata, doc_type):
    """Rasmiy sarlavha bloki: Respublika / Universitet / TASDIQLAYMAN — bo'sh
    maydonlar sariq placeholder bilan."""
    sect_key, fields = _muassasa_field_map(doc_type)

    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t1.add_run("O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI").bold = True

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    universitet = (metadata.get('universitet') or '').strip()
    if universitet:
        t2.add_run(universitet)
    else:
        add_placeholder_run(t2, 'Universitet nomi')

    if 'kafedra_nomi' in fields:
        t3 = doc.add_paragraph()
        t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kafedra = (metadata.get('kafedra_nomi') or '').strip()
        fakultet = (metadata.get('fakultet') or '').strip()
        if fakultet:
            t3.add_run(f"{fakultet} fakulteti, ")
        if kafedra:
            t3.add_run(f"{kafedra} kafedrasi")
        else:
            add_placeholder_run(t3, 'Kafedra nomi')

    doc.add_paragraph()
    approve = doc.add_paragraph()
    approve.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    approve.add_run('TASDIQLAYMAN\n')
    rahbar_label = RAHBAR_LABELS.get(doc_type, 'Rahbar F.I.Sh.')
    approve.add_run(f'{rahbar_label} __________________ ')
    rahbar_fio = (metadata.get('rahbar_fio') or '').strip()
    if rahbar_fio:
        approve.add_run(rahbar_fio)
    else:
        add_placeholder_run(approve, rahbar_label)
    approve.add_run('\n"____" _______________ ' +
                     ((metadata.get('oquv_yili') or '').split('-')[0] or '20____') + ' yil')
    doc.add_paragraph()


def _add_title(doc, metadata, label):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{metadata.get('fan_nomi', '')} fanidan {label.upper()}")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph()


def _add_info_table(doc, metadata, doc_type):
    cfg = DOCUMENT_TYPES[doc_type]
    rows = []
    for key, label in (('soha', 'Bilim sohasi (OAK)'), ('mutaxassislik', 'Mutaxassislik'),
                        ('talim_bosqichi', "Ta'lim bosqichi")):
        val = metadata.get(key)
        if isinstance(val, str):
            val = val.strip()
        if val:
            rows.append((label, val))
    if 'kredit_soat' in cfg.get('sections', []):
        hours = sum(int(metadata.get(f) or 0) for f in SECTIONS['kredit_soat']['hours_fields'])
        if metadata.get('kredit') or hours:
            rows.append(('Kredit / soatlar',
                        f"{metadata.get('kredit', 3)} kredit, jami {hours} soat "
                        f"(ma'ruza {metadata.get('maruza_soat', 0)}, amaliy {metadata.get('amaliy_soat', 0)}, "
                        f"mustaqil {metadata.get('mustaqil_soat', 0)})"))
    if rows:
        _table(doc, ['Ma\'lumot', 'Qiymat'], rows)

    _, fields = _muassasa_field_map(doc_type)
    if fields:
        extra_pairs = []
        for key in ('fan_kodi', 'bilim_sohasi_diplom', 'talim_sohasi_diplom', 'talim_yonalishi_diplom',
                    'tuzuvchi_fio', 'ilmiy_daraja_unvon', 'lavozim', 'email', 'telefon',
                    'taqrizchi1_fio', 'taqrizchi1_daraja', 'taqrizchi2_fio', 'taqrizchi2_daraja'):
            if key in fields:
                extra_pairs.append((fields[key]['label'], metadata.get(key)))
        _info_grid(doc, extra_pairs)


# ── har bir hujjat turi uchun kontent qismini yozib chiquvchi funksiyalar ────

def _render_sillabus_like(doc, content):
    if content.get('fanning_maqsadi'):
        _heading(doc, 'Fanning maqsadi')
        _para(doc, content['fanning_maqsadi'])
    if content.get('fanning_vazifalari'):
        _heading(doc, 'Fanning vazifalari')
        for v in content['fanning_vazifalari']:
            _para(doc, v, bullet=True)
    if content.get('boglanish_fanlar'):
        _heading(doc, 'Boshqa fanlar bilan bog\'liqligi')
        for v in content['boglanish_fanlar']:
            _para(doc, v, bullet=True)
    if content.get('kutilayotgan_natijalar'):
        _heading(doc, 'Kutilayotgan natijalar')
        _table(doc, ['Kod', 'Natija'],
               [(n.get('kod', ''), n.get('matn', '')) for n in content['kutilayotgan_natijalar']])
    if content.get('mavzular_rejasi'):
        _heading(doc, 'Mavzular rejasi')
        _table(doc, ['Hafta', 'Mavzu', 'Mazmuni', 'Ma\'ruza', 'Amaliy'],
               [(m.get('hafta', ''), m.get('mavzu', ''), m.get('mazmuni', ''),
                 m.get('maruza_soati', ''), m.get('amaliy_soati', ''))
                for m in content['mavzular_rejasi']])
    bm = content.get('baholash_mezoni')
    if bm:
        _heading(doc, 'Baholash mezoni')
        _table(doc, ['Nazorat turi', 'Ulush (%)'],
               [(k.replace('_', ' ').capitalize(), v) for k, v in bm.items()])
    ad = content.get('adabiyotlar')
    if ad:
        _heading(doc, 'Foydalanilgan adabiyotlar')
        if ad.get('asosiy'):
            _para(doc, 'Asosiy adabiyotlar:')
            for a in ad['asosiy']:
                _para(doc, a, bullet=True)
        if ad.get('qoshimcha'):
            _para(doc, "Qo'shimcha adabiyotlar:")
            for a in ad['qoshimcha']:
                _para(doc, a, bullet=True)
        if ad.get('internet_resurslar'):
            _para(doc, 'Internet resurslar:')
            for a in ad['internet_resurslar']:
                _para(doc, a, bullet=True)


def _render_uslubiy_qollanma(doc, content):
    if content.get('kirish'):
        _heading(doc, 'Kirish')
        _para(doc, content['kirish'])
    if content.get('umumiy_korsatmalar'):
        _heading(doc, "Umumiy ko'rsatmalar")
        for v in content['umumiy_korsatmalar']:
            _para(doc, v, bullet=True)
    for m in content.get('mashgulotlar', []):
        _heading(doc, f"{m.get('tartib_raqami', '')}-mashg'ulot: {m.get('mavzu', '')}", size=13)
        if m.get('maqsad'):
            _para(doc, f"Maqsad: {m['maqsad']}")
        for t in m.get('topshiriqlar', []):
            _para(doc, t, bullet=True)
        if m.get('nazorat_savollari'):
            _para(doc, 'Nazorat savollari:')
            for s in m['nazorat_savollari']:
                _para(doc, s, bullet=True)
    ad = content.get('adabiyotlar')
    if ad:
        _heading(doc, 'Foydalanilgan adabiyotlar')
        for a in ad.get('asosiy', []):
            _para(doc, a, bullet=True)
        for a in ad.get('qoshimcha', []):
            _para(doc, a, bullet=True)


def _render_mustaqil_talim(doc, content):
    if content.get('kirish'):
        _heading(doc, 'Kirish')
        _para(doc, content['kirish'])
    for t in content.get('topshiriqlar', []):
        _heading(doc, f"{t.get('raqam', '')}-topshiriq: {t.get('mavzu', '')}", size=13)
        _para(doc, t.get('topshiriq_matni', ''))
        if t.get('hajm'):
            _para(doc, f"Hajm: {t['hajm']}")
        if t.get('baholash_mezoni'):
            _para(doc, f"Baholash mezoni: {t['baholash_mezoni']}")
        for a in t.get('adabiyotlar', []):
            _para(doc, a, bullet=True)
    if content.get('umumiy_baholash_rubrikasi'):
        _heading(doc, 'Umumiy baholash rubrikasi')
        _table(doc, ['Mezon', 'Ball'],
               [(r.get('mezon', ''), r.get('ball', '')) for r in content['umumiy_baholash_rubrikasi']])


def _render_baholash_mezonlari(doc, content):
    if content.get('umumiy_tamoyillar'):
        _heading(doc, 'Umumiy tamoyillar')
        _para(doc, content['umumiy_tamoyillar'])
    for bt in content.get('baholash_turlari', []):
        _heading(doc, f"{bt.get('tur', '')} (maks. {bt.get('maksimal_ball', '')} ball)", size=13)
        _table(doc, ['Mezon', 'Ball', 'Tavsif'],
               [(m.get('mezon', ''), m.get('ball', ''), m.get('tavsif', ''))
                for m in bt.get('mezonlar', [])])
    if content.get('baholash_shkalasi'):
        _heading(doc, 'Baholash shkalasi')
        _table(doc, ['Daraja', "Ball oralig'i", 'Tavsif'],
               [(s.get('daraja', ''), s.get("ball_oralig'i", s.get('ball_oraligi', '')), s.get('tavsif', ''))
                for s in content['baholash_shkalasi']])


def _render_test_savollari(doc, content):
    for q in content.get('savollar', []):
        _heading(doc, f"{q.get('raqam', '')}. {q.get('savol', '')}", size=12)
        variantlar = q.get('variantlar', {})
        togri = q.get('togri_javob', '')
        for letter in ('A', 'B', 'C', 'D', 'E'):
            val = variantlar.get(letter, '')
            if not val:
                continue
            mark = ' ✓' if letter == togri else ''
            _para(doc, f"{letter}) {val}{mark}", bullet=True)


def _render_imtihon_savollari(doc, content):
    for v in content.get('variantlar', []):
        _heading(doc, f"{v.get('variant', '')}-variant", size=13)
        for s in v.get('savollar', []):
            _para(doc, f"{s.get('raqam', '')}. {s.get('savol', '')}")
            if s.get('javob_namunasi'):
                _para(doc, f"Javob namunasi: {s['javob_namunasi']}")
        doc.add_page_break()


def _render_kurs_ishi_mavzulari(doc, content):
    _table(doc, ['№', 'Mavzu', 'Tavsif', 'Tavsiya etiladigan adabiyot yo\'nalishi'],
           [(m.get('raqam', ''), m.get('mavzu', ''), m.get('tavsif', ''),
             m.get('tavsiya_adabiyot_yonalishi', ''))
            for m in content.get('mavzular', [])])


_RENDERERS = {
    'sillabus': _render_sillabus_like,
    'ishchi-dastur': _render_sillabus_like,
    'fan-dasturi': _render_sillabus_like,
    'uslubiy-qollanma': _render_uslubiy_qollanma,
    'mustaqil-talim': _render_mustaqil_talim,
    'baholash-mezonlari': _render_baholash_mezonlari,
    'test-savollari': _render_test_savollari,
    'imtihon-savollari': _render_imtihon_savollari,
    'kurs-ishi-mavzulari': _render_kurs_ishi_mavzulari,
}


def build_docx(doc_type, content, metadata):
    """content — AI JSON (dict), metadata — foydalanuvchi kiritgan forma qiymatlari
    (flat — barcha bo'lim maydonlari bitta darajada, masalan 'universitet', 'kredit', ...).
    Qaytaradi: io.BytesIO (fayl kontenti)."""
    metadata = metadata or {}
    cfg = DOCUMENT_TYPES[doc_type]
    doc = _setup_document()
    _add_header_block(doc, metadata, doc_type)
    _add_title(doc, metadata, cfg['label'])
    _add_info_table(doc, metadata, doc_type)
    renderer = _RENDERERS.get(doc_type, _render_sillabus_like)
    renderer(doc, content or {})
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def save_docx(doc_type, content, metadata, out_path=None):
    """DOCX ni diskka yozadi (static/generated_docs/<uuid>.docx, yoki `out_path`
    berilgan bo'lsa o'sha yerga — namuna generatsiyasi uchun) va yo'lini qaytaradi."""
    import os
    buf = build_docx(doc_type, content, metadata)
    if out_path:
        path = out_path
        os.makedirs(os.path.dirname(path), exist_ok=True)
    else:
        os.makedirs(GENERATED_DOCS_DIR, exist_ok=True)
        path = os.path.join(GENERATED_DOCS_DIR, f"{uuid.uuid4()}.docx")
    with open(path, 'wb') as f:
        f.write(buf.getvalue())
    return path
